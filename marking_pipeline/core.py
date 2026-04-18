import io
import hashlib
import json
import os
import re
import time
import zipfile
from collections import defaultdict
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import pdfplumber
import pymupdf
import requests
from docx import Document
try:
    from docx2python import docx2python
except ImportError:  # pragma: no cover - optional dependency
    docx2python = None

# `pymupdf4llm` is optional, but importing it eagerly has proven unstable in this
# environment. Keep PDF structure extraction on the safer fallback path unless a
# future change reintroduces a guarded lazy import.
pymupdf4llm = None


ROOT_DIR = Path(__file__).resolve().parents[1]
DEFAULT_OLLAMA_URL = "http://localhost:11434/api/chat"
SUPPORTED_TEXT_SUFFIXES = {".pdf", ".docx", ".txt"}
DEFAULT_DETECTOR_MODEL_NAME = "mistral:7b"
DEFAULT_OLLAMA_TIMEOUT_SECONDS = 300

_STRUCTURE_CUE_PATTERN = re.compile(
    r"(?im)^\s*(?:part|question|section|task|item)\s+[a-z0-9ivx]+(?:\s*q\d+)?(?:\b|\s*[:.)\]-])"
)

DEFAULT_CONTEXT_PATTERNS = {
    "rubric": ("rubric",),
    "brief": ("brief", "assignment"),
    "marking_scheme": ("marking_scheme", "marking scheme", "scheme"),
    "graded_sample": ("graded_sample", "graded sample", "sample", "example"),
    "other": ("support", "context", "guidance", "instruction"),
}


@dataclass(frozen=True)
class MarkingContext:
    rubric_text: str
    brief_text: str
    marking_scheme_text: str
    graded_sample_text: str
    other_context_text: str
    max_mark: float


@dataclass(frozen=True)
class AssessmentBundle:
    name: str
    folder: Path
    submission_files: tuple[Path, ...]
    rubric_files: tuple[Path, ...]
    brief_files: tuple[Path, ...]
    marking_scheme_files: tuple[Path, ...]
    graded_sample_files: tuple[Path, ...]
    other_files: tuple[Path, ...]


@dataclass(frozen=True)
class SubmissionPart:
    label: str
    focus_hint: str = ""
    anchor_text: str = ""
    section_text: str = ""
    max_mark: float | None = None
    marking_guidance: str = ""
    question_text_exact: str = ""
    task_type: str = ""
    criterion_mode: str = ""


@dataclass(frozen=True)
class SubmissionDiagnostics:
    extracted_word_count: int
    detected_part_count: int
    detected_part_labels: tuple[str, ...]
    low_text: bool
    possible_extraction_issue: bool


@dataclass(frozen=True)
class AssessmentStructureSection:
    label: str
    parent_label: str = ""
    max_mark: float | None = None
    weight_text: str = ""
    question_text_exact: str = ""
    marking_instructions_exact: str = ""
    anchor_phrases: tuple[str, ...] = ()
    evidence_expectations: tuple[str, ...] = ()
    dependency_labels: tuple[str, ...] = ()
    children: tuple["AssessmentStructureSection", ...] = ()


@dataclass(frozen=True)
class AssessmentUnit:
    label: str
    max_mark: float | None = None
    parent_label: str = ""
    grading_mode: str = "deterministic"
    task_type: str = ""
    criterion_mode: str = ""
    dependency_group: str = ""
    question_text_exact: str = ""
    marking_guidance: str = ""
    rubric_text: str = ""
    rubric_confidence_0_to_100: float | None = None
    rubric_issues: tuple[str, ...] = ()


@dataclass(frozen=True)
class AssessmentMap:
    units: tuple[AssessmentUnit, ...]
    overall_max_mark: float
    scale_confidence_0_to_100: float


@dataclass(frozen=True)
class PreparedAssessmentMap:
    cache_key: str
    original_map: AssessmentMap
    prepared_map: AssessmentMap
    prepared_units: tuple["PreparedAssessmentUnit", ...] = ()
    verifier_model_name: str | None = None
    verification_applied: bool = False
    verifier_attempt_count: int = 0
    preparation_issues: tuple[str, ...] = ()
    preparation_confidence_0_to_100: float | None = None
    preparation_latency_seconds: float = 0.0


@dataclass(frozen=True)
class PreparedAssessmentUnit:
    label: str
    compact_criteria: tuple[str, ...] = ()
    verifier_confidence_0_to_100: float | None = None
    verifier_issues: tuple[str, ...] = ()
    verifier_attempt_count: int = 0
    verifier_accepted_refinements: tuple[str, ...] = ()
    verifier_rejected_refinements: tuple[str, ...] = ()


_PREPARED_ASSESSMENT_CACHE: dict[str, PreparedAssessmentMap] = {}
MAX_PREPARED_MAP_VERIFIER_ATTEMPTS = 3
MIN_ACCEPTABLE_PREPARATION_CONFIDENCE = 75.0


def decode_text_bytes(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode the text file with a supported encoding.")


def extract_text_from_pdf(file_obj: Any) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_obj: Any) -> str:
    raw = file_obj.read()
    if hasattr(file_obj, "seek"):
        file_obj.seek(0)
    xml_text = _extract_text_from_docx_xml(raw)
    if xml_text.strip():
        return xml_text.strip()
    doc = Document(io.BytesIO(raw))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def extract_docx_structure_text(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(f"{_docx_heading_marker(paragraph)}{text}")
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _docx_heading_marker(paragraph: Any) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    normalized = re.sub(r"\s+", " ", text)
    if len(normalized) > 140:
        return ""
    text_runs = [run for run in paragraph.runs if run.text.strip()]
    if not text_runs:
        return ""
    bold_runs = sum(1 for run in text_runs if bool(run.bold))
    size_values = [run.font.size.pt for run in text_runs if run.font.size is not None]
    max_size = max(size_values) if size_values else None
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
    heading_like_text = bool(re.match(r"^(part|question|section|appendix)\b", normalized, re.IGNORECASE))
    compact_question_like = bool(re.match(r"^(q(?:uestion)?\s*\d+[a-z]?|[ivxlcdm]+[\.\)]|(?:\d+|[A-Z])[\.\)])\b", normalized, re.IGNORECASE))
    title_case_like = bool(re.match(r"^[A-Z][A-Za-z0-9 ,:/&()'\"-]{0,90}$", normalized))
    all_caps_like = normalized.isupper() and len(normalized.split()) <= 12
    heading_signal = (
        heading_like_text
        or "heading" in style_name
        or (compact_question_like and len(normalized.split()) <= 12)
        or (title_case_like and (bold_runs == len(text_runs) or max_size is not None and max_size >= 13))
        or all_caps_like
        or (bold_runs == len(text_runs) and len(normalized.split()) <= 16)
        or (max_size is not None and max_size >= 14 and len(normalized.split()) <= 20)
    )
    return "[HEADING] " if heading_signal else ""


def choose_docx_scoring_text(primary_text: str, fallback_text: str) -> str:
    primary = primary_text.strip()
    fallback = fallback_text.strip()
    if not primary:
        return fallback
    latex_hits = primary.count("<latex>")
    media_hits = primary.count("----media/")
    if latex_hits > 0 and media_hits <= max(2, latex_hits // 8):
        return primary
    if media_hits > 0 and latex_hits == 0:
        return fallback or primary
    return primary or fallback


def extract_docx_scoring_text(raw: bytes) -> str:
    fallback_text = extract_text_from_docx(io.BytesIO(raw))
    if docx2python is None:
        return fallback_text
    with docx2python(io.BytesIO(raw)) as result:
        text = (result.text or "").strip()
    return choose_docx_scoring_text(text, fallback_text)


def read_file_bytes(name: str, raw: bytes) -> str:
    lowered = name.lower()
    if lowered.endswith(".txt"):
        return decode_text_bytes(raw).strip()
    if lowered.endswith(".pdf"):
        return extract_text_from_pdf(io.BytesIO(raw))
    if lowered.endswith(".docx"):
        return extract_text_from_docx(io.BytesIO(raw))
    raise ValueError(f"Unsupported file type for {name}.")


def read_uploaded_files_text(uploaded_files: list[Any] | None) -> str:
    if not uploaded_files:
        return ""
    texts = [read_file_bytes(uploaded_file.name, uploaded_file.getvalue()) for uploaded_file in uploaded_files]
    return "\n\n".join(text for text in texts if text).strip()


def read_path_text(path: Path) -> str:
    return read_file_bytes(path.name, path.read_bytes())


def cleanup_pymupdf4llm_text(text: str) -> str:
    cleaned_lines: list[str] = []
    for line in text.replace("\f", "\n").splitlines():
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if "picture [" in stripped and "intentionally omitted" in stripped:
            continue
        if stripped.isdigit():
            continue
        stripped = re.sub(r"^#{1,6}\s*", "", stripped)
        stripped = stripped.replace("**", "").replace("__", "")
        stripped = stripped.replace("`", "")
        stripped = stripped.replace("<br>", " ")
        if "|" in stripped and stripped.count("|") >= 2:
            cells = [cell.strip() for cell in stripped.split("|") if cell.strip()]
            stripped = " ".join(cells)
        stripped = re.sub(r"\s+", " ", stripped).strip()
        cleaned_lines.append(stripped)

    cleaned = "\n".join(cleaned_lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def extract_pdf_structure_text(raw: bytes) -> str:
    if pymupdf4llm is None:
        return ""
    doc = pymupdf.open(stream=raw, filetype="pdf")
    try:
        extracted = pymupdf4llm.to_markdown(doc, header=False, footer=False)
    finally:
        doc.close()
    return cleanup_pymupdf4llm_text(extracted)


def build_submission_texts_from_path(path: Path) -> tuple[str, str | None]:
    if path.suffix.lower() == ".docx":
        raw = path.read_bytes()
        script_text = extract_docx_scoring_text(raw)
        structure_text = extract_docx_structure_text(raw) or None
        return script_text, structure_text

    structure_text: str | None = None
    if path.suffix.lower() == ".pdf":
        raw = path.read_bytes()
        pdf_text = extract_pdf_structure_text(raw)
        script_text = pdf_text if pdf_text else read_path_text(path)
    else:
        script_text = read_path_text(path)
    return script_text, structure_text


def build_submission_texts_from_upload(uploaded_file: Any) -> tuple[str, str | None]:
    raw = uploaded_file.getvalue()
    lowered = uploaded_file.name.lower()
    if lowered.endswith(".docx"):
        script_text = extract_docx_scoring_text(raw)
        structure_text = extract_docx_structure_text(raw) or None
        return script_text, structure_text

    structure_text: str | None = None
    if lowered.endswith(".pdf"):
        pdf_text = extract_pdf_structure_text(raw)
        script_text = pdf_text if pdf_text else read_file_bytes(uploaded_file.name, raw)
    else:
        script_text = read_file_bytes(uploaded_file.name, raw)
    return script_text, structure_text


def read_paths_text(paths: list[Path] | tuple[Path, ...]) -> str:
    if not paths:
        return ""
    return "\n\n".join(text for text in (read_path_text(path) for path in paths) if text).strip()


def infer_max_mark_from_texts(*texts: str) -> float | None:
    patterns = (
        r"maximum\s+mark\s*(?:is|=|:)?\s*(\d+(?:\.\d+)?)",
        r"max(?:imum)?\s+mark\s*(?:is|=|:)?\s*(\d+(?:\.\d+)?)",
        r"single\s+overall\s+mark\s+out\s+of\s*(\d+(?:\.\d+)?)",
        r"mark\s+out\s+of\s*(\d+(?:\.\d+)?)",
        r"out\s+of\s*(\d+(?:\.\d+)?)",
        r"total\s+(?:mark|marks)\s*(?:is|=|:)?\s*(\d+(?:\.\d+)?)",
    )

    candidates: set[float] = set()
    for text in texts:
        if not text:
            continue
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                candidates.add(float(match.group(1)))

    if not candidates:
        return None
    if len(candidates) > 1:
        raise ValueError(
            "Found conflicting maximum marks in the uploaded marking documents: "
            + ", ".join(_format_number(value) for value in sorted(candidates))
        )
    return next(iter(candidates))


def prepare_marking_context(
    rubric_text: str,
    brief_text: str,
    marking_scheme_text: str,
    graded_sample_text: str,
    other_context_text: str,
) -> MarkingContext:
    rubric_text = rubric_text.strip()
    brief_text = brief_text.strip()
    marking_scheme_text = marking_scheme_text.strip()
    graded_sample_text = graded_sample_text.strip()
    other_context_text = other_context_text.strip()

    if not (rubric_text or marking_scheme_text):
        raise ValueError("Provide a rubric or marking scheme before grading.")

    max_mark = infer_max_mark_from_texts(rubric_text, marking_scheme_text, brief_text)
    if max_mark is None:
        raise ValueError(
            "Could not infer the maximum mark from the rubric or marking scheme. "
            "Include an explicit phrase such as 'Maximum mark: 20' or 'out of 85'."
        )

    return MarkingContext(
        rubric_text=rubric_text,
        brief_text=brief_text,
        marking_scheme_text=marking_scheme_text,
        graded_sample_text=graded_sample_text,
        other_context_text=other_context_text,
        max_mark=max_mark,
    )


def build_messages(script_text: str, context: MarkingContext, filename: str) -> list[dict[str, str]]:
    max_mark_text = _format_number(context.max_mark)

    context_parts = []
    if context.rubric_text:
        context_parts.append(f"RUBRIC:\n{context.rubric_text}")
    if context.brief_text:
        context_parts.append(f"ASSIGNMENT BRIEF:\n{context.brief_text}")
    if context.marking_scheme_text:
        context_parts.append(f"MARKING SCHEME:\n{context.marking_scheme_text}")
    if context.graded_sample_text:
        context_parts.append(
            "EXAMPLE GRADED SCRIPT (for tone and structure only):\n"
            f"{context.graded_sample_text}"
        )
    if context.other_context_text:
        context_parts.append(f"OTHER SUPPORTING DOCUMENTS:\n{context.other_context_text}")

    system = (
        "You are a fair and consistent university examiner. "
        "Use only the supplied marking documents and the student's submission. "
        "Do not invent missing criteria or assume a different marking scale."
    )

    user = (
        "You are marking a student's script.\n\n"
        f"{'\n\n'.join(context_parts)}\n\n"
        f"STUDENT FILE NAME: {filename}\n\n"
        "STUDENT ANSWER:\n"
        f"\"\"\"{script_text}\"\"\"\n\n"
        "Return only one JSON object with exactly these keys:\n"
        '- "total_mark": number\n'
        '- "max_mark": number\n'
        '- "overall_feedback": string\n\n'
        f"Rules:\n"
        f"1. Use max_mark = {max_mark_text}.\n"
        f"2. total_mark must be between 0 and {max_mark_text}.\n"
        "3. overall_feedback must be at least 120 words.\n"
        "4. overall_feedback must include at least two concrete strengths and two concrete weaknesses.\n"
        "5. If the assessment has multiple questions, explicitly mention each question.\n"
        "6. Do not include markdown fences or any text outside the JSON object."
    )

    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]


def build_part_messages(
    part: SubmissionPart,
    context: MarkingContext,
    filename: str,
    model_name: str | None = None,
) -> list[dict[str, str]]:
    max_mark_text = _format_number(part.max_mark) if part.max_mark is not None else "100"
    score_field = "provisional_score"
    score_rule = f"{score_field} must be between 0 and {max_mark_text}."
    if part.max_mark is None:
        score_field = "provisional_score_0_to_100"
        score_rule = f"{score_field} must be between 0 and 100."
    task_text = _build_part_task_text(part)
    question_text = _question_text_for_model(
        part.question_text_exact or part.focus_hint or part.label,
        max_chars=_question_text_limit_for_part(part),
    )
    task_goal = _task_goal_for_model(task_text, question_text)
    required_steps = _extract_required_steps_for_model(part)
    criteria = _build_prompt_criteria(part)
    scale_type, scale_labels = _criterion_scale_for_part(part)
    payload_criteria = [
        {
            "criterion_name": criterion["criterion_name"],
            "check": criterion["check"],
        }
        for criterion in criteria
    ]
    scoring_rules = _extract_prompt_scoring_rules(part.marking_guidance, part.max_mark)
    question_payload: dict[str, Any] = {
        "text": question_text,
        "task_type": part.task_type or "unspecified",
    }
    if task_goal:
        question_payload["task_goal"] = task_goal
    if required_steps:
        question_payload["required_steps"] = required_steps
    payload = {
        "section_id": part.label,
        "max_mark": part.max_mark if part.max_mark is not None else 100,
        "question": question_payload,
        "scoring": {
            "method": _part_scoring_method(part),
            "scale": {
                "type": scale_type,
                "labels": list(scale_labels),
            },
            "criteria": payload_criteria,
        },
        "student_response": {
            "text": part.section_text or part.focus_hint or part.label,
        },
    }
    if _use_qwen_scoring_scaffold(model_name):
        # Qwen-specific scoring scaffold: local probes showed that Qwen-family models
        # are more reliable when scope, decision, and typing rules are separated into
        # short local blocks instead of being implied by repeated prose.
        payload["scoring"]["rules"] = scoring_rules
        payload["scoring"]["scope_block"] = _build_qwen_scope_block()
        payload["scoring"]["decision_block"] = _build_qwen_decision_block(part, score_rule, scale_labels)
        payload["scoring"]["output_block"] = _build_qwen_output_block(score_field, scale_labels)
        system = (
            "Score only this section from the payload. Use no outside information. "
            "Follow the scope_block, decision_block, and output_block exactly. Return one JSON object only."
        )
    else:
        payload["scoring"]["rules"] = [
            score_rule,
            "Use only this section.",
            *scoring_rules,
        ]
        payload["output_contract"] = [
            f"Return keys: section_label, {score_field}, criterion_notes, strengths, weaknesses, evidence, coverage_comment.",
            "criterion_notes: list of {criterion_name, status, note}.",
            f"status values: {', '.join(scale_labels)}.",
            "note: 2 to 5 words; do not copy criterion text.",
            "strengths and weaknesses: lists with at least 2 items each.",
            "evidence: list with at least 1 item.",
            "coverage_comment: required.",
        ]
        system = (
            "Grade only this section from the payload. Use no outside information. Return one JSON object only."
        )
    user = (
        "scoring_payload:\n"
        f"{json.dumps(payload, ensure_ascii=True, indent=2)}\n"
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def _use_qwen_scoring_scaffold(model_name: str | None) -> bool:
    return "qwen" in (model_name or "").strip().lower()


def _build_qwen_scope_block() -> list[str]:
    return [
        "Score only this section.",
        "Use only the question, criteria, scoring rules, and student response in this payload.",
        "Do not import requirements from other questions or from outside knowledge.",
    ]


def _build_qwen_decision_block(
    part: SubmissionPart,
    score_rule: str,
    scale_labels: tuple[str, ...],
) -> list[str]:
    mode = _criterion_mode_for_part(part)
    top_label = scale_labels[-1]

    # Evidence-aware support rule: for derivations, focus on visible method/setup;
    # for other tasks, flag unsupported claims
    if part.task_type == "deterministic_derivation":
        support_rule = (
            "Award partial credit for visible correct setup, valid intermediate steps, and correct interpretation "
            "even if the final result is incomplete or hard to verify in the extracted text. "
            "Award zero or minimal credit only if the response shows no relevant method or no visible attempt."
        )
    elif mode == "abstract":
        support_rule = "Do not award high marks to broad but weakly supported discussion."
    else:
        support_rule = "Do not award marks for unsupported steps, missing components, or incorrect claims."

    return [
        f"Task mode: {mode}.",
        f"For each criterion, choose one status only from: {', '.join(scale_labels)}.",
        "First assign the criterion statuses. Then assign the numeric score.",
        f"If a required element is missing, do not use {top_label} for that criterion and do not award a top score.",
        support_rule,
        score_rule,
        "Make the numeric score match the criterion statuses and the visible evidence.",
    ]


def _build_qwen_output_block(score_field: str, scale_labels: tuple[str, ...]) -> list[str]:
    return [
        f"Return one JSON object with exactly these keys: section_label, {score_field}, criterion_notes, strengths, weaknesses, evidence, coverage_comment.",
        "criterion_notes must be a list of {criterion_name, status, note}.",
        f"criterion_notes.status must be one of: {', '.join(scale_labels)}.",
        "criterion_notes.note must be 2 to 5 words and must not copy the criterion text.",
        "strengths and weaknesses must each contain at least 2 items.",
        "evidence must contain at least 1 item.",
        "coverage_comment is required.",
    ]


def _part_scoring_method(part: SubmissionPart) -> str:
    if part.task_type in {
        "critique_and_revision",
        "evaluative_discussion",
        "synthesis_across_sources",
        "whole_submission_holistic",
    }:
        return "classification_then_score"
    return "criterion_evaluation"


def _question_text_limit_for_part(part: SubmissionPart) -> int:
    if _part_scoring_method(part) == "criterion_evaluation":
        return 1800
    return 650


def _build_prompt_criteria(part: SubmissionPart) -> list[dict[str, str]]:
    criteria = _extract_prompt_criteria(part.marking_guidance)
    if not criteria:
        criteria = _default_prompt_criteria(part)
    criteria = _enrich_prompt_criteria(part, criteria)
    payload: list[dict[str, str]] = []
    scale_type, scale_labels = _criterion_scale_for_part(part)
    for index, criterion in enumerate(criteria, start=1):
        normalized_criterion = _normalize_prompt_criterion_text(part, criterion)
        weak_anchor, strong_anchor = _criterion_anchors(normalized_criterion)
        payload.append(
            {
                "criterion_name": f"criterion_{index}",
                "check": normalized_criterion,
                "scale_type": scale_type,
                "scale_labels": ", ".join(scale_labels),
                "weak_anchor": weak_anchor,
                "strong_anchor": strong_anchor,
            }
        )
    return payload


def _enrich_prompt_criteria(part: SubmissionPart, criteria: list[str]) -> list[str]:
    normalized: set[str] = set()
    enriched: list[str] = []
    for criterion in _construct_criteria_for_task(
        part.task_type,
        part.question_text_exact,
        part.marking_guidance,
    ):
        cleaned = _normalize_prompt_criterion_text(part, str(criterion))
        key = cleaned.lower().rstrip(".")
        if not cleaned or key in normalized:
            continue
        normalized.add(key)
        enriched.append(cleaned)
        if len(enriched) >= 5:
            return enriched

    for criterion in criteria:
        cleaned = _normalize_prompt_criterion_text(part, str(criterion))
        key = cleaned.lower().rstrip(".")
        if not cleaned or key in normalized:
            continue
        normalized.add(key)
        enriched.append(cleaned)
        if len(enriched) >= 5:
            return enriched

    for criterion in _supplemental_prompt_criteria(part):
        cleaned = _normalize_prompt_criterion_text(part, str(criterion))
        key = cleaned.lower().rstrip(".")
        if not cleaned or key in normalized:
            continue
        normalized.add(key)
        enriched.append(cleaned)
        if len(enriched) >= 5:
            break
    return enriched


def _criterion_mode_for_part(part: SubmissionPart) -> str:
    if part.criterion_mode in {"deterministic", "abstract"}:
        return part.criterion_mode
    if _part_scoring_method(part) == "classification_then_score":
        return "abstract"
    if part.task_type in {"critique_and_revision", "evaluative_discussion", "synthesis_across_sources", "whole_submission_holistic"}:
        return "abstract"
    return "deterministic"


def _criterion_scale_for_part(part: SubmissionPart) -> tuple[str, tuple[str, ...]]:
    if _criterion_mode_for_part(part) == "abstract":
        return "likert_judgement", ("weak", "developing", "clear", "strong")
    return "verification_check", ("missing_or_wrong", "partial", "secure")


def _normalize_prompt_criterion_text(part: SubmissionPart, criterion: str) -> str:
    cleaned = " ".join(str(criterion).split()).strip()
    if not cleaned:
        return ""
    mode = _criterion_mode_for_part(part)
    lowered = cleaned.lower()
    if mode == "abstract":
        starters = (
            "evaluate whether ",
            "assess whether ",
            "determine whether ",
            "identify whether ",
            "check whether ",
            "check that ",
            "confirm that ",
        )
        for prefix in starters:
            if lowered.startswith(prefix):
                cleaned = "Judge how well " + cleaned[len(prefix):]
                break
        lowered = cleaned.lower()
        if lowered.startswith("evaluate how "):
            cleaned = "Judge how well " + cleaned[len("evaluate how "):]
        elif lowered.startswith("assess how "):
            cleaned = "Judge how well " + cleaned[len("assess how "):]
        elif lowered.startswith("determine how "):
            cleaned = "Judge how well " + cleaned[len("determine how "):]
    else:
        starters = (
            ("evaluate whether ", "Check that "),
            ("assess whether ", "Check that "),
            ("determine whether ", "Check that "),
            ("identify whether ", "Check that "),
            ("check whether ", "Check that "),
            ("check that ", "Check that "),
            ("confirm that ", "Confirm that "),
        )
        for prefix, replacement in starters:
            if lowered.startswith(prefix):
                cleaned = replacement + cleaned[len(prefix):]
                break
    cleaned = cleaned[0].upper() + cleaned[1:] if cleaned else cleaned
    return cleaned.rstrip(".") + "."


def _supplemental_prompt_criteria(part: SubmissionPart) -> list[str]:
    if _criterion_mode_for_part(part) == "abstract":
        common_structured = [
            "Judge how completely the answer covers the required steps or components.",
            "Judge how well the section supports its score with relevant reasoning or evidence.",
            "Judge how tightly the response stays focused on the requested task.",
            "Judge how clearly the most important point is made.",
        ]
    else:
        common_structured = [
            "Confirm that the required steps or components are present without material gaps.",
            "Confirm that the stated result is supported by enough method, working, or explanation to justify credit.",
            "Confirm that the response stays focused on the requested task instead of drifting into loosely related material.",
            "Confirm that the most important point is stated clearly rather than left implicit or buried.",
        ]
    task_type = part.task_type or ""
    if task_type == "critique_and_revision":
        return [
            "Judge how effectively the response prioritizes the most important errors or omissions.",
            "Judge how well the revision stays aligned with the original task after the corrections are made.",
            *common_structured,
        ]
    if task_type in {"evaluative_discussion", "synthesis_across_sources"}:
        return [
            "Judge how effectively the response prioritizes the most relevant points instead of lower-value material.",
            "Judge how proportionate the conclusion is to the support given in the section.",
            "Judge how clearly the argument maintains a line of judgement.",
            *common_structured,
        ]
    if task_type in {"measurement_and_data_design", "causal_identification", "regression_specification"}:
        return [
            "Confirm that the answer covers the required components without leaving a material gap in the setup.",
            "Confirm that the proposed approach fits the stated question rather than a nearby but different one.",
            "Confirm that key terms, variables, or assumptions are explicit enough to be usable.",
            *common_structured,
        ]
    if task_type == "deterministic_derivation":
        targeted: list[str] = []
        question_text = part.question_text_exact.lower()
        if "inequalit" in question_text or "condition" in question_text:
            targeted.append(
                "Confirm that the answer states the required inequalities or parameter conditions correctly."
            )
        if any(phrase in question_text for phrase in ("numerical values", "example", "give an example")):
            targeted.append(
                "Confirm that any numerical example is valid and actually satisfies the required conditions."
            )
        if any(
            phrase in question_text
            for phrase in (
                "for each family",
                "other families are doing",
                "change its decision",
            )
        ):
            targeted.append(
                "Confirm that the derivation checks each required case or family decision rather than skipping a decisive step."
            )
        return [*targeted, *common_structured]
    if _part_scoring_method(part) == "classification_then_score":
        return [
            "Judge how effectively the response prioritizes the most relevant points instead of lower-value material.",
            "Judge how proportionate the conclusion is to the support given in the section.",
            "Judge how clearly the argument maintains a line of judgement.",
            *common_structured,
        ]
    return common_structured


def _merge_unique_criteria(*groups: list[str]) -> list[str]:
    merged: list[str] = []
    seen: set[str] = set()
    for group in groups:
        for item in group:
            cleaned = " ".join(str(item).split()).strip()
            key = cleaned.lower().rstrip(".")
            if not cleaned or key in seen:
                continue
            seen.add(key)
            merged.append(cleaned.rstrip(".") + ".")
    return merged


def _extract_parenthesized_goal_items(question_text: str) -> list[str]:
    items: list[str] = []
    for match in re.finditer(r"\((\d+)\)\s*([^()]+?)(?=(?:\s*\(\d+\)|$))", question_text):
        item = " ".join(match.group(2).split()).strip(" ;,.")
        if item:
            items.append(item)
    return items


def _extract_task_constructs(task_type: str, question_text: str, marking_guidance: str) -> list[str]:
    text = " ".join(question_text.split())
    lowered = text.lower()
    guidance = " ".join(marking_guidance.split()).lower()
    constructs: list[str] = []

    def add(item: str) -> None:
        cleaned = " ".join(item.split()).strip(" ;,.")
        if cleaned:
            constructs.append(cleaned)

    parenthesized_items = _extract_parenthesized_goal_items(text)
    if parenthesized_items and any(term in lowered for term in ("your answer should contain", "answer should contain", "note:")):
        for item in parenthesized_items[:4]:
            add(f"includes {item}")
    elif any(term in lowered for term in ("prompt", "output", "rewrite", "rewritten", "own tone", "own words")):
        if "prompt" in lowered:
            add("includes the original prompt")
        if "output" in lowered:
            add("includes the model output being audited")
        if any(term in lowered for term in ("mistake", "inaccurac", "comment", "audit")):
            add("identifies and explains substantive inaccuracies or weaknesses in the output")
        if any(term in lowered for term in ("rewrite", "rewritten", "own tone", "own words", "edit the output")):
            add("provides a corrected rewrite in the student's own tone")

    if task_type == "critique_and_revision":
        if any(term in lowered for term in ("quote", "quotes", "page number", "page numbers", "report", "lecture", "course material", "reference")):
            add("supports the critique and corrections with specific source evidence or page references")
    elif task_type == "measurement_and_data_design":
        if "for each prediction" in lowered:
            add("covers each prediction from the theory rather than only a subset")
        add("maps each theoretical concept to an observable variable or indicator")
        if any(term in lowered for term in ("data source", "data sources", "dataset", "source of data", "link to the source", "find some data sources")):
            add("proposes suitable data sources for the measures being used")
        if "link" in lowered and "source" in lowered:
            add("provides a usable source link for the proposed data")
        if any(term in lowered for term in ("which variable", "combination of variables", "dataset")):
            add("identifies which dataset variable or variable combination would be used")
    elif task_type == "regression_specification":
        if any(term in lowered for term in ("choose one of the theory", "prediction from your answer")):
            add("matches the regression specification to one stated theory prediction")
        if any(term in lowered for term in ("data you identified", "question 2", "data identified in question 2")):
            add("uses variables that are consistent with the measures or data identified earlier")
        if any(term in lowered for term in ("explain carefully each term", "each term in your equation", "define each term")):
            add("defines each term in the equation clearly")
        if any(term in lowered for term in ("coefficient of interest", "main coefficient")):
            add("identifies the main coefficient of interest")
        if "hypothesis" in lowered:
            add("states the hypothesis that would test the chosen prediction")
    elif task_type == "welfare_reasoning":
        if any(term in lowered for term in ("increase social welfare", "decrease social welfare", "under which conditions")):
            add("states the conditions under which VAT increases social welfare and the conditions under which it decreases social welfare")
        if any(term in lowered for term in ("utility of each family", "each family changes")):
            add("shows how the utility of each relevant family changes when VAT increases")
        add("distinguishes welfare effects across the relevant families or groups")
    elif task_type == "synthesis_across_sources":
        if any(term in lowered for term in ("main claims", "claims in the output")):
            add("identifies the main claims being assessed across the relevant materials")
        if any(term in lowered for term in ("support, qualify, or contradict", "support qualify", "contradict the claims")):
            add("shows whether the relevant sources support, qualify, or contradict the claims")
        if any(term in lowered for term in ("quote", "quotes", "page number", "page numbers")):
            add("uses specific quotations or page references where the task asks for them")
        if any(term in lowered for term in ("video", "videos", "youtube")):
            add("covers the arguments raised in the relevant source materials rather than only one source")
        if "not included in the model" in lowered or "not included in the framework" in lowered:
            add("identifies arguments or factors raised in the sources that the model does not capture")
        if any(term in lowered for term in ("change the model", "extend the model", "analyse these points")):
            add("proposes concrete changes or extensions to the model to analyse the omitted points")
    elif task_type == "evaluative_discussion":
        if any(term in lowered for term in ("good idea", "whether you think")):
            add("reaches a clear and justified overall judgement")
        if any(term in lowered for term in ("use the results from the analysis", "lecture", "course concepts", "concept studied")):
            add("uses the relevant model results, evidence, or course concepts to justify the judgement")
        if any(term in lowered for term in ("which groups would benefit", "lose more", "costs and benefits")):
            add("distinguishes who gains, who loses, and the main tradeoffs")
        if any(term in lowered for term in ("prompt", "output")):
            add("includes the required AI prompt and generated output")
        if "main changes" in lowered:
            add("covers the main changes introduced by the policy")
    elif task_type == "deterministic_derivation":
        if "condition" in lowered or "inequalit" in lowered:
            add("states the required conditions or inequalities correctly")
        if any(term in lowered for term in ("example", "numerical values")):
            add("provides a valid numerical example when the task asks for one")
        if any(term in lowered for term in ("for each family", "other families are doing", "change its decision")):
            add("checks each required case or family decision rather than skipping a decisive step")
    elif task_type == "comparative_statics":
        if any(term in lowered for term in ("each parameter", "changing each")):
            add("covers the effect of each relevant parameter rather than only a subset")
        if any(term in lowered for term in ("more likely", "scenario", "outcome")):
            add("links each parameter change to the stated scenario or outcome comparison")

    if "parent marking instructions:" in guidance and any(term in guidance for term in ("quote", "page number", "reference", "source evidence")):
        add("uses the source evidence required by the parent marking instructions")

    normalized: list[str] = []
    seen: set[str] = set()
    for item in constructs:
        key = item.lower()
        if key not in seen:
            normalized.append(item)
            seen.add(key)
    return normalized[:5]


def _construct_criteria_for_task(task_type: str, question_text: str, marking_guidance: str) -> list[str]:
    criteria: list[str] = []
    for construct in _extract_task_constructs(task_type, question_text, marking_guidance):
        criteria.append(f"Evaluate whether the response {construct}.")
    return criteria[:5]


def _extract_required_steps_for_model(part: SubmissionPart) -> list[str]:
    # Use the full cleaned question here so late instructions still become explicit
    # required steps even when the prompt carries a shorter display version.
    question_text = _question_text_for_model(part.question_text_exact, max_chars=None)
    if not question_text:
        return []
    lowered = question_text.lower()
    steps: list[str] = []
    if part.task_type == "deterministic_derivation":
        if "inequalit" in lowered or "condition" in lowered:
            steps.append("state the required inequalities or parameter conditions")
        if any(phrase in lowered for phrase in ("for each family", "other families are doing", "change its decision")):
            steps.append("check the relevant family-by-family deviation conditions")
        if any(phrase in lowered for phrase in ("numerical values", "give an example", "example of some")):
            steps.append("give a numerical example that satisfies the stated conditions")
        if "explain carefully" in lowered or "why they result" in lowered:
            steps.append("explain why the conditions imply the stated scenario")
    elif part.task_type == "comparative_statics":
        if "each of the parameters" in lowered or "changing each" in lowered:
            steps.append("cover each relevant parameter rather than only a subset")
        if "more likely" in lowered or "scenario" in lowered:
            steps.append("link each parameter change to the scenario comparison")
    elif part.task_type == "critique_and_revision":
        if "prompt" in lowered:
            steps.append("include the original prompt")
        if "output" in lowered:
            steps.append("include the model output and audit it")
        if any(phrase in lowered for phrase in ("page numbers", "quotes", "course material", "report")):
            steps.append("support the audit with cited evidence")
        if "rewrite" in lowered or "own tone" in lowered:
            steps.append("provide a corrected rewrite in the student's own tone")
    normalized: list[str] = []
    seen: set[str] = set()
    for step in steps:
        key = step.strip().lower()
        if key and key not in seen:
            normalized.append(step)
            seen.add(key)
    return normalized[:4]


def _criterion_anchors(criterion: str) -> tuple[str, str]:
    lowered = criterion.lower()
    if any(term in lowered for term in ("evidence", "reference", "page", "source", "support", "quote", "cited")):
        return (
            "Support is missing, vague, or not tied to the point.",
            "Support is specific and linked to the point being made.",
        )
    if any(term in lowered for term in ("judgement", "tradeoff", "qualification", "counterargument", "appraisal", "proportionate")):
        return (
            "Discussion stays descriptive or one-sided.",
            "Judgement is clear, balanced, and appropriately qualified.",
        )
    if any(term in lowered for term in ("working", "algebra", "derivation", "equation", "specification", "variable", "method")):
        return (
            "Method is incomplete or materially wrong.",
            "Method is complete, consistent, and materially correct.",
        )
    if any(term in lowered for term in ("interpret", "explain", "reasoning", "mechanism", "why", "comparison")):
        return (
            "Explanation is vague or does not follow from the setup.",
            "Explanation is clear and follows from the setup.",
        )
    if any(term in lowered for term in ("covers", "complete", "components", "steps", "gap")):
        return (
            "Required parts are missing or only partly attempted.",
            "Required parts are covered without material gaps.",
        )
    if any(term in lowered for term in ("addresses", "states", "defines", "identifies", "response")):
        return (
            "Misses the task or treats it only loosely.",
            "Directly addresses the task in the way asked.",
        )
    return (
        "Performance on this criterion is weak or incomplete.",
        "Performance on this criterion is secure and well supported.",
    )


def _question_text_for_model(question_text: str, max_chars: int | None = 650) -> str:
    raw = question_text.strip()
    if not raw:
        return ""
    cleaned = re.sub(r"(\w)-\s+(\w)", r"\1\2", raw)
    cleaned = " ".join(cleaned.split())
    cleaned = re.sub(r"\[\s*suggested\s+word\s+count\s*:[^\]]*\]", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(
        r"^(?:part|question|section)\s+[a-z0-9ivx]+(?:\s+q[a-z0-9ivx]+)?\s*(?:\(\d+(?:\.\d+)?\s*marks?\))?\s*",
        "",
        cleaned,
        flags=re.IGNORECASE,
    ).strip()
    if max_chars is None:
        return cleaned.rstrip(".")
    return _trim_goal_fragment(cleaned, max_chars=max_chars)


def _task_goal_for_model(task_goal: str, question_text: str) -> str:
    # Use max_chars=None so that an already-built task_goal is not re-truncated
    # to 650 chars here.  If the upstream builder already appended "..." (because
    # the source exceeded its own 1200-char cap), treat the result as unusable and
    # let the model rely on question.text alone.
    if task_goal.strip().endswith("..."):
        return ""
    cleaned_goal = _question_text_for_model(task_goal, max_chars=None)
    cleaned_question = _question_text_for_model(question_text, max_chars=None)
    if not cleaned_goal:
        return ""
    if len(cleaned_goal) < 24 and cleaned_goal.endswith(("(", "[", ":", ";", ",")):
        return ""
    goal_key = re.sub(r"\s+", " ", cleaned_goal.lower()).strip(" .")
    question_key = re.sub(r"\s+", " ", cleaned_question.lower()).strip(" .")
    if goal_key and question_key:
        if goal_key == question_key:
            return ""
        if goal_key in question_key and len(goal_key) >= max(20, int(len(question_key) * 0.45)):
            return ""
    return cleaned_goal


def _build_part_task_text(part: SubmissionPart) -> str:
    task_from_type = _task_goal_from_task_type(part.task_type, part.question_text_exact)
    if part.question_text_exact.strip():
        fuller_question_task = _question_text_for_model(part.question_text_exact, max_chars=1200)
        if part.task_type == "critique_and_revision" and fuller_question_task:
            return fuller_question_task
        if task_from_type and (
            "..." in task_from_type
            or len(task_from_type) < min(80, max(20, len(fuller_question_task) // 3))
            or _looks_damaged_question_text(task_from_type)
        ):
            return fuller_question_task or task_from_type
    if task_from_type:
        return task_from_type
    focus = part.focus_hint.strip()
    guidance_focus = _extract_task_focus_from_guidance(part.marking_guidance)
    if focus:
        return focus
    if guidance_focus:
        return guidance_focus
    return "Assess the content that best matches this section label."


def _extract_task_focus_from_guidance(marking_guidance: str) -> str:
    text = marking_guidance.strip()
    if not text:
        return ""
    match = re.search(r"Task focus:\s*(.+)", text)
    if match is not None:
        return match.group(1).strip().rstrip(".")
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        lowered = stripped.lower()
        if lowered.startswith(("unit:", "available marks:", "compact scoring criteria:", "detailed rubric:")):
            continue
        if lowered.startswith("use this rubric") or lowered.startswith("use uk-style quality bands"):
            continue
        if stripped.startswith("- "):
            continue
        return stripped.rstrip(".")
    return ""


def _extract_prompt_criteria(marking_guidance: str) -> list[str]:
    text = marking_guidance.strip()
    if not text:
        return []
    criteria: list[str] = []
    in_compact_block = False
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            if in_compact_block and criteria:
                break
            continue
        lowered = stripped.lower()
        if lowered == "compact scoring criteria:":
            in_compact_block = True
            continue
        if lowered.startswith("detailed rubric:"):
            if criteria:
                break
            in_compact_block = False
            continue
        if not stripped.startswith("- "):
            continue
        item = stripped[2:].strip()
        if not item:
            continue
        if in_compact_block:
            criteria.append(item)
        elif lowered.startswith(("- place work", "- use the top band", "- award explicit partial credit", "- rank attempts")):
            break
        elif item.lower().startswith(("evaluate whether", "assess whether", "check whether", "identify whether")):
            criteria.append(item)
        if len(criteria) >= 5:
            break
    if criteria:
        return criteria
    if "\n" not in text and len(text.split()) <= 40:
        return [text.rstrip(".") + "."]
    return criteria


def _default_prompt_criteria(part: SubmissionPart) -> list[str]:
    task_text = _build_part_task_text(part).lower()
    if _criterion_mode_for_part(part) == "abstract":
        return [
            f"Judge how directly the response addresses {task_text}.",
            "Judge how clearly the reasoning supports the conclusion reached.",
            "Judge how fully the section covers the task and justifies the score awarded.",
        ]
    return [
        f"Confirm that the response directly addresses {task_text}.",
        "Confirm that the reasoning, method, or explanation is accurate enough to justify credit.",
        "Confirm that the required steps or components are present without a material gap.",
    ]


def _extract_prompt_scoring_rules(marking_guidance: str, max_mark: float | None) -> list[str]:
    rules: list[str] = []
    guidance_text = marking_guidance.lower()
    if any(band in guidance_text for band in ("first", "2:1", "2:2", "third/pass", "missing/unfinished", "fail")):
        rules.append("Choose the quality band first, then place the score within it.")
    if any(term in guidance_text for term in ("partial credit", "algebra", "equation", "deriv", "method", "interpretation")):
        rules.append("Give partial credit for correct setup, method, or interpretation.")
    if any(term in guidance_text for term in ("reference", "references", "page number", "page numbers", "quote", "quotes")):
        rules.append("Reward specific references when the task asks for them.")
    if any(term in guidance_text for term in ("rewrite", "own tone", "own words", "correct any mistakes", "edit the output")):
        rules.append("Reward justified correction and rewriting when asked.")
    # M2 addition: explicit guidance for no-attempt cases
    rules.append("If a response contains no relevant attempt, no attempt at the required method, or only unrelated content, award zero or minimal credit.")
    deduped: list[str] = []
    seen: set[str] = set()
    for rule in rules:
        key = rule.lower().rstrip(".")
        if key in seen:
            continue
        seen.add(key)
        deduped.append(rule)
    return deduped


def _task_goal_from_task_type(task_type: str, question_text: str) -> str:
    derived_goal = _derive_task_goal_from_question(task_type, question_text)
    if derived_goal:
        return derived_goal
    if task_type == "critique_and_revision":
        return "complete the requested draft, audit, and revision steps using the supplied materials"
    if task_type == "model_specification":
        return "state the required formal specification and explain what each term means"
    if task_type == "comparative_statics":
        return "explain how parameter changes shift the likelihood or direction of the relevant outcomes"
    if task_type == "deterministic_derivation":
        return "derive the required conditions and justify why they produce the stated scenario"
    if task_type == "explanation_interpretation":
        return "explain the required mechanism, parameter effects, or interpretation"
    if task_type == "welfare_reasoning":
        return "explain the welfare implications, relevant conditions, and effects on the affected parties"
    if task_type == "evaluative_discussion":
        return "make and justify an evaluative judgement using the relevant analysis or concepts"
    if task_type == "prediction_generation":
        return "state the requested predictions or implications clearly and correctly"
    if task_type == "measurement_and_data_design":
        return "show how the relevant variables or constructs could be measured and linked to suitable data"
    if task_type == "regression_specification":
        return "write the regression equation, define its terms, and state the testable hypothesis"
    if task_type == "causal_identification":
        return "explain the endogeneity risks and propose feasible remedies"
    if task_type == "synthesis_across_sources":
        return "use the relevant framework and materials to evaluate claims, omissions, and possible extensions"
    normalized = " ".join(question_text.split()).strip()
    if not normalized:
        return ""
    first_sentence = re.split(r"(?<=[.!?])\s+", normalized, maxsplit=1)[0].strip()
    return first_sentence[:220].rstrip()


def _derive_task_goal_from_question(task_type: str, question_text: str) -> str:
    cleaned = _clean_question_text_for_goal(question_text)
    if not cleaned:
        return ""
    numbered_items = _extract_numbered_goal_items(cleaned)
    if task_type == "critique_and_revision" and numbered_items:
        selected = [
            _trim_goal_fragment(item, max_chars=180)
            for item in numbered_items[:3]
            if item.strip()
        ]
        return "; ".join(fragment for fragment in selected if fragment)
    if numbered_items:
        lead_item = _trim_goal_fragment(numbered_items[0], max_chars=260)
        if lead_item:
            return lead_item
    cleaned_without_numbering = _trim_goal_fragment(cleaned, max_chars=max(len(cleaned), 320))
    first_sentence = re.split(r"(?<=[.!?])\s+", cleaned_without_numbering, maxsplit=1)[0].strip()
    return _trim_goal_fragment(first_sentence, max_chars=260)


def _clean_question_text_for_goal(question_text: str) -> str:
    text = " ".join(question_text.split()).strip()
    if not text:
        return ""
    patterns = (
        r"^(?:part|question|section)\s+[a-z0-9ivx]+(?:\s+q[a-z0-9ivx]+)?\s*(?:\(\d+(?:\.\d+)?\s*marks?\))?\s*",
    )
    changed = True
    while changed:
        changed = False
        for pattern in patterns:
            updated = re.sub(pattern, "", text, flags=re.IGNORECASE).strip()
            if updated != text:
                text = updated
                changed = True
    return text


def _extract_numbered_goal_items(question_text: str) -> list[str]:
    matches = list(re.finditer(r"(?:^|\s)(\d+)\.\s+", question_text))
    if len(matches) < 2:
        return []
    items: list[str] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(question_text)
        item = question_text[start:end].strip(" ;.")
        if item:
            items.append(item)
    return items


def _trim_goal_fragment(text: str, max_chars: int) -> str:
    normalized = re.sub(r"^\d+\.\s*(?:\[\d+(?:\.\d+)?\s*marks?\])?\s*", "", text.strip(), flags=re.IGNORECASE)
    normalized = normalized.rstrip(".")
    if len(normalized) <= max_chars:
        return normalized
    trimmed = normalized[: max_chars - 3].rstrip(" ,;:")
    return trimmed + "..."


def build_structure_messages(script_text: str, filename: str) -> list[dict[str, str]]:
    return build_structure_messages_with_guidance(script_text, filename, "")


def _build_structure_detection_chunks(script_text: str, max_chunks: int = 8) -> str:
    text = script_text.replace("\r\n", "\n")
    if not text.strip():
        return ""

    spans: list[tuple[int, int]] = []
    spans.append((0, min(len(text), 700)))
    for match in _STRUCTURE_CUE_PATTERN.finditer(text):
        start = max(0, match.start() - 100)
        end = min(len(text), match.start() + 420)
        spans.append((start, end))
        if len(spans) >= max_chunks:
            break

    merged: list[tuple[int, int]] = []
    for start, end in sorted(spans):
        if not merged or start > merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))

    chunks = []
    for index, (start, end) in enumerate(merged[:max_chunks], start=1):
        snippet = text[start:end].strip()
        if not snippet:
            continue
        chunks.append(f"[Chunk {index} @ {start}:{end}]\\n{snippet}")
    return "\n\n---\n\n".join(chunks)


def _build_subpart_detection_chunks(section_text: str, child_specs: list[SubmissionPart], max_chunks: int = 8) -> str:
    text = section_text.replace("\r\n", "\n")
    if not text.strip():
        return ""

    spans: list[tuple[int, int]] = [(0, min(len(text), 650))]
    for child in child_specs:
        candidate = (child.anchor_text or child.label).strip()
        if not candidate:
            continue
        pattern = re.compile(rf"(?im)^\\s*{re.escape(candidate)}(?:\\b|\\s*[:.)\\]-])")
        for match in pattern.finditer(text):
            start = max(0, match.start() - 90)
            end = min(len(text), match.start() + 360)
            spans.append((start, end))
            if len(spans) >= max_chunks:
                break
        if len(spans) >= max_chunks:
            break

    merged: list[tuple[int, int]] = []
    for start, end in sorted(spans):
        if not merged or start > merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))

    chunks = []
    for index, (start, end) in enumerate(merged[:max_chunks], start=1):
        snippet = text[start:end].strip()
        if not snippet:
            continue
        chunks.append(f"[Chunk {index} @ {start}:{end}]\\n{snippet}")
    return "\n\n---\n\n".join(chunks)


def _preferred_detector_model(primary_model_name: str) -> str:
    preferred = DEFAULT_DETECTOR_MODEL_NAME.strip()
    if preferred:
        return preferred
    return primary_model_name


def _call_structure_detector_with_fallback(
    primary_model_name: str,
    messages: list[dict[str, str]],
    ollama_url: str,
) -> dict[str, Any]:
    detector_model = _preferred_detector_model(primary_model_name)
    attempted: list[str] = []
    last_error: Exception | None = None
    for candidate in (detector_model, primary_model_name):
        candidate = candidate.strip()
        if not candidate or candidate in attempted:
            continue
        attempted.append(candidate)
        try:
            return _call_ollama_json(
                model_name=candidate,
                messages=messages,
                ollama_url=ollama_url,
                profile="detector_small",
            )
        except Exception as exc:
            last_error = exc
    if last_error is not None:
        raise last_error
    raise ValueError("No detection model available for structure extraction call.")


def build_structure_messages_with_guidance(script_text: str, filename: str, structure_guidance: str) -> list[dict[str, str]]:
    chunked_answer = _build_structure_detection_chunks(script_text)
    system = (
        "You are a fast structure detector for small local models. "
        "Use only the supplied chunks, keep reasoning local, and do not grade."
    )
    guidance_block = ""
    if structure_guidance.strip():
        guidance_block = (
            "STRUCTURE HINTS FROM MARKING DOCUMENTS:\n"
            f"{structure_guidance.strip()}\n\n"
        )
    user = (
        f"STUDENT FILE NAME: {filename}\n\n"
        f"{guidance_block}"
        "TARGETED SUBMISSION CHUNKS (not full script):\n"
        f"\"\"\"{chunked_answer}\"\"\"\n\n"
        "Return only one JSON object with exactly this structure:\n"
        '{ "sections": [ { "label": string, "focus_hint": string, "anchor_text": string } ] }\n\n'
        "Rules:\n"
        "1. Use structure hints only when they are specific and relevant; otherwise rely on the chunks.\n"
        "2. Prefer coarse structure when uncertain; split only when chunk evidence is clear.\n"
        "3. label should be short, e.g., 'Question 1', 'Part 2', 'Part 2 Q3'.\n"
        "4. focus_hint must be <= 18 words and action-specific.\n"
        "5. anchor_text must be a short verbatim quote from a chunk near that section start.\n"
        "6. Return at most 16 sections.\n"
        "7. If separation is unclear, return one section labelled 'Whole Submission'.\n"
        "8. Do not include markdown fences or any text outside JSON."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_subpart_structure_messages(parent_part: SubmissionPart, filename: str, child_specs: list[SubmissionPart]) -> list[dict[str, str]]:
    chunked_parent_text = _build_subpart_detection_chunks(parent_part.section_text, child_specs)
    child_labels = ", ".join(
        f"{child.label} ({_format_number(child.max_mark) if child.max_mark is not None else '?'})"
        for child in child_specs
    )
    system = (
        "You are a fast subpart detector for small local models. "
        "Only use the supplied chunks and keep output compact."
    )
    user = (
        f"STUDENT FILE NAME: {filename}\n"
        f"PARENT SECTION: {parent_part.label}\n"
        f"EXPECTED CHILD UNITS FROM MARKING DOCUMENTS: {child_labels}\n\n"
        "PARENT SECTION CHUNKS (not full section):\n"
        f"\"\"\"{chunked_parent_text}\"\"\"\n\n"
        "Return only one JSON object with exactly this structure:\n"
        '{ "sections": [ { "label": string, "focus_hint": string, "anchor_text": string } ] }\n\n'
        "Rules:\n"
        "1. Use child labels only when the student's section text clearly separates them.\n"
        "2. If the split is unclear, return one section with the parent label.\n"
        "3. anchor_text must be a short exact quote copied verbatim from a chunk near the start of that child section.\n"
        "4. focus_hint must be <= 14 words and specific.\n"
        "5. Do not invent labels beyond the parent and expected child units.\n"
        "6. Return at most the number of expected child units.\n"
        "7. Do not include markdown fences or any text outside the JSON object."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_moderation_messages(
    script_text: str,
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
    context: MarkingContext,
    filename: str,
) -> list[dict[str, str]]:
    context_text = _build_context_text(context)
    payload = []
    for part, analysis in zip(parts, part_analyses):
        payload.append(
            {
                "section_label": part.label,
                "section_max_mark": part.max_mark,
                "current_provisional_score": analysis.get("provisional_score"),
                "current_provisional_score_0_to_100": analysis.get("provisional_score_0_to_100"),
                "strengths": analysis.get("strengths", []),
                "weaknesses": analysis.get("weaknesses", []),
                "evidence": analysis.get("evidence", []),
                "coverage_comment": analysis.get("coverage_comment", ""),
            }
        )
    analyses_text = json.dumps(payload, ensure_ascii=True, indent=2)
    system = (
        "You are moderating section-by-section marks within a single student's submission. "
        "Compare the sections against each other for internal consistency. "
        "Adjust a section only when the current score is clearly too harsh or too generous relative to the evidence and the other sections. "
        "The final overall mark will be computed separately by arithmetic, so you must only return section-level adjustments."
    )
    user = (
        f"{context_text}\n\n"
        f"STUDENT FILE NAME: {filename}\n\n"
        "FULL STUDENT ANSWER:\n"
        f"\"\"\"{script_text}\"\"\"\n\n"
        "CURRENT SECTION ANALYSES:\n"
        f"{analyses_text}\n\n"
        "Return only one JSON object with exactly this structure:\n"
        '{ "adjusted_sections": [ { "section_label": string, "adjusted_provisional_score": number|null, "adjusted_provisional_score_0_to_100": number|null, "rationale": string } ] }\n\n'
        "Rules:\n"
        "1. Include every section exactly once.\n"
        "2. Use adjusted_provisional_score only for sections that have an explicit max mark.\n"
        "3. Use adjusted_provisional_score_0_to_100 only when the section is on a percentage scale.\n"
        "4. Keep changes bounded and evidence-based; do not rewrite the whole assessment.\n"
        "5. If a current section score already looks internally consistent, keep it unchanged.\n"
        "6. Do not include markdown fences or any text outside the JSON object."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_synthesis_messages(
    script_text: str,
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
    context: MarkingContext,
    filename: str,
) -> list[dict[str, str]]:
    max_mark_text = _format_number(context.max_mark)
    part_labels = ", ".join(part.label for part in parts)
    analyses_text = json.dumps(part_analyses, ensure_ascii=True, indent=2)
    context_text = _build_context_text(context)
    part_weights_text = ", ".join(
        f"{part.label}={_format_number(part.max_mark)}" for part in parts if part.max_mark is not None
    )
    system = (
        "You are a fair and consistent university examiner completing the final synthesis stage. "
        "Use the supplied marking documents, the full student submission, and the per-section analyses. "
        "You must reward and penalize consistently across sections."
    )
    user = (
        "You are synthesizing a final mark from section-by-section analysis.\n\n"
        f"{context_text}\n\n"
        f"STUDENT FILE NAME: {filename}\n"
        f"DETECTED SECTIONS: {part_labels}\n\n"
        f"{'SECTION MARKS: ' + part_weights_text + chr(10) + chr(10) if part_weights_text else ''}"
        "FULL STUDENT ANSWER:\n"
        f"\"\"\"{script_text}\"\"\"\n\n"
        "SECTION ANALYSES:\n"
        f"{analyses_text}\n\n"
        "Return only one JSON object with exactly these keys:\n"
        '- "overall_feedback": string\n'
        '- "covered_parts": array of strings\n'
        '- "strengths": array of strings\n'
        '- "weaknesses": array of strings\n\n'
        "Rules:\n"
        f"1. The overall maximum mark is {max_mark_text}, but do not return any mark fields.\n"
        "2. Use the section analyses and section marks to support feedback; do not ignore a weaker section simply because other sections are strong.\n"
        "3. overall_feedback must be at least 140 words.\n"
        "4. overall_feedback must explicitly mention every detected section label.\n"
        "5. covered_parts must list every detected section label.\n"
        "6. Do not include markdown fences or any text outside the JSON object."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_calibration_messages(
    assessment_name: str,
    context: MarkingContext,
    model_label: str,
    student_summaries: list[dict[str, Any]],
) -> list[dict[str, str]]:
    max_mark_text = _format_number(context.max_mark)
    context_text = _build_context_text(context)
    summaries_text = json.dumps(student_summaries, ensure_ascii=True, indent=2)
    system = (
        "You are calibrating marks across a cohort after provisional grading. "
        "Your job is to improve consistency across students answering the same assessment. "
        "Do not invent evidence beyond the supplied summaries."
    )
    user = (
        f"ASSESSMENT: {assessment_name}\n"
        f"MODEL LABEL: {model_label}\n\n"
        f"{context_text}\n\n"
        "PROVISIONAL STUDENT SUMMARIES:\n"
        f"{summaries_text}\n\n"
        "Return only one JSON object with exactly this structure:\n"
        '{ "calibrated_results": [ { "filename": string, "adjusted_mark": number, "rationale": string } ] }\n\n'
        "Rules:\n"
        f"1. Every adjusted_mark must be between 0 and {max_mark_text}.\n"
        "2. Include every filename exactly once.\n"
        "3. Keep changes modest and evidence-based.\n"
        "4. Use the relative quality signals in the summaries to improve consistency across students and sections.\n"
        "5. Do not include markdown fences or any text outside the JSON object."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_verification_messages(
    script_text: str,
    context: MarkingContext,
    filename: str,
    grading_result: dict[str, Any],
) -> list[dict[str, str]]:
    context_text = _build_context_text(context)
    grading_summary = json.dumps(
        {
            "total_mark": grading_result.get("total_mark"),
            "max_mark": grading_result.get("max_mark"),
            "covered_parts": grading_result.get("covered_parts", []),
            "strengths": grading_result.get("strengths", []),
            "weaknesses": grading_result.get("weaknesses", []),
            "overall_feedback": grading_result.get("overall_feedback", ""),
        },
        ensure_ascii=True,
        indent=2,
    )
    system = (
        "You are verifying an existing assessment grade. "
        "Check whether the mark and feedback are plausible given the submission and the marking documents. "
        "Do not regrade from scratch unless needed to explain a discrepancy."
    )
    user = (
        f"{context_text}\n\n"
        f"STUDENT FILE NAME: {filename}\n\n"
        "FULL STUDENT ANSWER:\n"
        f"\"\"\"{script_text}\"\"\"\n\n"
        "PROPOSED GRADING RESULT:\n"
        f"{grading_summary}\n\n"
        "Return only one JSON object with exactly these keys:\n"
        '- "agreement": string\n'
        '- "confidence_0_to_100": number\n'
        '- "issues": array of strings\n'
        '- "recommendation": string\n\n'
        "Rules:\n"
        "1. agreement must be one of: agree, minor_concern, major_concern.\n"
        "2. issues should be concrete and brief.\n"
        "3. recommendation should say whether to accept, review, or regrade.\n"
        "4. Do not include markdown fences or any text outside the JSON object."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def build_regrade_messages(
    script_text: str,
    context: MarkingContext,
    filename: str,
    prior_result: dict[str, Any],
    verifier_result: dict[str, Any],
) -> list[dict[str, str]]:
    context_text = _build_context_text(context)
    prior_summary = json.dumps(
        {
            "total_mark": prior_result.get("total_mark"),
            "max_mark": prior_result.get("max_mark"),
            "covered_parts": prior_result.get("covered_parts", []),
            "strengths": prior_result.get("strengths", []),
            "weaknesses": prior_result.get("weaknesses", []),
            "overall_feedback": prior_result.get("overall_feedback", ""),
        },
        ensure_ascii=True,
        indent=2,
    )
    verifier_summary = json.dumps(verifier_result, ensure_ascii=True, indent=2)
    max_mark_text = _format_number(context.max_mark)
    system = (
        "You are revising a grading decision after verifier feedback. "
        "Reassess the mark and feedback carefully, using the verifier concerns only when they are valid. "
        "Do not defend the previous answer by reflex."
    )
    user = (
        f"{context_text}\n\n"
        f"STUDENT FILE NAME: {filename}\n\n"
        "FULL STUDENT ANSWER:\n"
        f"\"\"\"{script_text}\"\"\"\n\n"
        "PRIOR GRADING RESULT:\n"
        f"{prior_summary}\n\n"
        "VERIFIER FEEDBACK:\n"
        f"{verifier_summary}\n\n"
        "Return only one JSON object with exactly these keys:\n"
        '- "total_mark": number\n'
        '- "max_mark": number\n'
        '- "overall_feedback": string\n'
        '- "covered_parts": array of strings\n'
        '- "strengths": array of strings\n'
        '- "weaknesses": array of strings\n\n'
        "Rules:\n"
        f"1. Use max_mark = {max_mark_text}.\n"
        f"2. total_mark must be between 0 and {max_mark_text}.\n"
        "3. overall_feedback must be at least 140 words.\n"
        "4. Address verifier concerns when they are supported by the submission and marking documents.\n"
        "5. If verifier concerns are not supported, keep the stronger original judgement.\n"
        "6. Do not include markdown fences or any text outside the JSON object."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def call_ollama(
    script_text: str,
    context: MarkingContext,
    filename: str,
    model_name: str,
    rubric_verifier_model_name: str | None = None,
    part_verifier_model_name: str | None = None,
    structure_script_text: str | None = None,
    prepared_assessment_map: PreparedAssessmentMap | None = None,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> dict[str, Any]:
    return call_ollama_comparative_first(
        script_text=script_text,
        context=context,
        filename=filename,
        model_name=model_name,
        rubric_verifier_model_name=rubric_verifier_model_name,
        part_verifier_model_name=part_verifier_model_name,
        structure_script_text=structure_script_text,
        prepared_assessment_map=prepared_assessment_map,
        ollama_url=ollama_url,
    )


def call_ollama_comparative_first(
    script_text: str,
    context: MarkingContext,
    filename: str,
    model_name: str,
    rubric_verifier_model_name: str | None = None,
    part_verifier_model_name: str | None = None,
    structure_script_text: str | None = None,
    prepared_assessment_map: PreparedAssessmentMap | None = None,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> dict[str, Any]:
    timings: dict[str, float] = {
        "assessment_prepare": 0.0,
        "structure_detect": 0.0,
        "part_refine": 0.0,
        "part_analysis": 0.0,
        "moderation": 0.0,
        "finalize": 0.0,
    }
    script_text = script_text.strip()
    if not script_text:
        raise ValueError(
            "No text could be extracted from the submission. "
            "OCR is not implemented yet, so use a text-based PDF, DOCX, or CSV answer."
        )
    use_alternate_structure_text = bool((structure_script_text or "").strip()) and bool(extract_expected_parts_from_context(context))
    structure_input_text = (structure_script_text or "").strip() if use_alternate_structure_text else script_text

    if prepared_assessment_map is None:
        stage_started = time.perf_counter()
        prepared_assessment_map = prepare_assessment_map(
            context=context,
            verifier_model_name=rubric_verifier_model_name,
            ollama_url=ollama_url,
        )
        timings["assessment_prepare"] = round(time.perf_counter() - stage_started, 2)
    else:
        timings["assessment_prepare"] = round(prepared_assessment_map.preparation_latency_seconds, 2)
    assessment_map = prepared_assessment_map.prepared_map
    total_started = time.perf_counter()
    stage_started = time.perf_counter()
    detected_parts = detect_submission_parts(structure_input_text, filename, model_name, context=context, ollama_url=ollama_url)
    timings["structure_detect"] = round(time.perf_counter() - stage_started, 2)
    stage_started = time.perf_counter()
    parts = segment_submission_parts(script_text, detected_parts)
    parts = refine_submission_granularity(parts, context, filename, model_name, ollama_url=ollama_url)
    parts = apply_assessment_map_to_submission_parts(parts, assessment_map)
    parts = apply_prepared_artifact_to_submission_parts(parts, prepared_assessment_map)
    timings["part_refine"] = round(time.perf_counter() - stage_started, 2)
    diagnostics = build_submission_diagnostics(script_text, parts)

    part_analyses = []
    stage_started = time.perf_counter()
    for part in parts:
        if not part.section_text.strip():
            part_analyses.append(build_missing_part_analysis(part))
            continue
        analysis = _run_part_analysis_with_retry(
            part=part,
            context=context,
            filename=filename,
            model_name=model_name,
            ollama_url=ollama_url,
        )
        if part_verifier_model_name:
            verifier_result = verify_part_analysis(
                part=part,
                context=context,
                filename=filename,
                part_analysis=analysis,
                model_name=part_verifier_model_name,
                ollama_url=ollama_url,
            )
            analysis = apply_part_verifier_result(analysis, verifier_result)
        part_analyses.append(analysis)
    timings["part_analysis"] = round(time.perf_counter() - stage_started, 2)

    if len(part_analyses) > 1:
        stage_started = time.perf_counter()
        part_analyses = moderate_linked_part_analyses(
            script_text=script_text,
            context=context,
            filename=filename,
            parts=parts,
            part_analyses=part_analyses,
            assessment_map=assessment_map,
            model_name=model_name,
            ollama_url=ollama_url,
        )
        timings["moderation"] = round(time.perf_counter() - stage_started, 2)

    if part_analyses:
        stage_started = time.perf_counter()
        result = build_local_final_result(
            context=context,
            diagnostics=diagnostics,
            parts=parts,
            part_analyses=part_analyses,
        )
        timings["finalize"] = round(time.perf_counter() - stage_started, 2)
        return _attach_latency_metrics(result, timings, total_started, parts)

    synthesis_messages = build_synthesis_messages(script_text, parts, part_analyses, context, filename)
    stage_started = time.perf_counter()
    result = _run_final_result_with_retry(
        model_name=model_name,
        messages=synthesis_messages,
        expected_max_mark=context.max_mark,
        diagnostics=diagnostics,
        parts=parts,
        part_analyses=part_analyses,
        ollama_url=ollama_url,
    )
    timings["finalize"] = round(time.perf_counter() - stage_started, 2)
    return _attach_latency_metrics(result, timings, total_started, parts)


def extract_message_content(response_json: dict[str, Any]) -> str:
    if not isinstance(response_json, dict):
        raise ValueError("Ollama returned a non-object JSON response.")
    message = response_json.get("message")
    if not isinstance(message, dict):
        raise ValueError(f"Ollama response did not include a message object: {response_json}")
    content = message.get("content")
    if not isinstance(content, str) or not content.strip():
        raise ValueError(f"Ollama response did not include message.content text: {response_json}")
    return content.strip()


def parse_json_object(content: str) -> dict[str, Any]:
    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        data = _parse_first_json_object(content)
    if not isinstance(data, dict):
        raise ValueError(f"Model response was not a JSON object: {content}")
    return data


def detect_submission_parts(
    script_text: str,
    filename: str,
    model_name: str,
    context: MarkingContext | None = None,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> list[SubmissionPart]:
    context_fast_path = _detect_submission_parts_from_context(script_text, context)
    if context_fast_path is not None:
        return context_fast_path
    structure_guidance = extract_structure_guidance(context) if context is not None else ""
    data = _call_structure_detector_with_fallback(
        primary_model_name=model_name,
        messages=build_structure_messages_with_guidance(script_text, filename, structure_guidance),
        ollama_url=ollama_url,
    )
    return reconcile_detected_parts(normalize_detected_parts(data), context)


def describe_structure_detection_mode(
    script_text: str,
    context: MarkingContext | None,
) -> dict[str, Any]:
    context_fast_path = _detect_submission_parts_from_context(script_text, context)
    expected_parts = extract_expected_parts_from_context(context) if context is not None else []
    if context_fast_path is not None:
        return {
            "mode": "context_fast_path",
            "expected_part_count": len(expected_parts),
            "detected_part_count": len(context_fast_path),
        }
    return {
        "mode": "model_call",
        "expected_part_count": len(expected_parts),
        "detected_part_count": None,
    }


def _detect_submission_parts_from_context(
    script_text: str,
    context: MarkingContext | None,
) -> list[SubmissionPart] | None:
    if context is None:
        return None
    expected_parts = extract_expected_parts_from_context(context)
    if not expected_parts:
        return None
    if len(expected_parts) == 1:
        return expected_parts

    normalized_text = script_text.replace("\r\n", "\n")
    anchored_count = sum(1 for part in expected_parts if _find_part_anchor_position(normalized_text, part) != -1)
    if anchored_count >= max(2, len(expected_parts) // 2):
        return expected_parts
    return None


def normalize_detected_parts(data: dict[str, Any]) -> list[SubmissionPart]:
    sections = data.get("sections")
    if not isinstance(sections, list) or not sections:
        return [SubmissionPart(label="Whole Submission", focus_hint="Assess the full script as one piece of work.", anchor_text="")]

    parts: list[SubmissionPart] = []
    seen_labels: set[str] = set()
    for item in sections:
        if not isinstance(item, dict):
            continue
        label = str(item.get("label", "")).strip()
        if not label or _is_placeholder_section_label(label):
            continue
        focus_hint = str(item.get("focus_hint", "")).strip()
        anchor_text = str(item.get("anchor_text", "")).strip()
        if label in seen_labels:
            continue
        seen_labels.add(label)
        parts.append(SubmissionPart(label=label, focus_hint=focus_hint, anchor_text=anchor_text))

    if not parts:
        return [SubmissionPart(label="Whole Submission", focus_hint="Assess the full script as one piece of work.", anchor_text="")]
    return parts


def reconcile_detected_parts(detected_parts: list[SubmissionPart], context: MarkingContext | None) -> list[SubmissionPart]:
    if context is None:
        return detected_parts

    expected_parts = extract_expected_parts_from_context(context)
    if not expected_parts:
        return detected_parts

    expected_by_key = {_normalize_label_key(part.label): part for part in expected_parts}
    detected_by_key = {
        _normalize_label_key(part.label): part
        for part in detected_parts
        if _normalize_label_key(part.label) in expected_by_key
    }
    enriched = []
    for expected in expected_parts:
        part = detected_by_key.get(_normalize_label_key(expected.label), expected)
        enriched.append(
            SubmissionPart(
                label=expected.label,
                focus_hint=part.focus_hint or expected.focus_hint,
                anchor_text=part.anchor_text or expected.anchor_text,
                section_text=part.section_text,
                max_mark=expected.max_mark,
                marking_guidance=expected.marking_guidance,
                question_text_exact=expected.question_text_exact,
                task_type=expected.task_type,
            )
        )
    return enriched


def segment_submission_parts(script_text: str, detected_parts: list[SubmissionPart]) -> list[SubmissionPart]:
    if len(detected_parts) == 1 and detected_parts[0].label == "Whole Submission":
        return [SubmissionPart(label="Whole Submission", focus_hint=detected_parts[0].focus_hint, anchor_text="", section_text=script_text.strip(), max_mark=detected_parts[0].max_mark, marking_guidance=detected_parts[0].marking_guidance, question_text_exact=detected_parts[0].question_text_exact, task_type=detected_parts[0].task_type)]

    normalized_text = script_text.replace("\r\n", "\n")
    toc_range = _detect_toc_block_range(normalized_text, detected_parts)
    anchors: list[tuple[int, SubmissionPart]] = []
    used_positions: set[int] = set()
    for part in detected_parts:
        position = _find_part_anchor_position(normalized_text, part, toc_range=toc_range)
        if position == -1 or position in used_positions:
            continue
        used_positions.add(position)
        anchors.append((position, part))

    if len(anchors) < max(1, len(detected_parts) // 2):
        return [
            SubmissionPart(
                label=part.label,
                focus_hint=part.focus_hint,
                anchor_text=part.anchor_text,
                section_text=script_text.strip(),
                max_mark=part.max_mark,
                marking_guidance=part.marking_guidance,
                question_text_exact=part.question_text_exact,
                task_type=part.task_type,
            )
            for part in detected_parts
        ]

    anchors.sort(key=lambda item: item[0])
    segmented: list[SubmissionPart] = []
    for index, (start, part) in enumerate(anchors):
        end = anchors[index + 1][0] if index + 1 < len(anchors) else len(normalized_text)
        section_text = _select_local_chunk_with_fallback(
            normalized_text,
            start,
            end,
            part.question_text_exact or part.focus_hint,
        )
        segmented.append(
            SubmissionPart(
                label=part.label,
                focus_hint=part.focus_hint,
                anchor_text=part.anchor_text,
                section_text=section_text,
                max_mark=part.max_mark,
                marking_guidance=part.marking_guidance,
                question_text_exact=part.question_text_exact,
                task_type=part.task_type,
            )
        )
    return segmented


def refine_submission_granularity(
    segmented_parts: list[SubmissionPart],
    context: MarkingContext | None,
    filename: str,
    model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
    include_unmarked_subparts: bool = False,
) -> list[SubmissionPart]:
    if context is None or not segmented_parts:
        return segmented_parts

    child_map = extract_expected_subparts_from_context(context, require_marks=not include_unmarked_subparts)
    if not child_map:
        return segmented_parts

    refined: list[SubmissionPart] = []
    for part in segmented_parts:
        child_specs = child_map.get(_normalize_label_key(part.label), [])
        if len(child_specs) < 2:
            refined.append(part)
            continue
        if not part.section_text.strip():
            refined.extend(_resolve_expected_child_parts(parent_part=part, child_parts=[] , child_specs=child_specs))
            continue
        segmented_children = _segment_subparts_within_section(part, child_specs)
        if segmented_children is None:
            segmented_children = _detect_subparts_with_model(part, filename, child_specs, model_name, ollama_url)
        if segmented_children is None:
            refined.extend(_resolve_expected_child_parts(parent_part=part, child_parts=[], child_specs=child_specs))
            continue
        refined.extend(_resolve_expected_child_parts(parent_part=part, child_parts=segmented_children, child_specs=child_specs))
    return refined


def extract_structure_guidance(context: MarkingContext) -> str:
    source_text = "\n\n".join(
        text.strip()
        for text in (context.rubric_text, context.brief_text, context.marking_scheme_text)
        if text and text.strip()
    )
    if not source_text:
        return ""

    lines = [line.strip() for line in source_text.splitlines() if line.strip()]
    useful_lines: list[str] = []
    patterns = (
        r"\b(question|part|section|subsection|subpart|task|item)\s+[a-z0-9ivx]+\b",
        r"\b\d+(?:\.\d+)?\s*(marks?|points?|%)\b",
        r"\b(weight|weights|weighting|worth|allocated|allocation|available)\b",
        r"\bout of\s+\d+\b",
    )
    for line in lines:
        if any(re.search(pattern, line, flags=re.IGNORECASE) for pattern in patterns):
            useful_lines.append(line)
        if len(useful_lines) >= 12:
            break

    return "\n".join(useful_lines)


_TOP_LEVEL_CONTEXT_PATTERN = re.compile(
    r"^(?P<label>(?:question|part|section|task|item)\s+(?:\d+|[ivx]+|[a-z]))\b(?P<tail>.*)$",
    flags=re.IGNORECASE,
)
_SUBPART_CONTEXT_PATTERN = re.compile(
    r"^(?P<token>(?:\d+|[a-z])(?:[\.\)]|\])|\((?:\d+|[a-z])\))\s*(?P<body>.*)$",
    flags=re.IGNORECASE,
)
_WEIGHT_TEXT_PATTERN = re.compile(
    r"(?P<value>\d+(?:\.\d+)?)\s*(?P<unit>marks?|points?|%)\b|\b(?:worth|allocated|allocation|available)\s+(?P<alt>\d+(?:\.\d+)?)\s*(?P<alt_unit>marks?|points?|%)\b",
    flags=re.IGNORECASE,
)
_DEPENDENCY_LABEL_PATTERN = re.compile(
    r"\b(?P<kind>part|question|section)\s+(?P<id>\d+|[ivx]+)(?:\s*q(?P<child>\d+))?\b",
    flags=re.IGNORECASE,
)


def extract_assessment_structure(context: MarkingContext) -> tuple[AssessmentStructureSection, ...]:
    prompt_sections = extract_assessment_prompt_structure(context)
    marking_scheme_sections = extract_marking_scheme_structure(context)
    if prompt_sections and marking_scheme_sections:
        return _reconcile_assessment_structure_sections(prompt_sections, marking_scheme_sections)
    if prompt_sections:
        return prompt_sections
    return marking_scheme_sections


def extract_assessment_prompt_structure(context: MarkingContext) -> tuple[AssessmentStructureSection, ...]:
    source_text = "\n\n".join(
        text.strip()
        for text in (context.rubric_text, context.brief_text)
        if text and text.strip()
    )
    if not source_text:
        return ()
    return _merge_assessment_structure_sections(_deterministic_extract_assessment_structure(source_text))


def extract_marking_scheme_structure(context: MarkingContext) -> tuple[AssessmentStructureSection, ...]:
    source_text = context.marking_scheme_text.strip()
    if not source_text:
        return ()
    return _merge_assessment_structure_sections(_deterministic_extract_assessment_structure(source_text))


def _deterministic_extract_assessment_structure(source_text: str) -> tuple[AssessmentStructureSection, ...]:
    lines = [line.strip() for line in source_text.replace("\r\n", "\n").splitlines() if line.strip()]
    sections: list[AssessmentStructureSection] = []
    current_label = ""
    current_lines: list[str] = []

    def flush_section() -> None:
        nonlocal current_label, current_lines
        if not current_label:
            return
        sections.append(_build_structure_section(_canonical_context_label(current_label), current_lines))
        current_label = ""
        current_lines = []

    for line in lines:
        top_match = _TOP_LEVEL_CONTEXT_PATTERN.match(line)
        if top_match is not None and _is_top_level_context_header(line, top_match.group("tail")):
            flush_section()
            current_label = top_match.group("label").strip()
            current_lines = [line]
            continue
        if current_label:
            current_lines.append(line)
    flush_section()
    return tuple(section for section in sections if section.label)


def _build_structure_section(label: str, lines: list[str]) -> AssessmentStructureSection:
    header = lines[0] if lines else label
    body_lines = lines[1:] if len(lines) > 1 else []
    question_lines, marking_lines = _split_question_and_marking_lines(body_lines)
    children = _extract_structure_children(label, question_lines)
    question_text = "\n".join([header, *question_lines]).strip() if question_lines else header.strip()
    marking_text = "\n".join(marking_lines).strip()
    combined_text = "\n".join(text for text in (question_text, marking_text) if text).strip()
    return AssessmentStructureSection(
        label=label,
        max_mark=_extract_max_mark_text(header),
        weight_text=_extract_weight_text(header),
        question_text_exact=question_text,
        marking_instructions_exact=marking_text,
        anchor_phrases=_extract_anchor_lines(combined_text),
        evidence_expectations=_extract_evidence_lines(combined_text),
        dependency_labels=_extract_dependency_labels(combined_text, current_label=label),
        children=children,
    )


def _merge_assessment_structure_sections(
    sections: tuple[AssessmentStructureSection, ...],
) -> tuple[AssessmentStructureSection, ...]:
    merged: dict[str, AssessmentStructureSection] = {}
    order: list[str] = []
    for section in sections:
        key = _normalize_label_key(section.label)
        if key not in merged:
            merged[key] = section
            order.append(key)
            continue
        merged[key] = _merge_assessment_structure_section_pair(merged[key], section)
    return tuple(merged[key] for key in order if merged[key].label)


def _reconcile_assessment_structure_sections(
    prompt_sections: tuple[AssessmentStructureSection, ...],
    marking_scheme_sections: tuple[AssessmentStructureSection, ...],
) -> tuple[AssessmentStructureSection, ...]:
    scheme_by_key = {_normalize_label_key(section.label): section for section in marking_scheme_sections}
    prompt_by_key = {_normalize_label_key(section.label): section for section in prompt_sections}
    order: list[str] = []
    for section in marking_scheme_sections:
        key = _normalize_label_key(section.label)
        if key not in order:
            order.append(key)
    for section in prompt_sections:
        key = _normalize_label_key(section.label)
        if key not in order:
            order.append(key)

    reconciled: list[AssessmentStructureSection] = []
    for key in order:
        prompt_section = prompt_by_key.get(key)
        scheme_section = scheme_by_key.get(key)
        if prompt_section is not None and scheme_section is not None:
            reconciled.append(_reconcile_assessment_structure_section_pair(prompt_section, scheme_section))
        elif prompt_section is not None:
            reconciled.append(prompt_section)
        elif scheme_section is not None:
            reconciled.append(scheme_section)
    return tuple(section for section in reconciled if section.label)


def _reconcile_assessment_structure_section_pair(
    prompt_section: AssessmentStructureSection,
    scheme_section: AssessmentStructureSection,
) -> AssessmentStructureSection:
    return AssessmentStructureSection(
        label=scheme_section.label or prompt_section.label,
        parent_label=scheme_section.parent_label or prompt_section.parent_label,
        max_mark=scheme_section.max_mark if scheme_section.max_mark is not None else prompt_section.max_mark,
        weight_text=scheme_section.weight_text or prompt_section.weight_text,
        question_text_exact=_prefer_canonical_section_question_text(
            scheme_section.question_text_exact,
            prompt_section.question_text_exact,
        ),
        marking_instructions_exact=_prefer_richer_structure_text(
            scheme_section.marking_instructions_exact,
            prompt_section.marking_instructions_exact,
        ),
        anchor_phrases=_merge_unique_text_items(prompt_section.anchor_phrases, scheme_section.anchor_phrases),
        evidence_expectations=_merge_unique_text_items(
            prompt_section.evidence_expectations,
            scheme_section.evidence_expectations,
        ),
        dependency_labels=_merge_unique_text_items(
            prompt_section.dependency_labels,
            scheme_section.dependency_labels,
            current_label=prompt_section.label or scheme_section.label,
        ),
        children=_reconcile_structure_children(prompt_section.children, scheme_section.children),
    )


def _reconcile_structure_children(
    prompt_children: tuple[AssessmentStructureSection, ...],
    scheme_children: tuple[AssessmentStructureSection, ...],
) -> tuple[AssessmentStructureSection, ...]:
    scheme_by_key = {_normalize_label_key(child.label): child for child in scheme_children}
    prompt_by_key = {_normalize_label_key(child.label): child for child in prompt_children}
    order: list[str] = []
    for child in scheme_children:
        key = _normalize_label_key(child.label)
        if key not in order:
            order.append(key)
    for child in prompt_children:
        key = _normalize_label_key(child.label)
        if key not in order:
            order.append(key)

    merged_children: list[AssessmentStructureSection] = []
    for key in order:
        prompt_child = prompt_by_key.get(key)
        scheme_child = scheme_by_key.get(key)
        if prompt_child is not None and scheme_child is not None:
            merged_children.append(_reconcile_assessment_structure_section_pair(prompt_child, scheme_child))
        elif prompt_child is not None:
            merged_children.append(prompt_child)
        elif scheme_child is not None:
            merged_children.append(scheme_child)
    return tuple(child for child in merged_children if child.label)


def _merge_assessment_structure_section_pair(
    existing: AssessmentStructureSection,
    incoming: AssessmentStructureSection,
) -> AssessmentStructureSection:
    question_text = _prefer_richer_structure_text(existing.question_text_exact, incoming.question_text_exact)
    marking_text = _prefer_richer_structure_text(existing.marking_instructions_exact, incoming.marking_instructions_exact)
    combined_text = "\n".join(text for text in (question_text, marking_text) if text).strip()
    return AssessmentStructureSection(
        label=existing.label,
        parent_label=existing.parent_label or incoming.parent_label,
        max_mark=existing.max_mark if existing.max_mark is not None else incoming.max_mark,
        weight_text=existing.weight_text or incoming.weight_text,
        question_text_exact=question_text,
        marking_instructions_exact=marking_text,
        anchor_phrases=_merge_unique_text_items(existing.anchor_phrases, incoming.anchor_phrases),
        evidence_expectations=_merge_unique_text_items(existing.evidence_expectations, incoming.evidence_expectations),
        dependency_labels=_merge_unique_text_items(
            existing.dependency_labels,
            incoming.dependency_labels,
            current_label=existing.label,
        ),
        children=_merge_structure_children(existing.children, incoming.children),
    )


def _prefer_richer_structure_text(existing: str, incoming: str) -> str:
    existing_text = existing.strip()
    incoming_text = incoming.strip()
    if not existing_text:
        return incoming_text
    if not incoming_text:
        return existing_text
    if existing_text == incoming_text:
        return existing_text
    if len(incoming_text.splitlines()) > len(existing_text.splitlines()):
        return incoming_text
    if len(incoming_text) > len(existing_text) and len(incoming_text.splitlines()) >= len(existing_text.splitlines()):
        return incoming_text
    return existing_text


def _prefer_canonical_section_question_text(scheme_text: str, prompt_text: str) -> str:
    scheme_clean = scheme_text.strip()
    prompt_clean = prompt_text.strip()
    if not scheme_clean:
        return prompt_clean
    if not prompt_clean:
        return scheme_clean
    if _looks_header_only_question_text(scheme_clean) and not _looks_header_only_question_text(prompt_clean):
        return prompt_clean
    if _looks_damaged_question_text(scheme_clean) and not _looks_damaged_question_text(prompt_clean):
        return prompt_clean
    return scheme_clean


def _looks_header_only_question_text(text: str) -> bool:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return True
    if len(lines) == 1 and len(lines[0]) <= 40:
        return True
    return len(lines) == 1 and bool(_TOP_LEVEL_CONTEXT_PATTERN.match(lines[0]))


def _looks_damaged_question_text(text: str) -> bool:
    normalized = " ".join(text.split())
    if "..." in normalized:
        return True
    # CamelCase merges: e.g. "conditionsOnThe" — two or more hits indicates
    # words from different case contexts have been joined by the extractor.
    if len(re.findall(r"[a-z]{6,}[A-Z][a-z]+", normalized)) >= 2:
        return True
    # All-lowercase merges: a run of 16+ consecutive lowercase characters
    # typically means two or more words have been joined without a space.
    # Threshold at 16 avoids most valid single English words.
    return bool(re.search(r"[a-z]{16,}", normalized))


def _question_needs_richer_local_context(question_text: str) -> bool:
    cleaned = _clean_question_text_for_goal(question_text)
    if not cleaned:
        return False
    numbered_items = _extract_numbered_goal_items(cleaned)
    if len(numbered_items) >= 2:
        return True
    lowered = cleaned.lower()
    phrases = (
        "for each",
        "give an example",
        "explain why",
        "define each",
        "state the hypothesis",
        "more likely than",
    )
    return any(phrase in lowered for phrase in phrases)


def _section_text_looks_damaged_or_incomplete(section_text: str, question_text: str = "") -> bool:
    normalized = " ".join(section_text.split())
    if not normalized:
        return True
    if _looks_damaged_question_text(normalized):
        return True
    if len(normalized) < 120:
        return True
    if _question_needs_richer_local_context(question_text) and len(normalized) < 260:
        return True
    return False


def _expand_local_chunk(source_text: str, start: int, end: int, question_text: str = "") -> str:
    if not source_text:
        return ""
    base_padding = 260
    if _question_needs_richer_local_context(question_text):
        base_padding = 420
    expanded_start = max(0, start - base_padding)
    expanded_end = min(len(source_text), end + base_padding)
    if expanded_start > 0:
        newline_before = source_text.rfind("\n", expanded_start, start)
        if newline_before != -1:
            expanded_start = newline_before + 1
    if expanded_end < len(source_text):
        newline_after = source_text.find("\n", end, expanded_end)
        if newline_after != -1:
            expanded_end = newline_after
    return source_text[expanded_start:expanded_end].strip()


def _select_local_chunk_with_fallback(source_text: str, start: int, end: int, question_text: str = "") -> str:
    section_text = source_text[start:end].strip()
    if not _section_text_looks_damaged_or_incomplete(section_text, question_text):
        return section_text
    expanded = _expand_local_chunk(source_text, start, end, question_text)
    return expanded or section_text


def _merge_unique_text_items(
    existing: tuple[str, ...],
    incoming: tuple[str, ...],
    current_label: str = "",
) -> tuple[str, ...]:
    merged: list[str] = []
    seen: set[str] = set()
    current_key = _normalize_label_key(current_label) if current_label else ""
    for item in (*existing, *incoming):
        text = item.strip()
        if not text:
            continue
        normalized = _normalize_label_key(text)
        if normalized == current_key or normalized in seen:
            continue
        merged.append(text)
        seen.add(normalized)
    return tuple(merged)


def _merge_structure_children(
    existing: tuple[AssessmentStructureSection, ...],
    incoming: tuple[AssessmentStructureSection, ...],
) -> tuple[AssessmentStructureSection, ...]:
    merged: dict[str, AssessmentStructureSection] = {}
    order: list[str] = []
    for child in (*existing, *incoming):
        key = _normalize_label_key(child.label)
        if key not in merged:
            merged[key] = child
            order.append(key)
            continue
        merged[key] = _merge_assessment_structure_section_pair(merged[key], child)
    return tuple(merged[key] for key in order if merged[key].label)


def _split_question_and_marking_lines(lines: list[str]) -> tuple[list[str], list[str]]:
    question_lines: list[str] = []
    marking_lines: list[str] = []
    in_marking = False
    capture_continuation = False
    for line in lines:
        lowered = line.lower()
        if lowered.startswith("answer:"):
            in_marking = True
            remainder = line.split(":", 1)[1].strip()
            if remainder and _is_marking_instruction_line(remainder):
                marking_lines.append(remainder)
                capture_continuation = True
            continue
        if in_marking:
            if _is_marking_instruction_line(line) or (capture_continuation and line.startswith(("-", "•"))):
                marking_lines.append(line)
                capture_continuation = True
            else:
                capture_continuation = False
            continue
        question_lines.append(line)
    return question_lines, marking_lines


def _extract_structure_children(parent_label: str, question_lines: list[str]) -> tuple[AssessmentStructureSection, ...]:
    if not question_lines:
        return ()
    groups: list[list[str]] = []
    current: list[str] = []
    for line in question_lines:
        if _SUBPART_CONTEXT_PATTERN.match(line):
            if current:
                groups.append(current)
            current = [line]
        elif current:
            current.append(line)
    if current:
        groups.append(current)
    if len(groups) < 2:
        return ()
    children: list[AssessmentStructureSection] = []
    for group in groups:
        first_line = group[0]
        match = _SUBPART_CONTEXT_PATTERN.match(first_line)
        if match is None:
            continue
        token = _normalize_subpart_token(match.group("token"))
        child_label = _build_context_subpart_label(parent_label, token)
        question_text = "\n".join(group).strip()
        children.append(
            AssessmentStructureSection(
                label=child_label,
                parent_label=parent_label,
                max_mark=_extract_max_mark_text(first_line),
                weight_text=_extract_weight_text(first_line),
                question_text_exact=question_text,
                marking_instructions_exact="",
                anchor_phrases=_extract_anchor_lines(question_text),
                evidence_expectations=_extract_evidence_lines(question_text),
                dependency_labels=_extract_dependency_labels(question_text, current_label=child_label),
                children=(),
            )
        )
    return tuple(children)


def _canonical_context_label(label: str) -> str:
    parts = label.split()
    if len(parts) < 2:
        return label.strip()
    kind = parts[0].title()
    token = parts[1].upper() if parts[1].isalpha() else parts[1]
    return f"{kind} {token}"


def _is_top_level_context_header(line: str, tail: str) -> bool:
    if _extract_max_mark_text(line) is not None or _extract_weight_text(line):
        return True
    normalized_tail = " ".join(tail.split()).strip(" :.-")
    if not normalized_tail:
        return True
    return len(line) <= 32


def _normalize_subpart_token(token: str) -> str:
    cleaned = token.strip().strip("()[]").rstrip(".").rstrip(")").strip()
    return cleaned.upper() if cleaned.isalpha() else cleaned


def _build_context_subpart_label(parent_label: str, token: str) -> str:
    if token.isdigit():
        return f"{parent_label} Q{token}"
    return f"{parent_label} {token}"


def _extract_max_mark_text(text: str) -> float | None:
    match = _WEIGHT_TEXT_PATTERN.search(text)
    if match is None:
        return None
    value = match.group("value") or match.group("alt")
    unit = (match.group("unit") or match.group("alt_unit") or "").lower()
    if not value or unit == "%":
        return None
    return float(value)


def _extract_weight_text(text: str) -> str:
    match = _WEIGHT_TEXT_PATTERN.search(text)
    if match is None:
        return ""
    value = match.group("value") or match.group("alt") or ""
    unit = match.group("unit") or match.group("alt_unit") or ""
    if not value or not unit:
        return ""
    return f"{value} {unit}".strip()


def _is_marking_instruction_line(line: str) -> bool:
    lowered = line.strip().lower()
    if not lowered:
        return False
    starters = (
        "students should",
        "student should",
        "marks should",
        "mark should",
        "for full mark",
        "for a high grade",
        "points that should",
        "key elements",
        "references should",
        "references to",
        "note:",
        "high grade",
    )
    if lowered.startswith(starters):
        return True
    phrases = (
        "marks should be allocated",
        "should be allocated",
        "are expected to",
        "should mention",
        "should provide",
        "for full mark",
        "high grade would",
        "key elements for a high grade",
    )
    return any(phrase in lowered for phrase in phrases)


def _extract_anchor_lines(text: str) -> tuple[str, ...]:
    cues = (
        "use ",
        "using ",
        "comment on whether",
        "discuss",
        "explain",
        "identify",
        "support",
        "qualify",
        "contradict",
        "rewrite",
        "in your own",
        "use the results",
    )
    anchors: list[str] = []
    for line in [piece.strip() for piece in text.splitlines() if piece.strip()]:
        lowered = line.lower()
        if any(cue in lowered for cue in cues):
            anchors.append(line)
        if len(anchors) >= 4:
            break
    return tuple(anchors)


def _extract_evidence_lines(text: str) -> tuple[str, ...]:
    cues = (
        "quote",
        "quotes",
        "page",
        "slide",
        "reference",
        "references",
        "data source",
        "link",
        "evidence",
        "report",
        "course material",
        "lecture",
    )
    evidence_lines: list[str] = []
    for line in [piece.strip() for piece in text.splitlines() if piece.strip()]:
        lowered = line.lower()
        if any(cue in lowered for cue in cues):
            evidence_lines.append(line)
        if len(evidence_lines) >= 4:
            break
    return tuple(evidence_lines)


def _extract_dependency_labels(text: str, current_label: str = "") -> tuple[str, ...]:
    labels: list[str] = []
    current_key = _normalize_label_key(current_label) if current_label else ""
    seen_keys: set[str] = set()
    for match in _DEPENDENCY_LABEL_PATTERN.finditer(text):
        kind = match.group("kind").title()
        identifier = match.group("id")
        child = match.group("child")
        label = f"{kind} {identifier}"
        if child:
            label = f"{label} Q{child}"
        normalized = _normalize_label_key(label)
        if normalized == current_key or normalized in seen_keys:
            continue
        labels.append(label)
        seen_keys.add(normalized)
    return tuple(labels)


def _build_structure_marking_guidance(
    section: AssessmentStructureSection,
    parent: AssessmentStructureSection | None = None,
) -> str:
    lines: list[str] = []

    marking_text = section.marking_instructions_exact.strip()
    if marking_text:
        lines.append("Marking instructions:")
        lines.append(marking_text)

    if parent is not None:
        parent_marking = parent.marking_instructions_exact.strip()
        if parent_marking and parent_marking != marking_text:
            lines.append("Parent marking instructions:")
            lines.append(parent_marking)

    if section.anchor_phrases:
        lines.append("Anchor phrases:")
        lines.extend(f"- {phrase}" for phrase in section.anchor_phrases)

    if section.evidence_expectations:
        lines.append("Evidence expectations:")
        lines.extend(f"- {expectation}" for expectation in section.evidence_expectations)

    if section.dependency_labels:
        lines.append(f"Dependency links: {', '.join(section.dependency_labels)}")

    if not marking_text and parent is not None and section.question_text_exact.strip() and section.max_mark is None:
        task_type = infer_task_type(section.label, section.question_text_exact)
        grading_mode = infer_grading_mode(
            section.label,
            parent.marking_instructions_exact or section.question_text_exact,
            task_type=task_type,
        )
        lines.append("Child-specific scoring guidance:")
        lines.append(
            build_unit_rubric_text(
                label=section.label,
                max_mark=section.max_mark,
                grading_mode=grading_mode,
                task_type=task_type,
                question_text_exact=section.question_text_exact,
                marking_guidance=section.question_text_exact,
            )
        )

    if not lines:
        if section.question_text_exact.strip():
            return section.question_text_exact.strip()
        if parent is not None and parent.question_text_exact.strip():
            return parent.question_text_exact.strip()
        return ""
    return "\n".join(lines).strip()


def _resolve_child_task_type(parent: AssessmentStructureSection, child: AssessmentStructureSection) -> str:
    child_task_type = infer_task_type(child.label, child.question_text_exact or parent.question_text_exact)
    parent_task_type = infer_task_type(parent.label, parent.question_text_exact or parent.marking_instructions_exact)
    generic_child_types = {"evaluative_discussion", "deterministic_derivation", "whole_submission_holistic"}
    if (
        child_task_type in generic_child_types
        and parent_task_type not in generic_child_types
    ):
        return parent_task_type
    return child_task_type


def _resolve_section_children_with_marks(section: AssessmentStructureSection) -> tuple[AssessmentStructureSection, ...]:
    children = list(section.children)
    if len(children) < 2 or section.max_mark is None:
        return tuple(children)
    marked_total = sum(float(child.max_mark) for child in children if child.max_mark is not None)
    unresolved_children = [child for child in children if child.max_mark is None]
    if not unresolved_children:
        return tuple(children)
    remaining = float(section.max_mark) - marked_total
    if remaining <= 0:
        return tuple(children)
    per_child = round(remaining / len(unresolved_children), 2)
    resolved: list[AssessmentStructureSection] = []
    for child in children:
        if child.max_mark is not None:
            resolved.append(child)
            continue
        resolved.append(replace(child, max_mark=per_child))
    return tuple(resolved)


def extract_expected_parts_from_context(context: MarkingContext) -> list[SubmissionPart]:
    sections = extract_assessment_structure(context)
    parts: list[SubmissionPart] = []
    for section in sections:
        marking_guidance = _build_structure_marking_guidance(section)
        focus_hint = next((line for line in section.question_text_exact.splitlines()[1:] if line.strip()), section.label)
        parts.append(
            SubmissionPart(
                label=section.label,
                focus_hint=focus_hint.strip(),
                anchor_text=section.label,
                max_mark=section.max_mark,
                marking_guidance=marking_guidance,
                question_text_exact=section.question_text_exact,
            )
        )
    return parts


def extract_expected_subparts_from_context(
    context: MarkingContext,
    require_marks: bool = True,
) -> dict[str, list[SubmissionPart]]:
    child_map: dict[str, list[SubmissionPart]] = {}
    for section in extract_assessment_structure(context):
        resolved_children = _resolve_section_children_with_marks(section)
        safe_children = [
            child
            for child in resolved_children
            if (child.max_mark is not None) or not require_marks
        ]
        if len(safe_children) < 2:
            continue
        child_map[_normalize_label_key(section.label)] = [
            SubmissionPart(
                label=child.label,
                focus_hint=child.question_text_exact.splitlines()[0].strip() if child.question_text_exact.strip() else child.label,
                anchor_text=child.label.split()[-1] if " Q" in child.label else child.label,
                max_mark=child.max_mark,
                marking_guidance=_build_structure_marking_guidance(child, parent=section),
                question_text_exact=child.question_text_exact,
                task_type=_resolve_child_task_type(section, child),
                criterion_mode=infer_criterion_mode(_resolve_child_task_type(section, child)),
            )
            for child in safe_children
        ]
    return child_map


def build_submission_diagnostics(script_text: str, parts: list[SubmissionPart] | None = None) -> SubmissionDiagnostics:
    words = script_text.split()
    extracted_word_count = len(words)
    resolved_parts = parts or [SubmissionPart(label="Whole Submission", focus_hint="Assess the full script as one piece of work.")]
    possible_extraction_issue = bool(re.search(r"(.)\1{8,}", script_text)) or extracted_word_count < 80
    return SubmissionDiagnostics(
        extracted_word_count=extracted_word_count,
        detected_part_count=len(resolved_parts),
        detected_part_labels=tuple(part.label for part in resolved_parts),
        low_text=extracted_word_count < 120,
        possible_extraction_issue=possible_extraction_issue,
    )


def split_submission_into_parts(script_text: str) -> list[SubmissionPart]:
    return [SubmissionPart(label="Whole Submission", focus_hint="Assess the full script as one piece of work.")]


def build_assessment_map_cache_key(
    context: MarkingContext,
    verifier_model_name: str | None = None,
) -> str:
    payload = {
        "rubric_text": context.rubric_text,
        "brief_text": context.brief_text,
        "marking_scheme_text": context.marking_scheme_text,
        "graded_sample_text": context.graded_sample_text,
        "other_context_text": context.other_context_text,
        "max_mark": context.max_mark,
        "verifier_model_name": verifier_model_name or "",
    }
    serialized = json.dumps(payload, ensure_ascii=True, sort_keys=True)
    return hashlib.sha256(serialized.encode("utf-8")).hexdigest()


def clear_prepared_assessment_map_cache() -> None:
    _PREPARED_ASSESSMENT_CACHE.clear()


def _is_verifier_failure_issue(issue: str) -> bool:
    return issue.strip().startswith("rubric_verifier_failed:")


def _collect_preparation_issues(assessment_map: AssessmentMap) -> tuple[str, ...]:
    issues: list[str] = []
    for unit in assessment_map.units:
        issues.extend(issue for issue in unit.rubric_issues if issue.strip())
    return tuple(issues)


def _compute_preparation_confidence(assessment_map: AssessmentMap) -> float | None:
    unit_confidences = [
        unit.rubric_confidence_0_to_100
        for unit in assessment_map.units
        if unit.rubric_confidence_0_to_100 is not None
    ]
    if not unit_confidences:
        return None
    return round(sum(unit_confidences) / len(unit_confidences), 2)


def _unit_refinement_score(unit: AssessmentUnit) -> float:
    confidence = unit.rubric_confidence_0_to_100 if unit.rubric_confidence_0_to_100 is not None else 0.0
    issue_penalty = float(len(unit.rubric_issues))
    failure_penalty = 25.0 if any(_is_verifier_failure_issue(issue) for issue in unit.rubric_issues) else 0.0
    return round(confidence - issue_penalty - failure_penalty, 2)


def _summarize_unit_refinement(previous: AssessmentUnit, candidate: AssessmentUnit, attempt_number: int) -> str:
    previous_confidence = previous.rubric_confidence_0_to_100 if previous.rubric_confidence_0_to_100 is not None else 0.0
    candidate_confidence = candidate.rubric_confidence_0_to_100 if candidate.rubric_confidence_0_to_100 is not None else 0.0
    previous_issue_count = len(previous.rubric_issues)
    candidate_issue_count = len(candidate.rubric_issues)
    changed_text = previous.rubric_text.strip() != candidate.rubric_text.strip()
    text_change = "text_changed" if changed_text else "text_unchanged"
    return (
        f"attempt_{attempt_number}: confidence {previous_confidence:.2f}->{candidate_confidence:.2f}, "
        f"issues {previous_issue_count}->{candidate_issue_count}, {text_change}"
    )


def _is_acceptable_prepared_map(assessment_map: AssessmentMap) -> bool:
    confidences = [
        unit.rubric_confidence_0_to_100
        for unit in assessment_map.units
        if unit.rubric_confidence_0_to_100 is not None
    ]
    if not confidences:
        return True
    if any(any(_is_verifier_failure_issue(issue) for issue in unit.rubric_issues) for unit in assessment_map.units):
        return False
    return min(confidences) >= MIN_ACCEPTABLE_PREPARATION_CONFIDENCE


def _build_prepared_assessment_units(
    assessment_map: AssessmentMap,
    verifier_attempt_count: int,
    accepted_refinements: dict[str, list[str]],
    rejected_refinements: dict[str, list[str]],
) -> tuple[PreparedAssessmentUnit, ...]:
    prepared_units: list[PreparedAssessmentUnit] = []
    for unit in assessment_map.units:
        unit_key = _normalize_label_key(unit.label)
        prepared_units.append(
            PreparedAssessmentUnit(
                label=unit.label,
                compact_criteria=tuple(_extract_classification_criteria(unit)),
                verifier_confidence_0_to_100=unit.rubric_confidence_0_to_100,
                verifier_issues=tuple(issue for issue in unit.rubric_issues if issue.strip()),
                verifier_attempt_count=verifier_attempt_count,
                verifier_accepted_refinements=tuple(accepted_refinements.get(unit_key, [])),
                verifier_rejected_refinements=tuple(rejected_refinements.get(unit_key, [])),
            )
        )
    return tuple(prepared_units)


def _merge_refined_assessment_maps(
    current_map: AssessmentMap,
    candidate_map: AssessmentMap,
    attempt_number: int,
    accepted_refinements: dict[str, list[str]],
    rejected_refinements: dict[str, list[str]],
) -> AssessmentMap:
    merged_units: list[AssessmentUnit] = []
    for current_unit, candidate_unit in zip(current_map.units, candidate_map.units):
        unit_key = _normalize_label_key(current_unit.label)
        changed = (
            current_unit.rubric_text.strip() != candidate_unit.rubric_text.strip()
            or current_unit.rubric_confidence_0_to_100 != candidate_unit.rubric_confidence_0_to_100
            or current_unit.rubric_issues != candidate_unit.rubric_issues
        )
        if _unit_refinement_score(candidate_unit) > _unit_refinement_score(current_unit):
            merged_units.append(candidate_unit)
            if changed:
                accepted_refinements[unit_key].append(
                    _summarize_unit_refinement(current_unit, candidate_unit, attempt_number)
                )
            continue
        merged_units.append(current_unit)
        if changed:
            rejected_refinements[unit_key].append(
                _summarize_unit_refinement(current_unit, candidate_unit, attempt_number)
            )
    return replace(current_map, units=tuple(merged_units))


def prepare_assessment_map(
    context: MarkingContext,
    verifier_model_name: str | None = None,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> PreparedAssessmentMap:
    cache_key = build_assessment_map_cache_key(
        context=context,
        verifier_model_name=verifier_model_name,
    )
    cached = _PREPARED_ASSESSMENT_CACHE.get(cache_key)
    if cached is not None:
        return cached

    started = time.perf_counter()
    original_map = build_assessment_map(context)
    prepared_map = original_map
    verification_applied = False
    verifier_attempt_count = 0
    accepted_refinements: dict[str, list[str]] = defaultdict(list)
    rejected_refinements: dict[str, list[str]] = defaultdict(list)

    if verifier_model_name:
        verification_applied = True
        prepared_map = original_map
        for attempt_number in range(1, MAX_PREPARED_MAP_VERIFIER_ATTEMPTS + 1):
            verifier_attempt_count = attempt_number
            candidate_map = verify_assessment_rubrics(
                assessment_map=prepared_map,
                context=context,
                verifier_model_name=verifier_model_name,
                ollama_url=ollama_url,
            )
            if attempt_number == 1:
                prepared_map = candidate_map
            else:
                prepared_map = _merge_refined_assessment_maps(
                    current_map=prepared_map,
                    candidate_map=candidate_map,
                    attempt_number=attempt_number,
                    accepted_refinements=accepted_refinements,
                    rejected_refinements=rejected_refinements,
                )
            if _is_acceptable_prepared_map(prepared_map):
                break
    preparation_confidence = _compute_preparation_confidence(prepared_map)
    preparation_issues = _collect_preparation_issues(prepared_map)
    preparation_latency_seconds = round(time.perf_counter() - started, 2)

    prepared = PreparedAssessmentMap(
        cache_key=cache_key,
        original_map=original_map,
        prepared_map=prepared_map,
        prepared_units=_build_prepared_assessment_units(
            assessment_map=prepared_map,
            verifier_attempt_count=verifier_attempt_count,
            accepted_refinements=accepted_refinements,
            rejected_refinements=rejected_refinements,
        ),
        verifier_model_name=verifier_model_name,
        verification_applied=verification_applied,
        verifier_attempt_count=verifier_attempt_count,
        preparation_issues=preparation_issues,
        preparation_confidence_0_to_100=preparation_confidence,
        preparation_latency_seconds=preparation_latency_seconds,
    )
    _PREPARED_ASSESSMENT_CACHE[cache_key] = prepared
    return prepared


def build_assessment_map(context: MarkingContext) -> AssessmentMap:
    sections = extract_assessment_structure(context)

    units: list[AssessmentUnit] = []
    if sections:
        for section in sections:
            marked_children = [child for child in _resolve_section_children_with_marks(section) if child.max_mark is not None]
            if len(marked_children) >= 2:
                for child in marked_children:
                    marking_guidance = _build_structure_marking_guidance(child, parent=section)
                    question_text = child.question_text_exact or section.question_text_exact or marking_guidance or child.label
                    task_type = infer_task_type(child.label, question_text)
                    criterion_mode = infer_criterion_mode(task_type)
                    mode = infer_grading_mode(
                        child.label,
                        marking_guidance or section.marking_instructions_exact or section.question_text_exact or child.label,
                        task_type=task_type,
                    )
                    dependency_source = "\n".join(
                        text for text in (question_text, marking_guidance, section.question_text_exact) if text.strip()
                    )
                    dependency = infer_dependency_group(child.label, dependency_source, parent_label=section.label)
                    units.append(
                        AssessmentUnit(
                            label=child.label,
                            max_mark=child.max_mark,
                            parent_label=section.label,
                            grading_mode=mode,
                            task_type=task_type,
                            criterion_mode=criterion_mode,
                            dependency_group=dependency,
                            question_text_exact=question_text,
                            marking_guidance=marking_guidance,
                            rubric_text=build_unit_rubric_text(
                                label=child.label,
                                max_mark=child.max_mark,
                                grading_mode=mode,
                                task_type=task_type,
                                question_text_exact=question_text,
                                marking_guidance=marking_guidance,
                            ),
                        )
                    )
            else:
                marking_guidance = _build_structure_marking_guidance(section)
                question_text = section.question_text_exact or marking_guidance or section.label
                task_type = infer_task_type(section.label, question_text)
                criterion_mode = infer_criterion_mode(task_type)
                mode = infer_grading_mode(section.label, marking_guidance or section.label, task_type=task_type)
                dependency_source = "\n".join(text for text in (question_text, marking_guidance) if text.strip())
                dependency = infer_dependency_group(section.label, dependency_source, parent_label=section.label)
                units.append(
                    AssessmentUnit(
                        label=section.label,
                        max_mark=section.max_mark,
                        parent_label="",
                        grading_mode=mode,
                        task_type=task_type,
                        criterion_mode=criterion_mode,
                        dependency_group=dependency,
                        question_text_exact=question_text,
                        marking_guidance=marking_guidance,
                        rubric_text=build_unit_rubric_text(
                            label=section.label,
                            max_mark=section.max_mark,
                            grading_mode=mode,
                            task_type=task_type,
                            question_text_exact=question_text,
                            marking_guidance=marking_guidance,
                        ),
                    )
                )

    if not units:
        units.append(
            AssessmentUnit(
                label="Whole Submission",
                max_mark=context.max_mark,
                grading_mode="analytical",
                task_type="whole_submission_holistic",
                criterion_mode="abstract",
                dependency_group="",
                question_text_exact="Assess the full submission as one unit.",
                marking_guidance="Assess the full submission as one unit.",
                rubric_text=build_unit_rubric_text(
                    label="Whole Submission",
                    max_mark=context.max_mark,
                    grading_mode="analytical",
                    task_type="whole_submission_holistic",
                    question_text_exact="Assess the full submission as one unit.",
                    marking_guidance="Assess the full submission as one unit.",
                ),
            )
        )

    scale_confidence = 95.0 if any(unit.max_mark is not None for unit in units) else 60.0
    return AssessmentMap(
        units=tuple(units),
        overall_max_mark=context.max_mark,
        scale_confidence_0_to_100=scale_confidence,
    )


def build_rubric_matrix_markdown(context: MarkingContext, assessment_name: str = "Assessment") -> str:
    assessment_map = build_assessment_map(context)
    lines = [
        f"# {assessment_name} Rubric Matrix",
        "",
        f"Overall maximum mark: {_format_number(context.max_mark)}",
        "",
        "## Policy",
        "",
        "- This rubric is a classification matrix, not a transcript of the marking scheme.",
        "- Classification happens at part or subpart level first; the final mark is deterministic arithmetic from those unit scores.",
        "- Analytical units use UK-style bands. Deterministic units use evidence bands tied to correctness and method.",
        "- Within-band placement must be justified by concrete evidence from the script, not by generic prose.",
        "",
    ]
    grouped_units: dict[str, list[AssessmentUnit]] = {}
    for unit in assessment_map.units:
        grouped_units.setdefault(unit.parent_label or unit.label, []).append(unit)

    for group_label, units in grouped_units.items():
        lines.extend([f"## {group_label}", ""])
        for unit in units:
            lines.extend(_render_unit_rubric_matrix(unit))
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_rubric_verification_messages(unit: AssessmentUnit, context: MarkingContext) -> list[dict[str, str]]:
    max_mark_text = _format_number(unit.max_mark) if unit.max_mark is not None else _format_number(context.max_mark)
    structure_context = extract_structure_guidance(context)
    system = (
        "You are verifying and enriching one grading rubric unit before marking begins. "
        "You must preserve the unit label, max mark, grading mode, and dependency structure. "
        "You may only improve the rubric prose so it is clearer, more discriminating, and better aligned with the assessment documents. "
        'Return JSON with keys "rubric_text", "issues", and "confidence_0_to_100".'
    )
    user = (
        f"Unit label: {unit.label}\n"
        f"Parent label: {unit.parent_label or 'None'}\n"
        f"Unit max mark: {max_mark_text}\n"
        f"Grading mode: {unit.grading_mode}\n"
        f"Dependency group: {unit.dependency_group or 'independent'}\n\n"
        f"Current rubric text:\n{unit.rubric_text or unit.marking_guidance or '(empty)'}\n\n"
        f"Structure and marks evidence from the assessment documents:\n{structure_context or '(none)'}\n\n"
        "Verification rules:\n"
        "1. Do not assign a score and do not write a final mark.\n"
        "2. Do not change the label, max mark, part structure, or dependency structure.\n"
        "3. Use the assessment documents only when they add useful information about structure, marks, weights, or question numbering.\n"
        "4. If the unit is deterministic, enforce full-range discrimination, partial credit, method marks, equation or algebra checks, interpretation checks, and avoid compressing marks into the middle.\n"
        "5. If the unit is analytical, enforce UK classification language and distinctions: First, 2:1, 2:2, Third/Pass, Fail, Missing/unfinished. Map those bands to the unit mark range and make the top band rare and evidence-based.\n"
        "6. If a rule is already good, keep it.\n"
        "7. Put any concerns or missing evidence in issues.\n"
        "8. confidence_0_to_100 must be a number from 0 to 100.\n"
        "9. Return only JSON."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def normalize_verified_rubric(data: dict[str, Any], unit: AssessmentUnit) -> dict[str, Any]:
    rubric_text = data.get("rubric_text")
    if not isinstance(rubric_text, str) or not rubric_text.strip():
        raise ValueError(f"Verifier did not return rubric_text for {unit.label}: {data}")

    issues = data.get("issues", [])
    if issues is None:
        issues = []
    if not isinstance(issues, list) or any(not isinstance(item, str) or not item.strip() for item in issues):
        raise ValueError(f"Verifier returned invalid issues for {unit.label}: {data}")

    confidence = data.get("confidence_0_to_100")
    if isinstance(confidence, bool) or not isinstance(confidence, (int, float)):
        raise ValueError(f"Verifier returned invalid confidence for {unit.label}: {data}")
    confidence_value = round(float(confidence), 2)
    if confidence_value < 0 or confidence_value > 100:
        raise ValueError(f"Verifier returned out-of-range confidence for {unit.label}: {data}")

    return {
        "rubric_text": rubric_text.strip(),
        "issues": tuple(item.strip() for item in issues if item.strip()),
        "confidence_0_to_100": confidence_value,
    }


def verify_assessment_rubrics(
    assessment_map: AssessmentMap,
    context: MarkingContext,
    verifier_model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> AssessmentMap:
    verified_units: list[AssessmentUnit] = []
    for unit in assessment_map.units:
        messages = build_rubric_verification_messages(unit, context)
        normalized: dict[str, Any] | None = None
        last_error: Exception | None = None
        for attempt in range(2):
            try:
                data = _call_ollama_json(
                    model_name=verifier_model_name,
                    messages=messages,
                    ollama_url=ollama_url,
                    profile="rubric_verification",
                )
                normalized = normalize_verified_rubric(data, unit)
                break
            except Exception as exc:
                last_error = exc
                if attempt == 0:
                    messages = _build_rubric_verification_retry_messages(messages, unit, exc)
                    continue

        if normalized is None:
            fallback_issues = list(unit.rubric_issues)
            if last_error is not None:
                fallback_issues.append(f"rubric_verifier_failed: {last_error}")
            verified_units.append(
                replace(
                    unit,
                    rubric_confidence_0_to_100=0.0,
                    rubric_issues=tuple(fallback_issues),
                )
            )
            continue

        verified_units.append(
            replace(
                unit,
                rubric_text=normalized["rubric_text"],
                rubric_confidence_0_to_100=normalized["confidence_0_to_100"],
                rubric_issues=normalized["issues"],
            )
        )

    return replace(assessment_map, units=tuple(verified_units))


def infer_task_type(label: str, question_text: str) -> str:
    text = " ".join(f"{label}\n{question_text}".lower().split())
    if (
        any(term in text for term in ("generate", "draft", "produce"))
        and any(term in text for term in ("summary", "answer", "response"))
        and any(term in text for term in ("comment on whether", "review", "critique", "evaluate accuracy", "rewrite", "revise", "edit", "correct"))
    ):
        return "critique_and_revision"
    if (
        any(term in text for term in ("write out", "write down", "state", "specify", "formulate"))
        and any(term in text for term in ("payoff function", "utility function", "objective function", "formal model", "formal specification"))
    ):
        return "model_specification"
    if (
        any(term in text for term in ("parameter", "parameters"))
        and any(term in text for term in ("more likely", "less likely", "increases", "decreases", "higher", "lower"))
        and any(term in text for term in ("scenario", "outcome", "equilibrium"))
    ):
        return "comparative_statics"
    if (
        any(term in text for term in ("regression equation", "empirical specification", "estimating equation", "econometric specification"))
        and any(term in text for term in ("write down", "state", "specify", "formulate"))
    ):
        return "regression_specification"
    if any(term in text for term in ("derive", "show that", "solve for", "prove that")) and any(term in text for term in ("condition", "equilibrium", "solution")):
        return "deterministic_derivation"
    if any(term in text for term in ("welfare", "well-being", "wellbeing", "efficiency")) and any(term in text for term in ("condition", "implication", "effect", "impact")):
        return "welfare_reasoning"
    if any(term in text for term in ("list", "state", "give", "identify")) and any(term in text for term in ("prediction", "predictions", "implication", "implications", "testable hypothesis")):
        return "prediction_generation"
    if (
        any(term in text for term in ("measure", "measured", "measurement", "operationalise", "operationalize", "indicator", "observable", "observed"))
        and any(term in text for term in ("data source", "data sources", "dataset", "datasets", "real world", "practice", "in practice"))
    ):
        return "measurement_and_data_design"
    if (
        any(term in text for term in ("endogeneity", "omitted variable", "reverse causality", "selection bias", "simultaneity", "confounding"))
    ):
        return "causal_identification"
    if (
        (
            any(term in text for term in ("source", "sources", "material", "materials", "article", "articles", "video", "videos", "case", "cases"))
            and any(term in text for term in ("claim", "claims", "argument", "arguments", "compare", "contrast", "synthes", "shed light on"))
        )
        or (
            any(term in text for term in ("model", "framework", "theory"))
            and any(term in text for term in ("missing from", "omission", "omissions", "limitations", "extend", "extension"))
        )
    ):
        return "synthesis_across_sources"
    if any(term in text for term in ("discuss", "evaluate", "comment on whether", "whether you think", "critically assess", "good idea")):
        return "evaluative_discussion"
    if any(phrase in text for phrase in ("explain how changing", "explain carefully why", "discuss the intuition", "explain how each term", "explain how each variable", "interpret the coefficient", "interpret the result", "explain the mechanism")):
        return "explanation_interpretation"
    if any(term in text for term in ("summary", "overview", "comment", "assess")):
        return "evaluative_discussion"
    return "deterministic_derivation"


def infer_criterion_mode(task_type: str) -> str:
    if task_type in {
        "critique_and_revision",
        "evaluative_discussion",
        "synthesis_across_sources",
        "whole_submission_holistic",
    }:
        return "abstract"
    return "deterministic"


def _task_type_needs_model_refinement(part: SubmissionPart) -> bool:
    if not part.question_text_exact.strip():
        return False
    if " Q" not in part.label:
        return False
    return part.task_type in {"", "evaluative_discussion", "deterministic_derivation", "whole_submission_holistic"}


def _build_task_type_refinement_messages(part: SubmissionPart) -> list[dict[str, str]]:
    allowed_task_types = [
        "critique_and_revision",
        "model_specification",
        "comparative_statics",
        "regression_specification",
        "deterministic_derivation",
        "welfare_reasoning",
        "prediction_generation",
        "measurement_and_data_design",
        "causal_identification",
        "synthesis_across_sources",
        "evaluative_discussion",
        "explanation_interpretation",
        "whole_submission_holistic",
    ]
    system = (
        "You are classifying one assessment subtask. "
        "Return exactly one JSON object with keys "
        '"task_type" and "reason". '
        "Choose the single best task_type from the allowed list only."
    )
    user = (
        f"Section label: {part.label}\n"
        f"Current task_type guess: {part.task_type or '(none)'}\n"
        f"Question text:\n{part.question_text_exact.strip()}\n\n"
        f"Source guidance:\n{part.marking_guidance.strip() or '(none)'}\n\n"
        "Allowed task_type values:\n"
        + "\n".join(f"- {item}" for item in allowed_task_types)
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def refine_part_task_types_with_model(
    parts: list[SubmissionPart],
    model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> list[SubmissionPart]:
    refined: list[SubmissionPart] = []
    for part in parts:
        if not _task_type_needs_model_refinement(part):
            refined.append(part)
            continue
        try:
            data = _call_ollama_json(
                model_name=model_name,
                messages=_build_task_type_refinement_messages(part),
                ollama_url=ollama_url,
                profile="structure_extraction",
            )
            task_type = str(data.get("task_type", "")).strip()
            if task_type:
                refined.append(
                    replace(
                        part,
                        task_type=task_type,
                        criterion_mode=infer_criterion_mode(task_type),
                    )
                )
                continue
        except Exception:
            pass
        refined.append(part)
    return refined


def _matches_all(text: str, *terms: str) -> bool:
    return all(term in text for term in terms)


def infer_grading_mode(label: str, marking_guidance: str, task_type: str = "") -> str:
    if task_type in {
        "critique_and_revision",
        "comparative_statics",
        "evaluative_discussion",
        "synthesis_across_sources",
        "whole_submission_holistic",
    }:
        if task_type == "comparative_statics":
            return "deterministic"
        return "analytical"
    if task_type in {
        "model_specification",
        "deterministic_derivation",
        "explanation_interpretation",
        "welfare_reasoning",
        "prediction_generation",
        "measurement_and_data_design",
        "regression_specification",
        "causal_identification",
    }:
        return "deterministic"
    text = f"{label}\n{marking_guidance}".lower()
    analytical_patterns = (
        "summary",
        "discuss",
        "evaluate",
        "commentary",
        "in your own words",
        "good idea",
        "critically assess",
        "justify your conclusion",
        "comment on whether",
        "compare and contrast",
    )
    deterministic_patterns = (
        "derive",
        "equation",
        "regression",
        "condition",
        "hypothesis",
        "measure",
        "data source",
        "calculate",
        "show carefully",
        "parameter",
        "endogeneity",
    )
    if any(pattern in text for pattern in deterministic_patterns):
        return "deterministic"
    if any(pattern in text for pattern in analytical_patterns):
        return "analytical"
    return "deterministic"


def infer_dependency_group(label: str, marking_guidance: str, parent_label: str = "") -> str:
    text = " ".join(marking_guidance.lower().split())
    direct_reference_patterns = (
        r"\busing\s+(?:your|the|previous|earlier)\s+(?:answer|result|results|work|analysis|model|derivation|prediction|predictions|evidence)\b",
        r"\busing\s+the\s+\w+\s+(?:model|results|answer|analysis)\s+from\s+(?:question|part|section)\s+\w+\b",
        r"\bfrom\s+(?:your|the|previous|earlier)\s+(?:answer|result|results|work|analysis|model|derivation|prediction|predictions|evidence)\b",
        r"\bbased on\s+(?:your|the|previous|earlier)\s+(?:answer|result|results|work|analysis|model|derivation|prediction|predictions|evidence)\b",
        r"\bgiven\s+(?:your|the)\s+(?:answer|result|results|work|analysis|model|derivation)\b",
        r"\bas\s+(?:shown|derived|established)\s+in\s+(?:your|the)\s+(?:answer|work|analysis)\b",
        r"\b(question|part|section)\s+\w+\s+(?:above|earlier|previously)\b",
        r"\byour\s+(?:answer|result|results|work|analysis)\s+to\s+(?:question|part|section)\s+\w+\b",
        r"\bprediction[s]?\s+from\s+your\s+answer\b",
        r"\bfrom\s+the\s+model\s+above\b",
    )
    if any(re.search(pattern, text) for pattern in direct_reference_patterns):
        return _normalize_label_key(parent_label or label)
    return ""


def build_unit_rubric_text(
    label: str,
    max_mark: float | None,
    grading_mode: str,
    task_type: str,
    question_text_exact: str,
    marking_guidance: str,
) -> str:
    unit = AssessmentUnit(
        label=label,
        max_mark=max_mark,
        grading_mode=grading_mode,
        task_type=task_type,
        criterion_mode=infer_criterion_mode(task_type),
        question_text_exact=question_text_exact,
        marking_guidance=marking_guidance,
    )
    if grading_mode == "analytical":
        return build_analytical_rubric_text(unit)
    return build_deterministic_rubric_text(unit)


def build_deterministic_rubric_text(unit: AssessmentUnit) -> str:
    max_text = _format_number(unit.max_mark) if unit.max_mark is not None else "the available marks"
    criteria = _extract_classification_criteria(unit)
    return "\n".join(
        [
            f"Unit: {unit.label}",
            f"Task focus: {_build_task_focus(unit)}.",
            f"Available marks: 0 to {max_text}.",
            "Use this rubric to classify and score the attempt:",
            *[f"- {criterion}" for criterion in criteria],
            "- Award explicit partial credit for correct setup, valid intermediate steps, and correct interpretation even when the final answer is incomplete.",
            "- Place work lower when key steps are missing, algebra is wrong, the method is unjustified, or the interpretation does not follow from the working.",
            "- Rank attempts within the same band by accuracy, completeness of working, and clarity of interpretation.",
        ]
    ).strip()


def build_analytical_rubric_text(unit: AssessmentUnit) -> str:
    max_text = _format_number(unit.max_mark) if unit.max_mark is not None else "the available marks"
    criteria = _extract_classification_criteria(unit)
    return "\n".join(
        [
            f"Unit: {unit.label}",
            f"Task focus: {_build_task_focus(unit)}.",
            f"Available marks: 0 to {max_text}.",
            "Use UK-style quality bands: First, 2:1, 2:2, Third/Pass, Fail, Missing/unfinished.",
            "Use this rubric to classify and score the response:",
            *[f"- {criterion}" for criterion in criteria],
            "- Place work higher when the response is more precise, better supported, more relevant to the task, and more analytically convincing.",
            "- Place work lower when the response is vague, generic, descriptive, unsupported, or only partially relevant to the task.",
            "- Use the top band only for genuinely strong, well-supported work; do not compress all adequate answers into the middle of the scale.",
        ]
    ).strip()


def _render_unit_rubric_matrix(unit: AssessmentUnit) -> list[str]:
    max_text = _format_number(unit.max_mark) if unit.max_mark is not None else "variable"
    guidance_summary = _summarize_guidance_for_matrix(unit.marking_guidance)
    criteria_sentences = _extract_classification_criteria(unit)
    lines = [
        f"### {unit.label}",
        "",
        f"- Max mark: {max_text}",
        f"- Mode: {unit.grading_mode}",
        f"- Dependency group: {unit.dependency_group or 'independent'}",
        f"- Core criterion: {guidance_summary}",
        f"- Classification question: {_build_classification_question(unit)}",
        f"- Ranking rule: {_build_ranking_rule(unit)}",
        "",
        "#### Criteria Sentences",
        "",
    ]
    for sentence in criteria_sentences:
        lines.append(f"- {sentence}")
    lines.extend([
        "",
        "| Band | Descriptor | Within-band placement |",
        "| --- | --- | --- |",
    ])
    for band, descriptor, anchor in _band_rows_for_unit(unit):
        lines.append(f"| {band} | {descriptor} | {anchor} |")
    return lines


def _band_rows_for_unit(unit: AssessmentUnit) -> list[tuple[str, str, str]]:
    criterion = _build_task_focus(unit)
    if unit.grading_mode == "analytical":
        return [
            ("Missing/unfinished", f"Little or no usable response on {criterion}.", "Low: fragmentary or absent. Mid: partial attempt with minimal relevance. High: unfinished but shows some relevant substance."),
            ("Fail", f"Material is relevant in places but weak, unsupported, or substantially underdeveloped on {criterion}.", "Low: largely unsupported. Mid: some relevant points. High: close to pass but still too thin or inaccurate."),
            ("Third/Pass", f"Adequate response on {criterion} with basic relevance and limited development.", "Low: mostly descriptive. Mid: some explanation. High: coherent pass with clearer support."),
            ("2:2", f"Reasonably competent response on {criterion} with some support, but uneven precision or depth.", "Low: competent but patchy. Mid: solid lower-second standard. High: close to 2:1 but still uneven."),
            ("2:1", f"Strong response on {criterion} with clear support, relevant explanation, and useful appraisal.", "Low: strong core but some gaps. Mid: convincing upper-second standard. High: close to First with only limited weaknesses."),
            ("First", f"Excellent response on {criterion} that is precise, well-supported, and insightfully appraised.", "Low: clear First. Mid: strong First. High: rare exceptional answer with sustained precision and insight."),
        ]
    return [
        ("Missing", f"No usable attempt on {criterion}.", "Low: absent. Mid: fragmentary. High: minimal but relevant setup."),
        ("Weak", f"Response on {criterion} contains major errors or lacks the core method.", "Low: incorrect or missing method. Mid: some relevant setup. High: close to adequate but still materially wrong."),
        ("Adequate", f"Basic handling of {criterion} with some correct setup or partial method credit.", "Low: limited correct work. Mid: fair partial-credit answer. High: almost secure but with important gaps."),
        ("Secure", f"Mostly correct work on {criterion} with reasonable method and interpretation.", "Low: secure but patchy. Mid: reliable and mostly complete. High: close to strong with only minor issues."),
        ("Strong", f"Clear and largely correct handling of {criterion}, including accurate working and sensible interpretation where relevant.", "Low: strong but slightly uneven. Mid: consistently strong. High: nearly excellent."),
        ("Excellent", f"Full and accurate handling of {criterion} with precise method, complete reasoning, and convincing interpretation.", "Low: fully correct. Mid: fully correct and precise. High: exceptional clarity and completeness."),
    ]


def _summarize_guidance_for_matrix(marking_guidance: str) -> str:
    text = " ".join(marking_guidance.replace("\n", " ").split()).strip()
    if not text:
        return "Assess the unit against the stated task."
    match = re.split(r"(?<=[.!?])\s+", text, maxsplit=1)
    summary = match[0].strip()
    return summary[:220].rstrip()


def _extract_classification_criteria(unit: AssessmentUnit) -> list[str]:
    if unit.grading_mode == "analytical":
        return _build_analytical_criteria(unit)
    return _build_deterministic_criteria(unit)


def _default_classification_criteria(unit: AssessmentUnit) -> list[str]:
    if unit.grading_mode == "analytical":
        return [
            "Judge the quality of the argument, not just the amount written.",
            "Rank responses higher when they are more precise, better evidenced, and more analytically convincing.",
            "Use lower placement when the answer is descriptive, underdeveloped, unsupported, or only partially relevant.",
        ]
    return [
        "Judge the accuracy of the method, algebra, and final result.",
        "Rank responses higher when they show more complete working, fewer errors, and stronger interpretation.",
        "Use lower placement when important steps are missing or the reasoning is materially incorrect.",
    ]


def _build_analytical_criteria(unit: AssessmentUnit) -> list[str]:
    task_type_criteria = _criteria_for_task_type(unit)
    if task_type_criteria:
        return task_type_criteria[:5]
    text = unit.marking_guidance.lower()
    focus = _build_task_focus(unit)
    criteria = [
        f"Evaluate whether the response directly addresses {focus}.",
        "Evaluate whether the claims are accurate, relevant, and clearly explained rather than generic or impressionistic.",
        "Evaluate whether the answer uses supporting evidence, references, or course concepts where the task requires them.",
        "Evaluate whether the response shows appraisal, qualification, or correction rather than simple description.",
    ]
    if any(term in text for term in ("quote", "page", "slide", "reference", "evidence", "report")):
        criteria[2] = "Evaluate whether the answer uses specific supporting evidence, references, page numbers, or cited material where required."
    if any(term in text for term in ("rewrite", "own tone", "own words", "edit the output", "correct any mistakes")):
        criteria.append("Evaluate whether the response corrects mistakes and rewrites the material clearly in the student's own justified form.")
    return criteria[:5]


def _build_deterministic_criteria(unit: AssessmentUnit) -> list[str]:
    task_type_criteria = _criteria_for_task_type(unit)
    if task_type_criteria:
        return task_type_criteria[:5]
    text = unit.marking_guidance.lower()
    focus = _build_task_focus(unit)
    criteria = [
        f"Evaluate whether the attempt sets up and answers {focus} correctly.",
        "Evaluate whether the working is complete enough to justify the result, not just whether a final answer is stated.",
        "Evaluate whether the algebra, notation, equation choice, or derivation steps are accurate.",
        "Evaluate whether any explanation or interpretation follows correctly from the method and result.",
    ]
    if any(term in text for term in ("measure", "variable", "data source", "regression", "equation")):
        criteria.append("Evaluate whether variables, measurement choices, or formal specifications are clearly and correctly defined.")
    return criteria[:5]


def _build_classification_question(unit: AssessmentUnit) -> str:
    if unit.grading_mode == "analytical":
        return f"Which quality band best fits the student's response to {unit.label}, and where does it sit inside that band?"
    return f"Which correctness band best fits the student's attempt on {unit.label}, and where does it sit inside that band?"


def _build_ranking_rule(unit: AssessmentUnit) -> str:
    if unit.grading_mode == "analytical":
        return "Group scripts by quality band first, then rank them within the band by precision, support, relevance, and analytical depth."
    return "Group attempts by correctness band first, then rank them within the band by method accuracy, completeness of working, and interpretation."


def _criteria_for_task_type(unit: AssessmentUnit) -> list[str]:
    focus = _build_task_focus(unit)
    task_type = unit.task_type
    construct_criteria = _construct_criteria_for_task(
        task_type,
        unit.question_text_exact,
        unit.marking_guidance,
    )
    if task_type == "critique_and_revision":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the response completes {focus}.",
                "Evaluate whether the answer identifies substantive mistakes, omissions, or distortions rather than offering generic criticism.",
                "Evaluate whether the critique is grounded in the supplied source, prompt, or task requirements.",
                "Evaluate whether the revised version is clearer, more accurate, and appropriately rewritten rather than lightly edited.",
            ],
        )[:5]
    if task_type == "model_specification":
        return [
            f"Evaluate whether the response states {focus} correctly.",
            "Evaluate whether symbols, terms, and relationships are defined accurately and consistently.",
            "Evaluate whether the explanation makes clear what each part of the specification means in context.",
            "Evaluate whether any interpretation follows from the stated model rather than from unsupported commentary.",
        ]
    if task_type == "comparative_statics":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the response explains {focus} correctly.",
                "Evaluate whether the answer identifies the direction of the effect of changing each relevant parameter.",
                "Evaluate whether the explanation links parameter changes to the comparison between the relevant scenarios or outcomes.",
                "Evaluate whether the reasoning follows from the model rather than from unsupported intuition alone.",
            ],
        )[:5]
    if task_type == "deterministic_derivation":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the attempt sets up and answers {focus} correctly.",
                "Evaluate whether the working is complete enough to justify the result, not just whether a final answer is stated.",
                "Evaluate whether derivation steps, algebra, notation, and conditions are accurate.",
                "Evaluate whether any interpretation follows correctly from the method and result.",
            ],
        )[:5]
    if task_type == "explanation_interpretation":
        return [
            f"Evaluate whether the response explains {focus} clearly and accurately.",
            "Evaluate whether the explanation identifies the relevant mechanism, direction, or relationship instead of describing outcomes vaguely.",
            "Evaluate whether the interpretation matches the underlying model, evidence, or setup.",
            "Evaluate whether the reasoning is complete enough to show why the conclusion follows.",
        ]
    if task_type == "welfare_reasoning":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the response addresses {focus} using the relevant conditions or assumptions.",
                "Evaluate whether the answer distinguishes effects on the relevant parties, components, or margins rather than collapsing them into a single claim.",
                "Evaluate whether the welfare reasoning is internally consistent and follows from the setup.",
                "Evaluate whether the conclusion is qualified appropriately when conditions or tradeoffs matter.",
            ],
        )[:5]
    if task_type == "evaluative_discussion":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the response directly addresses {focus}.",
                "Evaluate whether the answer makes a defensible judgement rather than offering only description or summary.",
                "Evaluate whether the reasoning is supported with relevant evidence, examples, or concepts where the task requires them.",
                "Evaluate whether the discussion acknowledges tradeoffs, limits, counterarguments, or qualifications where they matter.",
            ],
        )[:5]
    if task_type == "prediction_generation":
        return [
            f"Evaluate whether the response states {focus} that follow from the setup.",
            "Evaluate whether each prediction is directionally clear and not merely a repetition of the question.",
            "Evaluate whether the predictions are tied to the stated model, assumptions, or mechanism.",
            "Evaluate whether the answer distinguishes strong predictions from conditional or qualified ones where necessary.",
        ]
    if task_type == "measurement_and_data_design":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the response defines {focus} clearly.",
                "Evaluate whether the answer maps concepts to observable variables or indicators in a plausible way.",
                "Evaluate whether proposed data sources are suitable for the measures being discussed.",
                "Evaluate whether the response explains how the source or measure captures the intended concept and notes material limitations where relevant.",
            ],
        )[:5]
    if task_type == "regression_specification":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the response states {focus} correctly.",
                "Evaluate whether variables, terms, and functional relationships are defined clearly and consistently.",
                "Evaluate whether the coefficient, comparison, or hypothesis of interest is identified correctly.",
                "Evaluate whether the specification matches the stated empirical question rather than a different one.",
            ],
        )[:5]
    if task_type == "causal_identification":
        return [
            f"Evaluate whether the response identifies {focus} correctly.",
            "Evaluate whether the answer explains the mechanism behind the identification problem rather than naming it only.",
            "Evaluate whether proposed remedies are relevant, feasible, and connected to the stated problem.",
            "Evaluate whether the response distinguishes stronger and weaker solutions or remaining limitations where relevant.",
        ]
    if task_type == "synthesis_across_sources":
        return _merge_unique_criteria(
            construct_criteria,
            [
                f"Evaluate whether the response completes {focus}.",
                "Evaluate whether the answer identifies the main claims, assumptions, or omissions across the relevant sources or materials.",
                "Evaluate whether theory, evidence, and source content are integrated rather than discussed in isolation.",
                "Evaluate whether the response explains what the model or framework can and cannot account for.",
            ],
        )[:5]
    return []


def _build_task_focus(unit: AssessmentUnit) -> str:
    if unit.task_type == "critique_and_revision":
        return "the required draft, accuracy audit, and corrected revision"
    if unit.task_type == "model_specification":
        return "the required payoff function and explanation of its terms"
    if unit.task_type == "comparative_statics":
        return "the required comparative-statics explanation of how parameter changes affect the relevant outcomes"
    if unit.task_type == "deterministic_derivation":
        return "the required derivation and resulting conditions"
    if unit.task_type == "explanation_interpretation":
        return "the required explanation and interpretation"
    if unit.task_type == "welfare_reasoning":
        return "the relevant welfare conditions and effects on the affected parties"
    if unit.task_type == "evaluative_discussion":
        return "the required evaluative discussion"
    if unit.task_type == "prediction_generation":
        return "the requested predictions or implications"
    if unit.task_type == "measurement_and_data_design":
        return "the required measurement choices and data-source mapping"
    if unit.task_type == "regression_specification":
        return "the required regression equation, coefficient of interest, and hypothesis"
    if unit.task_type == "causal_identification":
        return "the endogeneity concerns and feasible solutions"
    if unit.task_type == "synthesis_across_sources":
        return "the required synthesis of framework, evidence, and material claims"
    text = unit.marking_guidance.lower()
    if unit.grading_mode == "analytical":
        if any(term in text for term in ("comment", "evaluate", "discuss", "good idea")):
            return "the required evaluative discussion"
        if any(term in text for term in ("rewrite", "own tone", "own words")):
            return "the required corrected and rewritten response"
        return "the required analytical discussion"
    if any(term in text for term in ("equation", "regression")):
        return "the required equation or formal specification"
    if any(term in text for term in ("derive", "condition")):
        return "the required derivation and resulting conditions"
    if any(term in text for term in ("measure", "variable", "data source")):
        return "the required measurement or variable definition"
    if any(term in text for term in ("prediction", "list")):
        return "the requested predictions or listed implications"
    if any(term in text for term in ("explain", "parameter", "intuition", "interpret")):
        return "the required explanation and interpretation"
    return "the required analytical step"


def normalize_part_analysis(data: dict[str, Any], part: SubmissionPart) -> dict[str, Any]:
    expected_label = part.label
    expected_criteria = _build_prompt_criteria(part)
    score = data.get("provisional_score")
    score_field = "provisional_score"
    if score is None:
        score = data.get("provisional_score_0_to_100")
        score_field = "provisional_score_0_to_100"
    if isinstance(score, bool) or not isinstance(score, (int, float)):
        raise ValueError(f"Model did not return a numeric provisional score for {expected_label}: {data}")
    normalized_score = float(score)
    if normalized_score < 0:
        raise ValueError(f"Model returned a negative provisional score for {expected_label}: {data}")
    if score_field == "provisional_score_0_to_100" and normalized_score > 100:
        raise ValueError(f"Model returned a provisional score outside 0-100 for {expected_label}: {data}")

    criterion_notes = _normalize_criterion_notes(
        data.get("criterion_notes"),
        expected_criteria=expected_criteria,
        expected_label=expected_label,
    )
    strengths = _normalize_string_list(
        data.get("strengths"),
        "strengths",
        expected_label,
        minimum=2,
        pad_missing=True,
    )
    weaknesses = _normalize_string_list(
        data.get("weaknesses"),
        "weaknesses",
        expected_label,
        minimum=2,
        pad_missing=True,
    )
    evidence = _normalize_string_list(
        data.get("evidence"),
        "evidence",
        expected_label,
        minimum=1,
        pad_missing=True,
    )
    coverage_comment = data.get("coverage_comment")
    if not isinstance(coverage_comment, str) or not coverage_comment.strip():
        raise ValueError(f"Model did not return a coverage_comment for {expected_label}: {data}")

    if score_field == "provisional_score" and part.max_mark is not None and part.max_mark > 0:
        score_pct = (normalized_score / float(part.max_mark)) * 100.0
    else:
        score_pct = normalized_score
    grade_band = _normalize_grade_band(
        value=data.get("grade_band"),
        score_pct=score_pct,
        marking_guidance=part.marking_guidance,
    )
    within_band_position = _normalize_within_band_position(
        value=data.get("within_band_position"),
        score_pct=score_pct,
        grade_band=grade_band,
        marking_guidance=part.marking_guidance,
    )
    band_confidence = data.get("band_confidence_0_to_100")
    if isinstance(band_confidence, bool) or not isinstance(band_confidence, (int, float)):
        band_confidence_value = _default_band_confidence(grade_band, marking_guidance=part.marking_guidance)
    else:
        band_confidence_value = round(float(band_confidence), 2)
        if band_confidence_value < 0 or band_confidence_value > 100:
            band_confidence_value = _default_band_confidence(grade_band, marking_guidance=part.marking_guidance)

    return {
        "section_label": expected_label,
        "provisional_score_0_to_100": round(normalized_score, 2) if score_field == "provisional_score_0_to_100" else None,
        "provisional_score": round(normalized_score, 2) if score_field == "provisional_score" else None,
        "grade_band": grade_band,
        "within_band_position": within_band_position,
        "band_confidence_0_to_100": band_confidence_value,
        "criterion_notes": criterion_notes,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "evidence": evidence,
        "coverage_comment": coverage_comment.strip(),
    }


def _run_part_analysis_with_retry(
    part: SubmissionPart,
    context: MarkingContext,
    filename: str,
    model_name: str,
    ollama_url: str,
) -> dict[str, Any]:
    messages = build_part_messages(part, context, filename, model_name=model_name)
    last_error: Exception | None = None
    for attempt in range(2):
        analysis = _call_ollama_json(
            model_name=model_name,
            messages=messages,
            ollama_url=ollama_url,
            profile="part_scoring",
        )
        try:
            return normalize_part_analysis(analysis, part=part)
        except ValueError as exc:
            last_error = exc
            if attempt == 1 or not _should_retry_part_analysis(exc):
                raise
            messages = _build_part_analysis_retry_messages(messages, part, exc)
    if last_error is not None:
        raise last_error
    raise ValueError(f"Part analysis failed without a captured error for {part.label}.")


def build_part_verification_messages(
    part: SubmissionPart,
    context: MarkingContext,
    filename: str,
    part_analysis: dict[str, Any],
    model_name: str | None = None,
) -> list[dict[str, str]]:
    messages = build_part_messages(part, context, filename, model_name=model_name)
    payload_text = messages[1]["content"]
    analysis_summary = json.dumps(
        {
            "section_label": part_analysis.get("section_label"),
            "provisional_score": part_analysis.get("provisional_score"),
            "grade_band": part_analysis.get("grade_band"),
            "within_band_position": part_analysis.get("within_band_position"),
            "criterion_notes": part_analysis.get("criterion_notes", []),
            "strengths": part_analysis.get("strengths", []),
            "weaknesses": part_analysis.get("weaknesses", []),
            "evidence": part_analysis.get("evidence", []),
            "coverage_comment": part_analysis.get("coverage_comment", ""),
        },
        ensure_ascii=True,
        indent=2,
    )
    max_mark_text = _format_number(part.max_mark) if part.max_mark is not None else "100"
    system = (
        "You are verifying one section-level grading payload and score. "
        "Your job is to detect over-broad or over-generous criteria and scores that are too generous relative to the section evidence. "
        "Do not diagnose hidden payload issues such as missing required constructs. "
        "Instead, suggest tighter criterion refinements and moderate over-generous scoring."
    )
    user = (
        f"STUDENT FILE NAME: {filename}\n\n"
        "SECTION SCORING PAYLOAD:\n"
        f"{payload_text}\n\n"
        "PROPOSED SECTION ANALYSIS:\n"
        f"{analysis_summary}\n\n"
        "Return only one JSON object with exactly these keys:\n"
        '- "agreement": string\n'
        '- "confidence_0_to_100": number\n'
        '- "criteria_adjustments": array of strings\n'
        '- "overgenerous_criteria": array of strings\n'
        '- "score_too_generous": boolean\n'
        '- "max_reasonable_score": number\n'
        '- "recommendation": string\n\n'
        "Rules:\n"
        "1. agreement must be one of: agree, minor_concern, major_concern.\n"
        "2. criteria_adjustments must suggest tighter criterion wording only when the current criteria are too broad, too generic, or too generous.\n"
        "3. overgenerous_criteria must identify any existing criteria that are rewarding too much for this section.\n"
        "4. score_too_generous should be true only when the current score is materially too high given the section evidence.\n"
        f"5. max_reasonable_score must be between 0 and {max_mark_text}.\n"
        "6. recommendation should be one of: accept, lower_score, refine_criteria.\n"
        "7. Do not include markdown fences or any text outside the JSON object."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def normalize_part_verification_result(data: dict[str, Any], max_mark: float | None) -> dict[str, Any]:
    agreement = str(data.get("agreement", "")).strip()
    if agreement not in {"agree", "minor_concern", "major_concern"}:
        raise ValueError(f"Part verifier returned an invalid agreement value: {data}")
    confidence = data.get("confidence_0_to_100")
    if isinstance(confidence, bool) or not isinstance(confidence, (int, float)):
        raise ValueError(f"Part verifier returned an invalid confidence score: {data}")
    confidence_value = float(confidence)
    if confidence_value < 0 or confidence_value > 100:
        raise ValueError(f"Part verifier confidence is out of range: {data}")

    def normalize_string_list_field(field_name: str) -> list[str]:
        value = data.get(field_name, [])
        if not isinstance(value, list):
            raise ValueError(f"Part verifier returned invalid {field_name}: {data}")
        return [str(item).strip() for item in value if str(item).strip()]

    score_too_generous = data.get("score_too_generous")
    if not isinstance(score_too_generous, bool):
        raise ValueError(f"Part verifier returned invalid score_too_generous flag: {data}")
    max_reasonable_score = data.get("max_reasonable_score")
    if isinstance(max_reasonable_score, bool) or not isinstance(max_reasonable_score, (int, float)):
        raise ValueError(f"Part verifier returned invalid max_reasonable_score: {data}")
    max_reasonable_value = float(max_reasonable_score)
    upper_bound = float(max_mark) if max_mark is not None else 100.0
    if max_reasonable_value < 0 or max_reasonable_value > upper_bound:
        raise ValueError(f"Part verifier max_reasonable_score is out of range: {data}")
    recommendation = str(data.get("recommendation", "")).strip()
    if recommendation not in {"accept", "lower_score", "refine_criteria"}:
        raise ValueError(f"Part verifier returned invalid recommendation: {data}")
    return {
        "agreement": agreement,
        "confidence_0_to_100": round(confidence_value, 2),
        "criteria_adjustments": normalize_string_list_field("criteria_adjustments"),
        "overgenerous_criteria": normalize_string_list_field("overgenerous_criteria"),
        "score_too_generous": score_too_generous,
        "max_reasonable_score": round(max_reasonable_value, 2),
        "recommendation": recommendation,
    }


def verify_part_analysis(
    part: SubmissionPart,
    context: MarkingContext,
    filename: str,
    part_analysis: dict[str, Any],
    model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> dict[str, Any]:
    data = _call_ollama_json(
        model_name=model_name,
        messages=build_part_verification_messages(part, context, filename, part_analysis, model_name=model_name),
        ollama_url=ollama_url,
        profile="verification",
    )
    return normalize_part_verification_result(data, max_mark=part.max_mark)


def apply_part_verifier_result(
    part_analysis: dict[str, Any],
    verifier_result: dict[str, Any],
) -> dict[str, Any]:
    adjusted = dict(part_analysis)
    current_score = float(adjusted.get("provisional_score") or 0.0)
    capped_score = min(current_score, float(verifier_result.get("max_reasonable_score", current_score)))
    if (
        verifier_result.get("score_too_generous")
        and verifier_result.get("recommendation") == "lower_score"
        and capped_score < current_score
    ):
        adjusted["provisional_score"] = round(capped_score, 2)
        issue_text = "; ".join(
            item
            for item in (
                list(verifier_result.get("overgenerous_criteria", []))
                + list(verifier_result.get("criteria_adjustments", []))
            )
            if item
        )
        if issue_text:
            adjusted["coverage_comment"] = (
                f"{adjusted.get('coverage_comment', '').strip()} "
                f"Verifier cap applied because the earlier score was too generous: {issue_text}."
            ).strip()
    adjusted["verifier_result"] = verifier_result
    return adjusted


def build_missing_part_analysis(part: SubmissionPart) -> dict[str, Any]:
    guidance = part.marking_guidance
    criteria = _build_prompt_criteria(part)
    zero_score = 0.0
    grade_band = _normalize_grade_band(
        value="missing/unfinished",
        score_pct=zero_score,
        marking_guidance=guidance,
    )
    within_band_position = _normalize_within_band_position(
        value="low",
        score_pct=zero_score,
        grade_band=grade_band,
        marking_guidance=guidance,
    )
    max_mark_text = _format_number(part.max_mark) if part.max_mark is not None else "100"
    return {
        "section_label": part.label,
        "provisional_score_0_to_100": None,
        "provisional_score": zero_score,
        "grade_band": grade_band,
        "within_band_position": within_band_position,
        "band_confidence_0_to_100": 95.0,
        "criterion_notes": [
            {
                "criterion_name": criterion["criterion_name"],
                "status": "weak" if criterion.get("scale_type") == "likert_judgement" else "missing_or_wrong",
                "normalized_status": "missed",
                "scale_type": criterion.get("scale_type", ""),
                "scale_labels": [
                    label.strip()
                    for label in str(criterion.get("scale_labels", "")).split(",")
                    if label.strip()
                ],
                "note": "no usable response",
            }
            for criterion in criteria
        ],
        "strengths": [
            "No creditable response was identified for this section.",
            "The section remains available for explicit zero-credit accounting in the final arithmetic.",
        ],
        "weaknesses": [
            "The expected section content was not found in the extracted submission text.",
            f"No valid method, evidence, or answer was available to earn marks out of {max_mark_text}.",
        ],
        "evidence": [
            "No section text was matched to this expected assessment unit.",
        ],
        "coverage_comment": (
            "This expected assessment unit was not detected in the submission text. "
            "It is therefore treated as missing for deterministic arithmetic and awarded zero."
        ),
    }


def moderate_part_analyses_across_submission(
    script_text: str,
    context: MarkingContext,
    filename: str,
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
    model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> list[dict[str, Any]]:
    data = _call_ollama_json(
        model_name=model_name,
        messages=build_moderation_messages(script_text, parts, part_analyses, context, filename),
        ollama_url=ollama_url,
        profile="moderation",
    )
    return normalize_moderated_part_scores(data, parts, part_analyses)


def moderate_linked_part_analyses(
    script_text: str,
    context: MarkingContext,
    filename: str,
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
    assessment_map: AssessmentMap,
    model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> list[dict[str, Any]]:
    unit_by_key = {_normalize_label_key(unit.label): unit for unit in assessment_map.units}
    groups: dict[str, list[int]] = {}
    for index, part in enumerate(parts):
        unit = unit_by_key.get(_normalize_label_key(part.label))
        if unit is None or not unit.dependency_group:
            continue
        groups.setdefault(unit.dependency_group, []).append(index)

    if not groups:
        return part_analyses

    updated = [dict(item) for item in part_analyses]
    for _, indexes in groups.items():
        if len(indexes) < 2:
            continue
        grouped_parts = [parts[index] for index in indexes]
        grouped_analyses = [updated[index] for index in indexes]
        should_moderate, _ = _moderation_group_decision(grouped_parts, grouped_analyses)
        if not should_moderate:
            continue
        moderated = moderate_part_analyses_across_submission(
            script_text=script_text,
            context=context,
            filename=filename,
            parts=grouped_parts,
            part_analyses=grouped_analyses,
            model_name=model_name,
            ollama_url=ollama_url,
        )
        for index, moderated_item in zip(indexes, moderated):
            updated[index] = moderated_item
    return updated


def _should_moderate_group(
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
) -> bool:
    return _moderation_group_decision(parts, part_analyses)[0]


def describe_moderation_plan(
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
    assessment_map: AssessmentMap,
) -> list[dict[str, Any]]:
    unit_by_key = {_normalize_label_key(unit.label): unit for unit in assessment_map.units}
    groups: dict[str, list[int]] = {}
    for index, part in enumerate(parts):
        unit = unit_by_key.get(_normalize_label_key(part.label))
        if unit is None or not unit.dependency_group:
            continue
        groups.setdefault(unit.dependency_group, []).append(index)
    plan: list[dict[str, Any]] = []
    for group_name, indexes in groups.items():
        grouped_parts = [parts[index] for index in indexes]
        grouped_analyses = [part_analyses[index] for index in indexes]
        should_moderate, reason = _moderation_group_decision(grouped_parts, grouped_analyses)
        plan.append(
            {
                "dependency_group": group_name,
                "part_labels": [part.label for part in grouped_parts],
                "should_moderate": should_moderate,
                "reason": reason,
            }
        )
    return plan


def _moderation_group_decision(
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
) -> tuple[bool, str]:
    if len(parts) < 2 or len(parts) != len(part_analyses):
        return False, "insufficient_group_size"
    normalized_scores: list[float] = []
    band_levels: list[int] = []
    for part, analysis in zip(parts, part_analyses):
        if not part.section_text.strip():
            return False, f"missing_section_text:{part.label}"
        score_pct = _analysis_score_pct(part, analysis)
        if score_pct is None:
            return False, f"missing_score:{part.label}"
        normalized_scores.append(score_pct)
        band_level = _band_level(str(analysis.get("grade_band", "")).strip())
        if band_level is not None:
            band_levels.append(band_level)
    spread = max(normalized_scores) - min(normalized_scores)
    if spread > 15.0:
        return False, f"wide_score_spread:{spread:.2f}"
    if band_levels and (max(band_levels) - min(band_levels)) > 1:
        return False, "wide_band_gap"
    return True, "eligible"


def _analysis_score_pct(part: SubmissionPart, analysis: dict[str, Any]) -> float | None:
    score = analysis.get("provisional_score")
    if score is not None:
        if part.max_mark is None or float(part.max_mark) <= 0:
            return None
        return round((float(score) / float(part.max_mark)) * 100.0, 2)
    score_pct = analysis.get("provisional_score_0_to_100")
    if score_pct is None:
        return None
    return round(float(score_pct), 2)


def _band_level(grade_band: str) -> int | None:
    if not grade_band:
        return None
    order = {
        "Missing/unfinished": 0,
        "Missing": 0,
        "Fail": 1,
        "Weak": 1,
        "Third/Pass": 2,
        "Developing": 2,
        "2:2": 3,
        "Adequate": 3,
        "2:1": 4,
        "Secure": 4,
        "First": 5,
        "Strong": 5,
        "Excellent": 6,
    }
    return order.get(grade_band)


def normalize_moderated_part_scores(
    data: dict[str, Any],
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    adjusted_sections = data.get("adjusted_sections")
    if not isinstance(adjusted_sections, list):
        raise ValueError(f"Model returned invalid adjusted_sections for moderation: {data}")

    current_by_key = {_normalize_label_key(analysis["section_label"]): dict(analysis) for analysis in part_analyses}
    part_by_key = {_normalize_label_key(part.label): part for part in parts}
    original_by_key = {_normalize_label_key(analysis["section_label"]): analysis for analysis in part_analyses}
    moderated: dict[str, dict[str, Any]] = {}

    for item in adjusted_sections:
        if not isinstance(item, dict):
            raise ValueError(f"Moderation item was not an object: {item}")
        label = str(item.get("section_label", "")).strip()
        key = _normalize_label_key(label)
        if key not in current_by_key or key not in part_by_key:
            raise ValueError(f"Moderation returned an unknown section label: {label}")
        current = dict(current_by_key[key])
        part = part_by_key[key]
        rationale = item.get("rationale")
        if not isinstance(rationale, str) or not rationale.strip():
            raise ValueError(f"Moderation did not return a rationale for {label}: {item}")

        adjusted_abs = item.get("adjusted_provisional_score")
        adjusted_pct = item.get("adjusted_provisional_score_0_to_100")
        if adjusted_abs is not None and adjusted_pct is not None:
            raise ValueError(f"Moderation returned both absolute and percentage scores for {label}: {item}")

        original_score = original_by_key[key]
        before = original_score.get("provisional_score")
        if before is None:
            before = original_score.get("provisional_score_0_to_100")

        if part.max_mark is not None:
            if adjusted_abs is None:
                adjusted_value = float(before)
            else:
                if isinstance(adjusted_abs, bool) or not isinstance(adjusted_abs, (int, float)):
                    raise ValueError(f"Moderation did not return an absolute score for {label}: {item}")
                adjusted_value = round(float(adjusted_abs), 2)
                if adjusted_value < 0 or adjusted_value > float(part.max_mark):
                    raise ValueError(f"Moderation returned out-of-range score for {label}: {item}")
            current["provisional_score"] = adjusted_value
            current["provisional_score_0_to_100"] = None
        else:
            if adjusted_pct is None:
                adjusted_value = float(before)
            else:
                if isinstance(adjusted_pct, bool) or not isinstance(adjusted_pct, (int, float)):
                    raise ValueError(f"Moderation did not return a percentage score for {label}: {item}")
                adjusted_value = round(float(adjusted_pct), 2)
                if adjusted_value < 0 or adjusted_value > 100:
                    raise ValueError(f"Moderation returned out-of-range percentage for {label}: {item}")
            current["provisional_score"] = None
            current["provisional_score_0_to_100"] = adjusted_value

        after = current.get("provisional_score")
        if after is None:
            after = current.get("provisional_score_0_to_100")
        current["moderation_rationale"] = rationale.strip()
        current["moderation_delta"] = round(float(after) - float(before), 2)
        moderated[key] = current

    for analysis in part_analyses:
        key = _normalize_label_key(analysis["section_label"])
        if key in moderated:
            continue
        preserved = dict(analysis)
        preserved["moderation_rationale"] = "No moderation change returned."
        preserved["moderation_delta"] = 0.0
        moderated[key] = preserved

    return [moderated[_normalize_label_key(part.label)] for part in parts]


def normalize_marking_result(
    data: dict[str, Any],
    expected_max_mark: float,
    diagnostics: SubmissionDiagnostics | None = None,
    parts: list[SubmissionPart] | None = None,
    part_analyses: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    feedback = _normalize_feedback_text(data.get("overall_feedback") or data.get("feedback"))
    if not isinstance(feedback, str) or not feedback.strip():
        raise ValueError(f"Model did not return feedback text: {data}")

    total_mark = data.get("total_mark")
    model_max_mark = data.get("max_mark")
    if isinstance(model_max_mark, bool):
        raise ValueError(f"Model returned an invalid max_mark: {data}")

    normalized_total = None
    normalized_max = expected_max_mark

    if total_mark is not None:
        if isinstance(total_mark, bool) or not isinstance(total_mark, (int, float)):
            raise ValueError(f"Model did not return a numeric total_mark: {data}")
        normalized_total = float(total_mark)

    if isinstance(model_max_mark, (int, float)):
        if normalized_total is None:
            raise ValueError(f"Model returned max_mark without a numeric total_mark: {data}")
        normalized_model_max = float(model_max_mark)
        if normalized_model_max <= 0:
            raise ValueError(f"Model returned a non-positive max_mark: {data}")
        if abs(normalized_model_max - expected_max_mark) < 1e-9:
            normalized_max = normalized_model_max
        elif abs(normalized_model_max - 100.0) < 1e-9 and 0 <= normalized_total <= 100:
            normalized_total = round((normalized_total * expected_max_mark) / 100.0, 2)
        else:
            raise ValueError(
                "Model returned a max_mark that conflicts with the uploaded marking documents: "
                f"expected {_format_number(expected_max_mark)}, got {_format_number(normalized_model_max)}."
            )
    elif normalized_total is not None and normalized_total > expected_max_mark and normalized_total <= 100 and expected_max_mark != 100:
        normalized_total = round((normalized_total * expected_max_mark) / 100.0, 2)

    if normalized_total is not None and (normalized_total < 0 or normalized_total > expected_max_mark):
        raise ValueError(
            "Model returned a total_mark outside the expected range: "
            f"{_format_number(normalized_total)} / {_format_number(expected_max_mark)}."
        )

    strengths_value = data.get("strengths")
    weaknesses_value = data.get("weaknesses")
    strengths = _normalize_string_list(strengths_value, "strengths", "final result", minimum=2) if strengths_value is not None else []
    weaknesses = _normalize_string_list(weaknesses_value, "weaknesses", "final result", minimum=2) if weaknesses_value is not None else []

    covered_parts_raw = data.get("covered_parts")
    covered_parts = _normalize_string_list(covered_parts_raw, "covered_parts", "final result", minimum=1) if covered_parts_raw is not None else []
    if parts:
        covered_parts = [part.label for part in parts]

    if len(feedback.split()) < 120:
        raise ValueError("Model returned feedback that is shorter than the required minimum length.")

    if normalized_total is not None and normalized_total == expected_max_mark and _has_major_weaknesses(weaknesses, feedback):
        raise ValueError("Model awarded full marks while also describing substantial weaknesses.")

    if diagnostics and diagnostics.low_text:
        strengths.append(f"Extraction warning: submission contains only {diagnostics.extracted_word_count} words.")

    validation_notes = []
    if diagnostics:
        if diagnostics.possible_extraction_issue:
            validation_notes.append("possible_extraction_issue")
        if diagnostics.low_text:
            validation_notes.append("low_text")

    ai_total_mark = _clean_number(normalized_total) if normalized_total is not None else None
    math_total_mark = None
    ai_math_mark_delta = None

    if part_analyses:
        provisional_scores = [
            float(item["provisional_score"] if item.get("provisional_score") is not None else item["provisional_score_0_to_100"])
            for item in part_analyses
        ]
        moderation_deltas = [abs(float(item.get("moderation_delta", 0.0))) for item in part_analyses if item.get("moderation_delta") is not None]
        spread = max(provisional_scores) - min(provisional_scores)
        validation_notes.append(f"part_score_spread={spread:.2f}")
        if moderation_deltas:
            validation_notes.append(f"part_moderation_max_delta={max(moderation_deltas):.2f}")
        math_total = compute_total_mark_from_part_scores(expected_max_mark, parts or [], part_analyses)
        if math_total is None:
            raise ValueError("Could not compute deterministic total from part scores.")
        else:
            math_total_mark = _clean_number(math_total)
            normalized_total = math_total
            validation_notes.append("math_total_used")
            if ai_total_mark is not None:
                ai_math_mark_delta = _clean_number(round(ai_total_mark - math_total_mark, 2))
            if ai_total_mark is not None and abs(ai_total_mark - math_total_mark) > 1e-9:
                validation_notes.append(f"ai_total_disagreement={abs(ai_total_mark - math_total_mark):.2f}")
    elif normalized_total is None:
        raise ValueError("No deterministic part scores were available and the model did not return a total_mark.")

    return {
        "total_mark": _clean_number(normalized_total),
        "max_mark": _clean_number(normalized_max),
        "ai_total_mark": ai_total_mark,
        "math_total_mark": math_total_mark,
        "ai_math_mark_delta": ai_math_mark_delta,
        "overall_feedback": feedback.strip(),
        "strengths": strengths,
        "weaknesses": weaknesses,
        "covered_parts": covered_parts,
        "detected_part_count": diagnostics.detected_part_count if diagnostics else 1,
        "detected_part_labels": list(diagnostics.detected_part_labels) if diagnostics else [],
        "extracted_word_count": diagnostics.extracted_word_count if diagnostics else len(feedback.split()),
        "possible_extraction_issue": diagnostics.possible_extraction_issue if diagnostics else False,
        "validation_notes": validation_notes,
    }


def build_local_final_result(
    context: MarkingContext,
    diagnostics: SubmissionDiagnostics,
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
) -> dict[str, Any]:
    feedback = _render_feedback_from_part_analyses(parts, part_analyses, context.max_mark)
    strengths = _aggregate_part_bullets(part_analyses, field="strengths")
    weaknesses = _aggregate_part_bullets(part_analyses, field="weaknesses")
    data = {
        "overall_feedback": feedback,
        "covered_parts": [part.label for part in parts],
        "strengths": strengths[: max(2, len(strengths))],
        "weaknesses": weaknesses[: max(2, len(weaknesses))],
    }
    return normalize_marking_result(
        data,
        expected_max_mark=context.max_mark,
        diagnostics=diagnostics,
        parts=parts,
        part_analyses=part_analyses,
    )


def _attach_latency_metrics(
    result: dict[str, Any],
    timings: dict[str, float],
    total_started: float,
    parts: list[SubmissionPart],
) -> dict[str, Any]:
    total_seconds = round(time.perf_counter() - total_started, 2)
    part_count = max(len(parts), 1)
    result["latency_seconds_total"] = total_seconds
    result["latency_seconds_assessment_prepare"] = round(timings.get("assessment_prepare", 0.0), 2)
    result["latency_seconds_rubric_verify"] = result["latency_seconds_assessment_prepare"]
    result["latency_seconds_structure_detect"] = round(timings.get("structure_detect", 0.0), 2)
    result["latency_seconds_part_refine"] = round(timings.get("part_refine", 0.0), 2)
    result["latency_seconds_part_analysis"] = round(timings.get("part_analysis", 0.0), 2)
    result["latency_seconds_moderation"] = round(timings.get("moderation", 0.0), 2)
    result["latency_seconds_finalize"] = round(timings.get("finalize", 0.0), 2)
    result["latency_seconds_part_analysis_per_part_avg"] = round(result["latency_seconds_part_analysis"] / part_count, 2)
    validation_notes = list(result.get("validation_notes", []))
    validation_notes.append(f"latency_script_seconds={total_seconds:.2f}")
    validation_notes.append(f"latency_assessment_prepare_seconds={result['latency_seconds_assessment_prepare']:.2f}")
    validation_notes.append(f"latency_part_analysis_seconds={result['latency_seconds_part_analysis']:.2f}")
    result["validation_notes"] = validation_notes
    return result


def compute_total_mark_from_part_scores(
    expected_max_mark: float,
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
) -> float | None:
    if not parts or not part_analyses or len(parts) != len(part_analyses):
        return None

    part_marks = []
    for part, analysis in zip(parts, part_analyses):
        effective_part_max = float(part.max_mark) if part.max_mark is not None else None
        if len(parts) == 1 and effective_part_max is None:
            effective_part_max = float(expected_max_mark)

        score = analysis.get("provisional_score")
        if score is not None:
            part_marks.append(float(score))
            continue

        score_100 = analysis.get("provisional_score_0_to_100")
        if score_100 is None:
            return None
        if effective_part_max is None:
            return None
        part_marks.append((float(score_100) * effective_part_max) / 100.0)

    if not part_marks:
        return None

    summed_total = sum(part_marks)
    summed_total = round(summed_total, 2)
    return min(expected_max_mark, max(0.0, summed_total))


def _aggregate_part_bullets(part_analyses: list[dict[str, Any]], field: str) -> list[str]:
    items: list[str] = []
    seen: set[str] = set()
    for analysis in part_analyses:
        for raw in analysis.get(field, []):
            text = str(raw).strip()
            key = text.lower()
            if not text or key in seen:
                continue
            seen.add(key)
            items.append(text)
            if len(items) >= 6:
                return items
    return items


def _render_feedback_from_part_analyses(
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
    expected_max_mark: float,
) -> str:
    total = compute_total_mark_from_part_scores(expected_max_mark, parts, part_analyses)
    total_text = _format_number(total if total is not None else expected_max_mark)
    intro = (
        f"This script is classified from part-level evidence rather than a separate final-language pass. "
        f"The deterministic overall mark is {total_text} out of {_format_number(expected_max_mark)}, computed as the sum of the part scores."
    )
    section_lines: list[str] = []
    for part, analysis in zip(parts, part_analyses):
        section_score = analysis.get("provisional_score")
        if section_score is None and analysis.get("provisional_score_0_to_100") is not None and part.max_mark is not None:
            section_score = round((float(analysis["provisional_score_0_to_100"]) * float(part.max_mark)) / 100.0, 2)
        score_text = _format_number(float(section_score)) if section_score is not None else "unscored"
        max_text = _format_number(float(part.max_mark)) if part.max_mark is not None else "100"
        band = str(analysis.get("grade_band", "classified")).strip()
        position = str(analysis.get("within_band_position", "mid")).strip()
        coverage = str(analysis.get("coverage_comment", "")).strip()
        top_strength = ", ".join(str(item).strip() for item in analysis.get("strengths", [])[:2] if str(item).strip())
        top_weakness = ", ".join(str(item).strip() for item in analysis.get("weaknesses", [])[:2] if str(item).strip())
        section_lines.append(
            f"{part.label} scored {score_text}/{max_text} and was placed in the {band} band at the {position} end of that bracket. "
            f"{coverage} Stronger features here include {top_strength or 'clear relevant material'}, while the mark is held back by {top_weakness or 'remaining weaknesses in precision and support'}."
        )
    outro = (
        "The feedback therefore follows the classified rubric language for each part, rather than trying to invent a fresh overall judgement. "
        "Where a section sits inside its bracket depends on the evidence quality, completeness, precision, and support actually visible in that section."
    )
    return " ".join([intro, *section_lines, outro]).strip()


def calibrate_marks_across_students(
    provisional_results: list[dict[str, Any]],
    model_name: str,
    model_label: str,
    context: MarkingContext,
    assessment_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> dict[str, dict[str, Any]]:
    eligible = []
    for item in provisional_results:
        mark = item.get("total_mark")
        if isinstance(mark, (int, float)):
            eligible.append(
                {
                    "filename": str(item.get("filename", "")),
                    "provisional_mark": float(mark),
                    "detected_part_count": item.get("detected_part_count"),
                    "strengths": list(item.get("strengths", []))[:2],
                    "weaknesses": list(item.get("weaknesses", []))[:2],
                    "validation_notes": list(item.get("validation_notes", [])),
                }
            )

    if len(eligible) < 2:
        return {}

    data = _call_ollama_json(
        model_name=model_name,
        messages=build_calibration_messages(assessment_name, context, model_label, eligible),
        ollama_url=ollama_url,
        profile="moderation",
    )
    return normalize_calibration_result(data, eligible, context.max_mark)


def normalize_calibration_result(
    data: dict[str, Any],
    provisional_results: list[dict[str, Any]],
    max_mark: float,
) -> dict[str, dict[str, Any]]:
    calibrated = data.get("calibrated_results")
    if not isinstance(calibrated, list):
        raise ValueError(f"Calibration response did not include calibrated_results: {data}")

    expected = {str(item["filename"]): float(item["provisional_mark"]) for item in provisional_results}
    resolved: dict[str, dict[str, Any]] = {}
    for item in calibrated:
        if not isinstance(item, dict):
            continue
        filename = str(item.get("filename", "")).strip()
        adjusted_mark = item.get("adjusted_mark")
        rationale = str(item.get("rationale", "")).strip()
        if not filename or filename not in expected:
            continue
        if isinstance(adjusted_mark, bool) or not isinstance(adjusted_mark, (int, float)):
            raise ValueError(f"Calibration returned a non-numeric adjusted mark for {filename}: {item}")
        adjusted = float(adjusted_mark)
        if adjusted < 0 or adjusted > max_mark:
            raise ValueError(f"Calibration returned an out-of-range mark for {filename}: {item}")
        resolved[filename] = {
            "adjusted_mark": _clean_number(adjusted),
            "delta": round(adjusted - expected[filename], 2),
            "rationale": rationale,
        }

    missing = [filename for filename in expected if filename not in resolved]
    if missing:
        raise ValueError("Calibration did not return marks for: " + ", ".join(missing))
    return resolved


def verify_marking_result(
    script_text: str,
    context: MarkingContext,
    filename: str,
    grading_result: dict[str, Any],
    model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> dict[str, Any]:
    data = _call_ollama_json(
        model_name=model_name,
        messages=build_verification_messages(script_text, context, filename, grading_result),
        ollama_url=ollama_url,
        profile="verification",
    )
    return normalize_verification_result(data)


def regrade_marking_result(
    script_text: str,
    context: MarkingContext,
    filename: str,
    prior_result: dict[str, Any],
    verifier_result: dict[str, Any],
    model_name: str,
    ollama_url: str = DEFAULT_OLLAMA_URL,
) -> dict[str, Any]:
    parts = [
        SubmissionPart(label=label, focus_hint="")
        for label in prior_result.get("covered_parts", []) or ["Whole Submission"]
    ]
    diagnostics = build_submission_diagnostics(script_text, parts)
    messages = build_regrade_messages(script_text, context, filename, prior_result, verifier_result)
    return _run_final_result_with_retry(
        model_name=model_name,
        messages=messages,
        expected_max_mark=context.max_mark,
        diagnostics=diagnostics,
        parts=parts,
        part_analyses=[],
        ollama_url=ollama_url,
    )


def normalize_verification_result(data: dict[str, Any]) -> dict[str, Any]:
    agreement = str(data.get("agreement", "")).strip()
    if agreement not in {"agree", "minor_concern", "major_concern"}:
        raise ValueError(f"Verifier returned an invalid agreement value: {data}")
    confidence = data.get("confidence_0_to_100")
    if isinstance(confidence, bool) or not isinstance(confidence, (int, float)):
        raise ValueError(f"Verifier returned an invalid confidence score: {data}")
    confidence_value = float(confidence)
    if confidence_value < 0 or confidence_value > 100:
        raise ValueError(f"Verifier confidence is out of range: {data}")
    issues = data.get("issues")
    if not isinstance(issues, list):
        raise ValueError(f"Verifier returned invalid issues: {data}")
    normalized_issues = [str(item).strip() for item in issues if str(item).strip()]
    recommendation = str(data.get("recommendation", "")).strip()
    if not recommendation:
        raise ValueError(f"Verifier returned an empty recommendation: {data}")
    return {
        "agreement": agreement,
        "confidence_0_to_100": round(confidence_value, 2),
        "issues": normalized_issues,
        "recommendation": recommendation,
    }


def list_submission_files(folder: Path) -> list[Path]:
    return sorted(path for path in folder.iterdir() if path.is_file() and path.suffix.lower() in SUPPORTED_TEXT_SUFFIXES)


def parse_keyword_patterns(value: str) -> tuple[str, ...]:
    return tuple(part.strip().lower() for part in value.split(",") if part.strip())


def discover_assessment_bundles(
    root_folder: Path,
    rubric_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["rubric"],
    brief_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["brief"],
    marking_scheme_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["marking_scheme"],
    graded_sample_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["graded_sample"],
    other_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["other"],
) -> list[AssessmentBundle]:
    if not root_folder.exists():
        raise ValueError(f"Assessment root folder not found: {root_folder}")
    if not root_folder.is_dir():
        raise ValueError(f"Assessment root path is not a folder: {root_folder}")

    bundles: list[AssessmentBundle] = []
    for folder in sorted(path for path in root_folder.iterdir() if path.is_dir()):
        bundles.append(
            build_assessment_bundle(
                folder=folder,
                rubric_keywords=rubric_keywords,
                brief_keywords=brief_keywords,
                marking_scheme_keywords=marking_scheme_keywords,
                graded_sample_keywords=graded_sample_keywords,
                other_keywords=other_keywords,
            )
        )
    return bundles


def build_single_assessment_bundle(
    folder: Path,
    assessment_name: str | None = None,
    rubric_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["rubric"],
    brief_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["brief"],
    marking_scheme_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["marking_scheme"],
    graded_sample_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["graded_sample"],
    other_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["other"],
) -> AssessmentBundle:
    if not folder.exists():
        raise ValueError(f"Assessment folder not found: {folder}")
    if not folder.is_dir():
        raise ValueError(f"Assessment path is not a folder: {folder}")

    top_level_files = sorted(path for path in folder.iterdir() if path.is_file() and path.suffix.lower() in SUPPORTED_TEXT_SUFFIXES)
    rubric_files = _pick_matching_files(top_level_files, rubric_keywords)
    brief_files = _pick_matching_files(top_level_files, brief_keywords, exclude=rubric_files)
    marking_scheme_files = _pick_matching_files(top_level_files, marking_scheme_keywords, exclude=rubric_files + brief_files)
    graded_sample_files = _pick_matching_files(
        top_level_files,
        graded_sample_keywords,
        exclude=rubric_files + brief_files + marking_scheme_files,
    )
    other_files = _pick_matching_files(
        top_level_files,
        other_keywords,
        exclude=rubric_files + brief_files + marking_scheme_files + graded_sample_files,
    )

    child_submission_files = sorted(
        path
        for child in folder.iterdir()
        if child.is_dir()
        for path in child.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_TEXT_SUFFIXES
    )
    excluded_top_level = set(rubric_files + brief_files + marking_scheme_files + graded_sample_files + other_files)
    top_level_submissions = tuple(path for path in top_level_files if path not in excluded_top_level)
    submission_files = tuple(child_submission_files) + tuple(top_level_submissions)

    return AssessmentBundle(
        name=assessment_name or folder.name,
        folder=folder,
        submission_files=submission_files,
        rubric_files=rubric_files,
        brief_files=brief_files,
        marking_scheme_files=marking_scheme_files,
        graded_sample_files=graded_sample_files,
        other_files=other_files,
    )


def build_assessment_bundle(
    folder: Path,
    rubric_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["rubric"],
    brief_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["brief"],
    marking_scheme_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["marking_scheme"],
    graded_sample_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["graded_sample"],
    other_keywords: tuple[str, ...] = DEFAULT_CONTEXT_PATTERNS["other"],
) -> AssessmentBundle:
    files = _list_supported_files_recursive(folder)
    rubric_files = _pick_matching_files(files, rubric_keywords)
    brief_files = _pick_matching_files(files, brief_keywords, exclude=rubric_files)
    marking_scheme_files = _pick_matching_files(files, marking_scheme_keywords, exclude=rubric_files + brief_files)
    graded_sample_files = _pick_matching_files(
        files,
        graded_sample_keywords,
        exclude=rubric_files + brief_files + marking_scheme_files,
    )
    other_files = _pick_matching_files(
        files,
        other_keywords,
        exclude=rubric_files + brief_files + marking_scheme_files + graded_sample_files,
    )

    excluded = set(rubric_files + brief_files + marking_scheme_files + graded_sample_files + other_files)
    submission_files = tuple(path for path in files if path not in excluded)

    return AssessmentBundle(
        name=folder.name,
        folder=folder,
        submission_files=submission_files,
        rubric_files=rubric_files,
        brief_files=brief_files,
        marking_scheme_files=marking_scheme_files,
        graded_sample_files=graded_sample_files,
        other_files=other_files,
    )


def build_marking_context_from_bundle(bundle: AssessmentBundle) -> MarkingContext:
    return prepare_marking_context(
        rubric_text=read_paths_text(bundle.rubric_files),
        brief_text=read_paths_text(bundle.brief_files),
        marking_scheme_text=read_paths_text(bundle.marking_scheme_files),
        graded_sample_text=read_paths_text(bundle.graded_sample_files),
        other_context_text=read_paths_text(bundle.other_files),
    )


def apply_assessment_map_to_submission_parts(parts: list[SubmissionPart], assessment_map: AssessmentMap) -> list[SubmissionPart]:
    unit_by_key = {_normalize_label_key(unit.label): unit for unit in assessment_map.units}
    enriched: list[SubmissionPart] = []
    for part in parts:
        unit = unit_by_key.get(_normalize_label_key(part.label))
        if unit is None:
            enriched.append(part)
            continue
        enriched.append(
            SubmissionPart(
                label=part.label,
                focus_hint=part.focus_hint,
                anchor_text=part.anchor_text,
                section_text=part.section_text,
                max_mark=unit.max_mark if unit.max_mark is not None else part.max_mark,
                marking_guidance=unit.rubric_text or unit.marking_guidance or part.marking_guidance,
                question_text_exact=unit.question_text_exact or part.question_text_exact,
                task_type=unit.task_type or part.task_type,
                criterion_mode=unit.criterion_mode or part.criterion_mode,
            )
        )
    return enriched


def apply_prepared_artifact_to_submission_parts(
    parts: list[SubmissionPart],
    prepared_assessment_map: PreparedAssessmentMap,
) -> list[SubmissionPart]:
    prepared_unit_by_key = {
        _normalize_label_key(unit.label): unit
        for unit in prepared_assessment_map.prepared_units
    }
    enriched: list[SubmissionPart] = []
    for part in parts:
        prepared_unit = prepared_unit_by_key.get(_normalize_label_key(part.label))
        if prepared_unit is None or not prepared_unit.compact_criteria:
            enriched.append(part)
            continue
        compact_guidance = "\n".join(
            ["Compact scoring criteria:", *[f"- {criterion}" for criterion in prepared_unit.compact_criteria]]
        )
        merged_guidance = compact_guidance
        if part.marking_guidance.strip():
            merged_guidance = f"{compact_guidance}\n\nDetailed rubric:\n{part.marking_guidance.strip()}"
        enriched.append(
            SubmissionPart(
                label=part.label,
                focus_hint=part.focus_hint,
                anchor_text=part.anchor_text,
                section_text=part.section_text,
                max_mark=part.max_mark,
                marking_guidance=merged_guidance,
                question_text_exact=part.question_text_exact,
                task_type=part.task_type,
                criterion_mode=part.criterion_mode,
            )
        )
    return enriched


def _parse_first_json_object(content: str) -> dict[str, Any]:
    start = content.find("{")
    while start != -1:
        depth = 0
        in_string = False
        escape = False
        for index in range(start, len(content)):
            char = content[index]
            if in_string:
                if escape:
                    escape = False
                elif char == "\\":
                    escape = True
                elif char == '"':
                    in_string = False
                continue
            if char == '"':
                in_string = True
            elif char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    candidate = content[start : index + 1]
                    return json.loads(candidate)
        start = content.find("{", start + 1)
    raise ValueError(f"No JSON object found in model response: {content}")


def _clean_number(value: float) -> int | float:
    if float(value).is_integer():
        return int(value)
    return round(float(value), 2)


def _format_number(value: float) -> str:
    cleaned = _clean_number(value)
    return str(cleaned)


def _trim_support_text(text: str, max_chars: int = 1200) -> str:
    normalized = " ".join(text.split()).strip()
    if len(normalized) <= max_chars:
        return normalized
    return normalized[: max_chars - 3].rstrip() + "..."


def _build_context_text(context: MarkingContext) -> str:
    context_parts = []
    if context.rubric_text:
        context_parts.append(f"RUBRIC:\n{context.rubric_text}")
    if context.brief_text:
        context_parts.append(f"ASSIGNMENT BRIEF:\n{context.brief_text}")
    if context.marking_scheme_text:
        context_parts.append(f"MARKING SCHEME:\n{context.marking_scheme_text}")
    if context.graded_sample_text:
        context_parts.append(
            "EXAMPLE GRADED SCRIPT (for tone and structure only):\n"
            f"{context.graded_sample_text}"
        )
    if context.other_context_text:
        context_parts.append(f"OTHER SUPPORTING DOCUMENTS:\n{context.other_context_text}")
    return "\n\n".join(context_parts).strip()


def _build_part_support_text(context: MarkingContext, part: SubmissionPart) -> str:
    support_parts: list[str] = []
    structure_guidance = extract_structure_guidance(context)
    if structure_guidance:
        support_parts.append(f"STRUCTURE AND MARKS HINTS:\n{structure_guidance}")
    if context.brief_text.strip():
        support_parts.append(f"LOCAL ASSIGNMENT BRIEF SUPPORT:\n{_trim_support_text(context.brief_text)}")
    if context.marking_scheme_text.strip():
        support_parts.append(f"LOCAL MARKING SCHEME SUPPORT:\n{_trim_support_text(context.marking_scheme_text)}")
    if context.other_context_text.strip():
        support_parts.append(f"LOCAL OTHER SUPPORT:\n{_trim_support_text(context.other_context_text)}")
    if not support_parts and part.marking_guidance.strip():
        support_parts.append("Use only the section rubric and the section text.")
    return ("\n\n".join(support_parts) + "\n\n") if support_parts else ""


def _call_ollama_json(
    model_name: str,
    messages: list[dict[str, str]],
    ollama_url: str,
    profile: str = "default",
) -> dict[str, Any]:
    timeout_seconds = _ollama_timeout_seconds()
    payload = {
        "model": model_name,
        "format": "json",
        "messages": messages,
        "stream": False,
        "options": _ollama_options_for_profile(profile),
    }
    response = requests.post(ollama_url, json=payload, timeout=timeout_seconds)
    response.raise_for_status()
    response_json = response.json()
    content = extract_message_content(response_json)
    return parse_json_object(content)


def _ollama_timeout_seconds() -> int:
    raw = os.getenv("OLLAMA_HTTP_TIMEOUT_SECONDS", "").strip()
    if not raw:
        return DEFAULT_OLLAMA_TIMEOUT_SECONDS
    try:
        parsed = int(raw)
    except ValueError:
        return DEFAULT_OLLAMA_TIMEOUT_SECONDS
    return max(10, min(900, parsed))


def _ollama_options_for_profile(profile: str) -> dict[str, Any]:
    base = {
        "temperature": 0.0,
        "top_p": 0.85,
        "num_ctx": 8192,
        "seed": 7,
        "num_predict": 900,
    }
    if profile in {"structure_extraction", "rubric_verification", "verification"}:
        return {
            **base,
            "top_p": 0.8,
            "num_predict": 700,
        }
    if profile == "detector_small":
        return {
            **base,
            "num_ctx": 4096,
            "top_p": 0.75,
            "num_predict": 320,
        }
    if profile in {"part_scoring", "moderation", "final_result"}:
        return {
            **base,
            "temperature": 0.05,
            "top_p": 0.85,
            "num_predict": 1100,
        }
    return base


def _run_final_result_with_retry(
    model_name: str,
    messages: list[dict[str, str]],
    expected_max_mark: float,
    diagnostics: SubmissionDiagnostics,
    parts: list[SubmissionPart],
    part_analyses: list[dict[str, Any]],
    ollama_url: str,
) -> dict[str, Any]:
    last_error: Exception | None = None
    current_messages = list(messages)
    for attempt in range(2):
        data = _call_ollama_json(
            model_name=model_name,
            messages=current_messages,
            ollama_url=ollama_url,
            profile="final_result",
        )
        try:
            return normalize_marking_result(
                data,
                expected_max_mark=expected_max_mark,
                diagnostics=diagnostics,
                parts=parts,
                part_analyses=part_analyses,
            )
        except ValueError as exc:
            last_error = exc
            if attempt == 1 or not _should_retry_final_result(exc):
                raise
            current_messages = _build_final_result_retry_messages(current_messages, exc)
    if last_error is not None:
        raise last_error
    raise ValueError("Final result generation failed without a captured error.")


def _should_retry_final_result(exc: ValueError) -> bool:
    message = str(exc)
    retryable_patterns = (
        "shorter than the required minimum length",
        "did not explicitly cover all detected sections",
    )
    return any(pattern in message for pattern in retryable_patterns)


def _should_retry_part_analysis(exc: ValueError) -> bool:
    message = str(exc)
    retryable_patterns = (
        "criterion_notes",
        "too few items for strengths",
        "too few items for weaknesses",
        "too few items for evidence",
        "did not return a coverage_comment",
    )
    return any(pattern in message for pattern in retryable_patterns)


def _normalize_grade_band(value: Any, score_pct: float, marking_guidance: str) -> str:
    text = str(value).strip().lower()
    analytical = any(band in marking_guidance.lower() for band in ("first", "2:1", "2:2", "third/pass", "missing/unfinished", "fail"))
    if analytical:
        mapping = {
            "first": "First",
            "1st": "First",
            "2:1": "2:1",
            "upper second": "2:1",
            "2:2": "2:2",
            "lower second": "2:2",
            "third": "Third/Pass",
            "pass": "Third/Pass",
            "third/pass": "Third/Pass",
            "fail": "Fail",
            "missing": "Missing/unfinished",
            "missing/unfinished": "Missing/unfinished",
            "unfinished": "Missing/unfinished",
        }
        if text in mapping:
            return mapping[text]
        if score_pct >= 70:
            return "First"
        if score_pct >= 60:
            return "2:1"
        if score_pct >= 50:
            return "2:2"
        if score_pct >= 40:
            return "Third/Pass"
        if score_pct <= 20:
            return "Missing/unfinished"
        return "Fail"

    mapping = {
        "excellent": "Excellent",
        "strong": "Strong",
        "secure": "Secure",
        "adequate": "Adequate",
        "developing": "Developing",
        "weak": "Weak",
        "missing": "Missing",
    }
    if text in mapping:
        return mapping[text]
    if score_pct >= 85:
        return "Excellent"
    if score_pct >= 70:
        return "Strong"
    if score_pct >= 55:
        return "Secure"
    if score_pct >= 40:
        return "Adequate"
    if score_pct >= 25:
        return "Developing"
    if score_pct > 0:
        return "Weak"
    return "Missing"


def _normalize_within_band_position(value: Any, score_pct: float, grade_band: str, marking_guidance: str) -> str:
    text = str(value).strip().lower()
    if text in {"low", "mid", "high"}:
        return text
    analytical = any(band in marking_guidance.lower() for band in ("first", "2:1", "2:2", "third/pass", "missing/unfinished", "fail"))
    if analytical:
        if grade_band == "First":
            return _position_from_bounds(score_pct, 70.0, 100.0)
        if grade_band == "2:1":
            return _position_from_bounds(score_pct, 60.0, 69.99)
        if grade_band == "2:2":
            return _position_from_bounds(score_pct, 50.0, 59.99)
        if grade_band == "Third/Pass":
            return _position_from_bounds(score_pct, 40.0, 49.99)
        if grade_band == "Fail":
            return _position_from_bounds(score_pct, 20.01, 39.99)
        return _position_from_bounds(score_pct, 0.0, 20.0)
    if grade_band == "Excellent":
        return _position_from_bounds(score_pct, 85.0, 100.0)
    if grade_band == "Strong":
        return _position_from_bounds(score_pct, 70.0, 84.99)
    if grade_band == "Secure":
        return _position_from_bounds(score_pct, 55.0, 69.99)
    if grade_band == "Adequate":
        return _position_from_bounds(score_pct, 40.0, 54.99)
    if grade_band == "Developing":
        return _position_from_bounds(score_pct, 25.0, 39.99)
    if grade_band == "Weak":
        return _position_from_bounds(score_pct, 0.01, 24.99)
    return "low"


def _position_from_bounds(score_pct: float, lower: float, upper: float) -> str:
    if upper <= lower:
        return "mid"
    width = upper - lower
    if score_pct < lower + (width / 3):
        return "low"
    if score_pct > upper - (width / 3):
        return "high"
    return "mid"


def _default_band_confidence(grade_band: str, marking_guidance: str) -> float:
    analytical = any(band in marking_guidance.lower() for band in ("first", "2:1", "2:2", "third/pass", "missing/unfinished", "fail"))
    if analytical and grade_band in {"Fail", "Missing/unfinished"}:
        return 78.0
    if analytical:
        return 72.0
    return 75.0


def _build_rubric_verification_retry_messages(
    messages: list[dict[str, str]],
    unit: AssessmentUnit,
    exc: Exception,
) -> list[dict[str, str]]:
    retry_note = (
        "Previous attempt failed. Return exactly one valid JSON object with keys "
        '"rubric_text", "issues", and "confidence_0_to_100". '
        f"Do not truncate the JSON for {unit.label}. "
        f"Failure reason: {exc}"
    )
    return [*messages, {"role": "user", "content": retry_note}]


def _build_final_result_retry_messages(messages: list[dict[str, str]], exc: ValueError) -> list[dict[str, str]]:
    message = str(exc)
    if "shorter than the required minimum length" in message:
        retry_note = (
            "Previous attempt failed because overall_feedback was too short. "
            "Retry now and make overall_feedback at least 180 words, concrete, and explicitly tied to the section analyses and section labels. "
            "Return only the required JSON object."
        )
    elif "did not explicitly cover all detected sections" in message:
        retry_note = (
            "Previous attempt failed because not all detected sections were explicitly covered. "
            "Retry now and mention every detected section label in overall_feedback and covered_parts. "
            "Return only the required JSON object."
        )
    else:
        retry_note = (
            f"Previous attempt failed validation: {message}. "
            "Retry now and satisfy the validation rule. Return only the required JSON object."
        )
    return [*messages, {"role": "user", "content": retry_note}]


def _build_part_analysis_retry_messages(
    messages: list[dict[str, str]],
    part: SubmissionPart,
    exc: ValueError,
) -> list[dict[str, str]]:
    scale_labels = ", ".join(_criterion_scale_for_part(part)[1])
    retry_note = (
        f"Previous attempt for {part.label} failed validation: {exc}. "
        "Retry now and return one short criterion_notes item per criterion, at least 2 concrete strengths, at least 2 concrete weaknesses, "
        f"at least 1 concrete evidence item, and a non-empty coverage_comment. Use criterion_notes.status values from this scale only: {scale_labels}. "
        "Return only the required JSON object."
    )
    return [*messages, {"role": "user", "content": retry_note}]


def _normalize_feedback_text(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if not isinstance(value, dict):
        return ""
    content = value.get("content")
    if not isinstance(content, list):
        return ""
    lines: list[str] = []
    for item in content:
        if not isinstance(item, dict):
            continue
        label = str(item.get("section_label", "")).strip()
        analysis = str(item.get("analysis", "")).strip()
        feedback = str(item.get("feedback", "")).strip()
        parts = [piece for piece in (label, analysis, feedback) if piece]
        if parts:
            lines.append(": ".join(parts[:1]) + (f" - {' '.join(parts[1:])}" if len(parts) > 1 else ""))
    return " ".join(lines).strip()


def _normalize_string_list(
    value: Any,
    field_name: str,
    label: str,
    minimum: int,
    pad_missing: bool = False,
) -> list[str]:
    if not isinstance(value, list):
        raise ValueError(f"Model did not return a list for {field_name} in {label}: {value}")
    normalized = [str(item).strip() for item in value if str(item).strip()]
    if len(normalized) < minimum:
        if not pad_missing:
            raise ValueError(f"Model returned too few items for {field_name} in {label}: {value}")
        # Keep part-level scoring resilient when the model returns undersized bullet lists.
        while len(normalized) < minimum:
            if field_name == "strengths":
                normalized.append("Additional concrete strength was not provided in the model response.")
            elif field_name == "weaknesses":
                normalized.append("Additional concrete weakness was not provided in the model response.")
            elif field_name == "evidence":
                normalized.append("Concrete supporting evidence was not provided in the model response.")
            else:
                normalized.append("Additional required item was not provided in the model response.")
    return normalized


def _normalize_criterion_notes(
    value: Any,
    expected_criteria: list[dict[str, str]],
    expected_label: str,
) -> list[dict[str, str]]:
    expected_names = [str(item.get("criterion_name", "")).strip() for item in expected_criteria]
    if not all(expected_names):
        raise ValueError(f"Expected criteria were not fully defined for {expected_label}.")
    expected_scale_type = str(expected_criteria[0].get("scale_type", "")).strip() if expected_criteria else ""
    expected_scale_labels = tuple(
        label.strip()
        for label in str(expected_criteria[0].get("scale_labels", "")).split(",")
        if label.strip()
    ) if expected_criteria else ()
    if isinstance(value, dict):
        normalized_items: list[dict[str, Any]] = []
        for criterion_name, item in value.items():
            if not isinstance(item, dict):
                raise ValueError(f"Model returned invalid criterion_notes entries for {expected_label}: {value}")
            normalized_items.append(
                {
                    "criterion_name": str(criterion_name).strip(),
                    "status": item.get("status"),
                    "note": item.get("note"),
                }
            )
        value = normalized_items
    if not isinstance(value, list):
        raise ValueError(f"Model did not return a list for criterion_notes in {expected_label}: {value}")
    by_name: dict[str, dict[str, str]] = {}
    for item in value:
        if not isinstance(item, dict):
            raise ValueError(f"Model returned invalid criterion_notes entries for {expected_label}: {value}")
        criterion_name = str(item.get("criterion_name", "")).strip()
        if criterion_name not in expected_names or criterion_name in by_name:
            raise ValueError(f"Model returned invalid criterion_notes names for {expected_label}: {value}")
        status, normalized_status = _normalize_criterion_note_status(item.get("status"), expected_scale_type)
        if expected_scale_labels and status not in expected_scale_labels:
            raise ValueError(f"Model returned invalid criterion_notes status for {expected_label}: {value}")
        note = _normalize_criterion_note_text(item.get("note"))
        if not note:
            raise ValueError(f"Model returned empty criterion_notes note for {expected_label}: {value}")
        by_name[criterion_name] = {
            "criterion_name": criterion_name,
            "status": status,
            "normalized_status": normalized_status,
            "scale_type": expected_scale_type,
            "scale_labels": list(expected_scale_labels),
            "note": note,
        }
    if len(by_name) != len(expected_names):
        # M2 hardening: instead of crashing on incomplete criterion_notes,
        # fill missing criteria with conservative defaults ("weak" for likert, "missing_or_wrong" for scale)
        missing_status, missing_normalized = _normalize_criterion_note_status(
            "weak" if expected_scale_type == "likert_judgement" else "missing_or_wrong",
            expected_scale_type,
        )
        for criterion_name in expected_names:
            if criterion_name in by_name:
                continue
            by_name[criterion_name] = {
                "criterion_name": criterion_name,
                "status": missing_status,
                "normalized_status": missing_normalized,
                "scale_type": expected_scale_type,
                "scale_labels": list(expected_scale_labels),
                "note": "omitted by model",
            }
    return [by_name[name] for name in expected_names]


def _normalize_criterion_note_status(value: Any, scale_type: str) -> tuple[str, str]:
    text = str(value).strip().lower()
    if scale_type == "likert_judgement":
        native_mapping = {
            "weak": "weak",
            "developing": "developing",
            "clear": "clear",
            "strong": "strong",
            "missed": "weak",
            "met": "clear",
            "meets": "clear",
            "partial": "developing",
            "partly": "developing",
            "mixed": "developing",
            "secure": "strong",
            "strong_anchor": "strong",
            "weak_anchor": "weak",
            "absent": "weak",
        }
        normalized_mapping = {
            "weak": "missed",
            "developing": "partial",
            "clear": "met",
            "strong": "met",
        }
        native_status = native_mapping.get(text, text)
        return native_status, normalized_mapping.get(native_status, native_status)
    native_mapping = {
        "missing_or_wrong": "missing_or_wrong",
        "partial": "partial",
        "partly": "partial",
        "mixed": "partial",
        "secure": "secure",
        "missed": "missing_or_wrong",
        "met": "secure",
        "meets": "secure",
        "clear": "secure",
        "strong": "secure",
        "strong_anchor": "secure",
        "weak": "missing_or_wrong",
        "weak_anchor": "missing_or_wrong",
        "absent": "missing_or_wrong",
    }
    normalized_mapping = {
        "missing_or_wrong": "missed",
        "partial": "partial",
        "secure": "met",
    }
    native_status = native_mapping.get(text, text)
    return native_status, normalized_mapping.get(native_status, native_status)


def _normalize_criterion_note_text(value: Any) -> str:
    text = " ".join(str(value).strip().split())
    if not text:
        return ""
    text = re.sub(r"[.;:,]+$", "", text)
    words = text.split()
    if len(words) > 5:
        words = words[:5]
    normalized = " ".join(words).strip()
    return normalized[:48].strip()


def _has_major_weaknesses(weaknesses: list[str], feedback: str) -> bool:
    text = " ".join(weaknesses + [feedback]).lower()
    patterns = (
        "incorrect",
        "missing",
        "does not address",
        "fails to",
        "not properly specified",
        "superficial",
        "underdeveloped",
        "limited discussion",
    )
    return any(pattern in text for pattern in patterns)


def _is_placeholder_section_label(label: str) -> bool:
    normalized = _normalize_label_key(label)
    placeholders = {
        "section above",
        "section below",
        "answer above",
        "answer below",
        "above",
        "below",
    }
    return normalized in placeholders


def _extract_text_from_docx_xml(raw: bytes) -> str:
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    }
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            xml_bytes = archive.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile):
        return ""

    root = ET.fromstring(xml_bytes)
    paragraphs: list[str] = []
    for paragraph in root.findall(".//w:p", namespaces):
        text_parts: list[str] = []
        for node in paragraph.iter():
            tag = node.tag.rsplit("}", 1)[-1]
            if tag in {"t", "delText"} and node.text:
                text_parts.append(node.text)
            elif tag == "tab":
                text_parts.append("\t")
        paragraph_text = "".join(text_parts).strip()
        if paragraph_text:
            paragraphs.append(paragraph_text)
    return "\n".join(paragraphs)


def _restore_label_from_key(key: str) -> str:
    parts = key.split()
    if len(parts) >= 2:
        return f"{parts[0].title()} {parts[1].upper() if parts[1].isalpha() else parts[1]}"
    return key.title()


def _segment_subparts_within_section(parent_part: SubmissionPart, child_specs: list[SubmissionPart]) -> list[SubmissionPart] | None:
    section_text = parent_part.section_text.strip()
    if not section_text:
        return None

    anchors: list[tuple[int, SubmissionPart]] = []
    used_positions: set[int] = set()
    for child in child_specs:
        position = _find_subpart_anchor_position(section_text, child)
        if position == -1 or position in used_positions:
            continue
        used_positions.add(position)
        anchors.append((position, child))

    if len(anchors) < max(2, len(child_specs) // 2 + 1):
        return None

    anchors.sort(key=lambda item: item[0])
    segmented: list[SubmissionPart] = []
    for index, (start, child) in enumerate(anchors):
        end = anchors[index + 1][0] if index + 1 < len(anchors) else len(section_text)
        child_text = _select_local_chunk_with_fallback(
            section_text,
            start,
            end,
            child.question_text_exact or child.focus_hint,
        )
        segmented.append(
            SubmissionPart(
                label=child.label,
                focus_hint=child.focus_hint,
                anchor_text=child.anchor_text,
                section_text=child_text,
                max_mark=child.max_mark,
                marking_guidance=child.marking_guidance,
                question_text_exact=child.question_text_exact,
                task_type=child.task_type,
            )
        )
    return segmented


def _resolve_expected_child_parts(
    parent_part: SubmissionPart,
    child_parts: list[SubmissionPart],
    child_specs: list[SubmissionPart],
) -> list[SubmissionPart]:
    child_by_key = {_normalize_label_key(child.label): child for child in child_parts}
    parent_section_text = parent_part.section_text
    child_anchor_positions = {
        _normalize_label_key(child.label): _find_subpart_anchor_position(parent_section_text, child)
        for child in child_specs
    }
    resolved: list[SubmissionPart] = []
    for expected in child_specs:
        matched = child_by_key.get(_normalize_label_key(expected.label))
        section_text = matched.section_text if matched is not None else parent_section_text
        if (
            matched is not None
            and parent_section_text.strip()
            and _section_text_looks_damaged_or_incomplete(section_text, expected.question_text_exact or expected.focus_hint)
        ):
            fallback_start = child_anchor_positions.get(_normalize_label_key(expected.label), -1)
            if fallback_start != -1:
                following_positions = sorted(
                    position
                    for child in child_specs
                    if _normalize_label_key(child.label) != _normalize_label_key(expected.label)
                    for position in [child_anchor_positions.get(_normalize_label_key(child.label), -1)]
                    if position > fallback_start
                )
                fallback_end = following_positions[0] if following_positions else len(parent_section_text)
                section_text = _select_local_chunk_with_fallback(
                    parent_section_text,
                    fallback_start,
                    fallback_end,
                    expected.question_text_exact or expected.focus_hint,
                )
        if parent_section_text.strip() and expected.task_type == "critique_and_revision":
            current_start = child_anchor_positions.get(_normalize_label_key(expected.label), -1)
            if current_start != -1:
                following_positions = sorted(
                    position
                    for child in child_specs
                    if _normalize_label_key(child.label) != _normalize_label_key(expected.label)
                    for position in [child_anchor_positions.get(_normalize_label_key(child.label), -1)]
                    if position > current_start
                )
                cumulative_end = following_positions[0] if following_positions else len(parent_section_text)
                cumulative_text = _select_local_chunk_with_fallback(
                    parent_section_text,
                    0,
                    cumulative_end,
                    expected.question_text_exact or expected.focus_hint,
                )
                if cumulative_text:
                    section_text = cumulative_text
        resolved.append(
            SubmissionPart(
                label=expected.label,
                focus_hint=(matched.focus_hint if matched is not None else "") or expected.focus_hint,
                anchor_text=(matched.anchor_text if matched is not None else "") or expected.anchor_text,
                section_text=section_text,
                max_mark=expected.max_mark,
                marking_guidance=expected.marking_guidance,
                question_text_exact=expected.question_text_exact,
                task_type=expected.task_type,
                criterion_mode=expected.criterion_mode,
            )
        )
    return resolved


def _find_subpart_anchor_position(section_text: str, child: SubmissionPart) -> int:
    normalized_text = section_text.replace("\r\n", "\n")
    question_match = re.search(r"\bQ(?P<id>\d+)\b", child.label, flags=re.IGNORECASE)
    if not question_match:
        return -1
    question_id = question_match.group("id")
    patterns = [
        rf"(?im)^\s*Q{question_id}\b",
        rf"(?im)^\s*Question\s+{question_id}\b",
        rf"(?im)^\s*{question_id}\.\s",
        rf"(?im)^\s*{question_id}\)\s",
        rf"(?im)^\s*\({question_id}\)\s",
    ]
    for pattern in patterns:
        match = re.search(pattern, normalized_text)
        if match:
            return match.start()
    return -1


def _detect_subparts_with_model(
    parent_part: SubmissionPart,
    filename: str,
    child_specs: list[SubmissionPart],
    model_name: str,
    ollama_url: str,
) -> list[SubmissionPart] | None:
    # Refuse to call the model on very short sections — there is not enough
    # content to reliably split into subparts, and the call would waste time
    # or assign wrong text under the wrong label.
    _MIN_WORDS_PER_CHILD = 10
    word_count = len(parent_part.section_text.split())
    if word_count < len(child_specs) * _MIN_WORDS_PER_CHILD:
        return None

    data = _call_structure_detector_with_fallback(
        primary_model_name=model_name,
        messages=build_subpart_structure_messages(parent_part, filename, child_specs),
        ollama_url=ollama_url,
    )
    detected = normalize_detected_parts(data)
    if len(detected) == 1 and _normalize_label_key(detected[0].label) == _normalize_label_key(parent_part.label):
        return None

    child_by_key = {_normalize_label_key(child.label): child for child in child_specs}
    normalized_children: list[SubmissionPart] = []
    for child in detected:
        key = _normalize_label_key(child.label)
        expected = child_by_key.get(key)
        if expected is None:
            continue
        normalized_children.append(
            SubmissionPart(
                label=expected.label,
                focus_hint=child.focus_hint or expected.focus_hint,
                anchor_text=child.anchor_text or expected.anchor_text,
                max_mark=expected.max_mark,
                marking_guidance=expected.marking_guidance,
                question_text_exact=expected.question_text_exact,
                task_type=expected.task_type,
            )
        )
    if len(normalized_children) < 2:
        return None
    return _segment_subparts_within_section(parent_part, normalized_children)


def _list_supported_files_recursive(folder: Path) -> list[Path]:
    return sorted(path for path in folder.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_TEXT_SUFFIXES)


def _pick_matching_files(
    files: list[Path],
    keywords: tuple[str, ...],
    exclude: tuple[Path, ...] = (),
) -> tuple[Path, ...]:
    excluded = set(exclude)
    if not keywords:
        return ()
    matches = []
    for path in files:
        if path in excluded:
            continue
        haystack = f"{path.stem} {path.parent.name}".lower().replace("_", " ")
        if any(keyword in haystack for keyword in keywords):
            matches.append(path)
    return tuple(matches)


def _normalize_label_key(label: str) -> str:
    return " ".join(label.lower().split())


_TOC_MIN_LABELS = 3
_TOC_MAX_FRACTION = 0.18
# A TOC entry line ends with 2+ dotted leaders (or dashes/em-dashes) followed by
# an optional space and a 1-3 digit page number, with no further non-space text.
_TOC_LINE_RE = re.compile(r"(?m)^\s*\S[^\n]*?[.\u2026\u2014\u2013-]{2,}\s*\d{1,3}\s*$")


def _detect_toc_block_range(
    normalized_text: str,
    parts: list["SubmissionPart"],
) -> tuple[int, int] | None:
    """Return (start, end) character offsets of a probable table-of-contents block.

    Identifies TOC *entry* lines — lines whose content matches a known part label
    AND that end with dotted leaders + a page number (e.g. "Part 1 ......... 3").
    If _TOC_MIN_LABELS or more such lines appear in the first _TOC_MAX_FRACTION
    of the document, returns a tight range spanning just those lines (+ a small
    buffer), so that real body headers immediately following the TOC are NOT
    included in the exclusion zone.
    Returns None if no TOC cluster is found.
    """
    if not parts or not normalized_text:
        return None
    scan_limit = max(0, int(len(normalized_text) * _TOC_MAX_FRACTION))
    labels = [p.label.strip() for p in parts if p.label.strip()]
    if len(labels) < _TOC_MIN_LABELS:
        return None

    toc_matches = [
        m
        for m in _TOC_LINE_RE.finditer(normalized_text[:scan_limit])
        if any(
            re.search(rf"\b{re.escape(label)}\b", m.group(0), flags=re.IGNORECASE)
            for label in labels
        )
    ]
    if len(toc_matches) >= _TOC_MIN_LABELS:
        toc_start = toc_matches[0].start()
        toc_end = toc_matches[-1].end()
        return toc_start, toc_end
    return None


def _find_part_anchor_position(
    normalized_text: str,
    part: "SubmissionPart",
    toc_range: tuple[int, int] | None = None,
) -> int:
    """Return the character offset of the best anchor for *part*.

    Priority (descending):
    1. Line-start / header-style match **outside** the TOC block.
    2. Word-boundary match outside the TOC.
    3. Line-start / header-style match inside the TOC (last resort).
    4. Any word-boundary match inside the TOC.
    """
    toc_start, toc_end = toc_range if toc_range else (-1, -1)

    def _in_toc(pos: int) -> bool:
        return toc_start != -1 and toc_start <= pos < toc_end

    header_outside: int = -1
    word_outside: int = -1
    header_inside: int = -1
    word_inside: int = -1

    candidates = [part.anchor_text.strip(), part.label.strip()]
    for candidate in candidates:
        if not candidate:
            continue
        header_pattern = re.compile(
            rf"(?im)^\s*{re.escape(candidate)}(?:\b|\s*[:.)\]-])",
            flags=re.IGNORECASE,
        )
        for m in header_pattern.finditer(normalized_text):
            pos = m.start()
            if not _in_toc(pos):
                if header_outside == -1:
                    header_outside = pos
            else:
                if header_inside == -1:
                    header_inside = pos

        word_pattern = re.compile(rf"\b{re.escape(candidate)}\b", flags=re.IGNORECASE)
        for m in word_pattern.finditer(normalized_text):
            pos = m.start()
            if not _in_toc(pos):
                if word_outside == -1:
                    word_outside = pos
            else:
                if word_inside == -1:
                    word_inside = pos

    for pos in (header_outside, word_outside, header_inside, word_inside):
        if pos != -1:
            return pos
    return -1
